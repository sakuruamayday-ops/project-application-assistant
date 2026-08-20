#!/usr/bin/env python3
"""Fill a governed dual-report DOCX master from a private evidence fixture.

The fixture stays outside the public skill tree.  This module validates real
source anchors, fills the copied Word master, adds a source ledger, and refuses
to leave training placeholders in a completed draft.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, RGBColor


PLACEHOLDER = re.compile(r"［[^］\r\n]{1,120}］")
TRAINING_MARKERS = ("培训模板", "培训说明", "灰色提示文字")
INTERNAL_DRAFT_MARKERS = ("真实客户资料候选验收稿", "候选验收稿")
REPORT_TYPES = {"preassessment", "feasibility"}
CONCLUSION_ALIASES = {
    "待资料": "暂无法判断",
    "培育后申报": "有条件申报",
    "建议申报": "可申报",
}
CONCLUSIONS = {"可申报", "有条件申报", "不可申报", "暂无法判断"}
EVIDENCE_STATES = {
    "已具备",
    "企业提供待核",
    "待企业确认",
    "第三方平台查询，待企业确认",
    "当前检索层未命中",
    "明确未具备",
    "存在差距",
    "冲突待核",
    "不适用",
}
PORTABLE_CJK_FONT = "Noto Sans SC"
PUBLIC_DIGEST = re.compile(r"(?:SHA[\s_-]*256|(?<![0-9a-f])[0-9a-f]{64}(?![0-9a-f]))", re.IGNORECASE)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", "", value or "")


def extract_source_text(path: Path) -> str:
    suffix = path.suffix.casefold()
    if suffix == ".docx":
        document = Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        parts.extend(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        return "\n".join(parts)
    if suffix == ".pdf":
        import fitz

        document = fitz.open(path)
        try:
            return "\n".join(page.get_text("text") for page in document)
        finally:
            document.close()
    if suffix in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        try:
            return "\n".join(
                str(value)
                for sheet in workbook.worksheets
                for row in sheet.iter_rows(values_only=True)
                for value in row
                if value not in (None, "")
            )
        finally:
            workbook.close()
    if suffix in {".txt", ".md", ".csv", ".json", ".jsonl"}:
        return path.read_text(encoding="utf-8")
    raise ValueError(f"不支持的客户资料格式:{path.name}")


def _public_source_reference(material: dict[str, Any]) -> dict[str, str | None]:
    source_type = str(material.get("source_type") or "").strip().casefold()
    raw = str(
        material.get("url")
        or material.get("official_url")
        or material.get("link")
        or material.get("source")
        or ""
    ).strip()
    if source_type in {"knowledge-base", "knowledge_base", "gongchuang-knowledge"} or raw == "共创知识库":
        return {"label": "来源共创知识库", "url": None}
    if raw.startswith(("https://", "http://")):
        return {"label": raw, "url": raw}
    return {"label": raw or "用户提供", "url": None}


def validate_case_fixture(
    fixture: dict[str, Any],
    *,
    expected_project_id: str | None = None,
    public_root: Path | None = None,
) -> dict[str, Any]:
    required = {
        "project_id",
        "enterprise",
        "project_object",
        "suggested_year",
        "deadline",
        "conclusion",
        "conclusion_basis",
        "primary_gap",
        "next_action",
        "materials",
        "policies",
    }
    missing = sorted(required - set(fixture))
    if missing:
        raise ValueError("客户夹具缺字段:" + ",".join(missing))
    project_id = str(fixture["project_id"]).strip()
    if expected_project_id and project_id != expected_project_id:
        raise ValueError(f"项目路由与夹具不一致:{expected_project_id}/{project_id}")
    if not str(fixture["enterprise"]).strip():
        raise ValueError("企业名称不得为空")
    materials = fixture.get("materials")
    if not isinstance(materials, list) or not materials:
        raise ValueError("每类项目至少需要一份真实客户资料")
    normalized_sources: list[dict[str, Any]] = []
    for material in materials:
        if not isinstance(material, dict):
            raise ValueError("客户资料记录格式错误")
        path = Path(str(material.get("path") or "")).expanduser().resolve()
        if not path.is_file():
            raise FileNotFoundError(path)
        if public_root is not None:
            try:
                path.relative_to(public_root.resolve())
            except ValueError:
                pass
            else:
                raise ValueError("真实客户资料不得放入公共技能树")
        anchors = material.get("anchors")
        if not isinstance(anchors, list) or not anchors:
            raise ValueError(f"客户资料缺原文锚点:{path.name}")
        source_text = normalize_text(extract_source_text(path))
        checked_anchors: list[str] = []
        for anchor in anchors:
            anchor_text = str(anchor).strip()
            if len(normalize_text(anchor_text)) < 4:
                raise ValueError(f"原文锚点过短:{path.name}")
            if normalize_text(anchor_text) not in source_text:
                raise ValueError(f"原文锚点未命中:{path.name}:{anchor_text}")
            checked_anchors.append(anchor_text)
        normalized_sources.append(
            {
                "path": path,
                "name": path.name,
                "sha256": sha256_file(path),
                "anchors": checked_anchors,
                "role": str(material.get("role") or "企业资料"),
                "public_source": _public_source_reference(material),
            }
        )
    policies = fixture.get("policies")
    if not isinstance(policies, list) or not policies:
        raise ValueError("夹具必须至少锁定一项政策基线")
    for policy in policies:
        if not isinstance(policy, dict) or not str(policy.get("title") or "").strip():
            raise ValueError("政策基线缺文件名称")
        if not str(policy.get("locator") or "").strip():
            raise ValueError("政策基线缺原文位置或官方链接")
    for condition in fixture.get("conditions", []):
        state = str(condition.get("state") or "待企业确认")
        if state not in EVIDENCE_STATES:
            raise ValueError(f"不受控的证据状态:{state}")
    conclusion = CONCLUSION_ALIASES.get(str(fixture["conclusion"]).strip(), str(fixture["conclusion"]).strip())
    if conclusion not in CONCLUSIONS:
        raise ValueError("前期评估结论必须为可申报、有条件申报、不可申报或暂无法判断")
    return {**fixture, "conclusion": conclusion, "_validated_sources": normalized_sources}


def _set_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _apply_portable_cjk_font(document: Document) -> None:
    """Bind visible text to a macOS CJK font before PDF visual regression."""
    for paragraph in _iter_paragraphs(document):
        for run in paragraph.runs:
            if not run.text:
                continue
            run.font.name = PORTABLE_CJK_FONT
            run_properties = run._element.get_or_add_rPr()
            fonts = run_properties.get_or_add_rFonts()
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{attribute}"), PORTABLE_CJK_FONT)


def _iter_paragraphs(document: Document) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (section.header, section.footer):
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _replace_training_text(value: str, fixture: dict[str, Any], release_tag: str) -> str:
    replacements = {
        f"{release_tag} 培训模板 | 共创研究院": "共创研究院",
        "V1.6.5.1 培训模板 | 共创研究院": "共创研究院",
        f"{release_tag} 培训模板": "",
        "V1.6.5.1 培训模板": "",
        "企业名称：［填写］": f"企业名称：{fixture['enterprise']}",
        "报告日期：［填写］": f"报告日期：{fixture.get('report_date') or date.today().isoformat()}",
        "顾问：［填写］": f"顾问：{fixture.get('advisor') or '共创研究院'}",
        "灰色提示文字为培训说明，正式交付前应替换或删除。": "本稿已按真实客户资料回填，未取得的企业数据统一标注待企业确认。",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)
    version_pattern = rf"(?:{re.escape(release_tag)}|V\d+(?:\.\d+){{2,3}})"
    value = re.sub(
        rf"{version_pattern}\s*培训模板\s*\|\s*共创研究院",
        "共创研究院",
        value,
    )
    value = re.sub(rf"{version_pattern}\s*培训模板", "", value)
    return value


def _fallback_for_placeholder(token: str, context: str, fixture: dict[str, Any]) -> str:
    hint = token.strip("［］")
    combined = context + hint
    primary_source = fixture["_validated_sources"][0]["name"]
    if any(word in combined for word in ("截止", "官方时间")):
        return str(fixture["deadline"])
    if any(word in combined for word in ("申报年度", "完成年度")):
        return str(fixture["suggested_year"])
    if any(word in combined for word in ("产品", "申报对象")):
        return str(fixture["project_object"])
    if any(word in combined for word in ("文件、页码", "文件或台账", "证据", "来源")):
        return primary_source
    if any(word in combined for word in ("责任人", "核验人", "复核人")):
        return str(fixture.get("owner") or "企业项目负责人")
    if any(word in combined for word in ("完成节点", "建议时间", "年度/季度", "截止日")):
        return str(fixture.get("target_period") or f"{fixture['suggested_year']}年申报前")
    if any(word in combined for word in ("具体动作", "补强动作", "后续动作", "下一步")):
        return str(fixture["next_action"])
    if any(word in combined for word in ("差距", "短板")):
        return str(fixture["primary_gap"])
    if any(word in combined for word in ("评分", "预计得分", "确定分", "条件分", "满分")):
        return "待锁定当期评分表"
    if any(word in combined for word in ("数据", "企业值", "现状", "已知")):
        return str(fixture.get("headline_fact") or "当前资料未提供可复算数据")
    if any(word in combined for word in ("状态", "已具备/待")):
        return "待企业确认"
    if any(word in combined for word in ("口径", "统计期间")):
        return "以申报年度通知及企业专项底稿为准"
    if any(word in combined for word in ("交付物", "完成标准", "可核验成果")):
        return "形成可复核台账、报告及原始证据索引"
    if any(word in combined for word in ("核心原因", "依据", "限制")):
        return str(fixture["conclusion_basis"])
    return "待企业确认"


def _condition_match(fixture: dict[str, Any], condition_text: str) -> dict[str, Any] | None:
    normalized = normalize_text(condition_text)
    matches: list[tuple[int, dict[str, Any]]] = []
    for item in fixture.get("conditions", []):
        key = normalize_text(str(item.get("match") or ""))
        if key and (key in normalized or normalized in key):
            matches.append((len(key), item))
    return max(matches, default=(0, None), key=lambda pair: pair[0])[1]


def _fill_policy_table(table, fixture: dict[str, Any], feasibility: bool) -> None:
    policies = fixture["policies"]
    if len(table.rows[0].cells) >= 3:
        _set_text(table.rows[0].cells[2].paragraphs[0], "官方来源")

    def categories(policy: dict[str, Any]) -> set[str]:
        declared = str(policy.get("kind") or "").strip().casefold()
        result = {declared} if declared else set()
        title = str(policy.get("title") or "")
        if re.search(r"管理办法|评价办法|实施办法|工作指引|若干措施|实施细则", title):
            result.add("management")
        if re.search(r"组织申报|申报通知|年度.*通知", title):
            result.update({"annual-notice", "deadline"})
        if re.search(r"申报指南|评分附件|评分细则|榜单|指南", title):
            result.add("guidance")
        return result

    row_kinds = {
        "管理办法或评价办法": "management",
        "申报年度通知": "annual-notice",
        "申报指南与评分附件": "guidance",
        "截止日期": "deadline",
    }
    for row in table.rows[1:]:
        row_name = row.cells[0].text.strip()
        expected = row_kinds.get(row_name)
        policy = next((item for item in policies if expected in categories(item)), None)
        cells = row.cells
        if len(cells) == 4:
            if policy is None:
                cells[1].text = "当前未取得独立文件"
                cells[2].text = "—"
                cells[3].text = "待核验"
                continue
            cells[1].text = str(policy["title"])
            locator = str(policy["locator"])
            if locator.startswith(("https://", "http://")):
                _set_cell_text(cells[2], "")
                _add_hyperlink(cells[2].paragraphs[0], "查看官方原文", locator)
            else:
                cells[2].text = locator
            cells[3].text = str(policy.get("status") or "已锁定基线")
        elif len(cells) >= 5:
            if policy is None:
                cells[1].text = "当前未取得独立文件"
                cells[2].text = "—"
                cells[3].text = "—"
                cells[4].text = "待核验"
                continue
            cells[1].text = str(policy["title"])
            cells[2].text = str(policy.get("scope") or "以通知和附件适用范围为准")
            locator = str(policy["locator"])
            if locator.startswith(("https://", "http://")):
                _set_cell_text(cells[3], "")
                _add_hyperlink(cells[3].paragraphs[0], "查看官方原文", locator)
            else:
                cells[3].text = locator
            cells[4].text = str(policy.get("status") or "已锁定基线")


def _fill_summary_table(table, fixture: dict[str, Any], feasibility: bool) -> None:
    for row in table.rows[1:]:
        key = row.cells[0].text.strip()
        if "当前可行性" in key:
            row.cells[1].text = str(fixture["conclusion"])
            row.cells[2].text = str(fixture["conclusion_basis"])
        elif "申报结论" in key:
            row.cells[1].text = str(fixture["conclusion"])
            row.cells[2].text = str(fixture["conclusion_basis"])
        elif "建议申报年度" in key:
            row.cells[1].text = str(fixture["suggested_year"])
            row.cells[2].text = str(fixture.get("year_basis") or "结合当期窗口、证据完整度和补强周期")
        elif "申报层级" in key or "档次" in key:
            row.cells[1].text = str(fixture.get("level") or "待企业确认适用层级")
            row.cells[2].text = str(fixture.get("level_basis") or "以项目属地和当期申报表为准")
        elif "首要短板" in key:
            row.cells[1].text = str(fixture["primary_gap"])
            row.cells[2].text = str(fixture["next_action"])
        elif "预计满足硬门槛" in key:
            row.cells[1].text = str(fixture.get("gate_summary") or "待企业补齐数据后复算")
            row.cells[2].text = "按当期通知原文逐条确认"
        elif "预计评分" in key:
            row.cells[1].text = str(fixture.get("score") or "不估分")
            row.cells[2].text = str(fixture.get("score_basis") or "当期完整评分表或企业数据未闭合")
        elif "关键前置任务" in key:
            row.cells[1].text = str(fixture["next_action"])
            row.cells[2].text = str(fixture.get("target_period") or f"{fixture['suggested_year']}年申报前")


def _fill_condition_table(table, fixture: dict[str, Any], feasibility: bool) -> None:
    primary_source = fixture["_validated_sources"][0]["name"]
    for row in table.rows[1:]:
        cells = row.cells
        condition = cells[0].text.strip()
        match = _condition_match(fixture, condition) or {}
        value = str(match.get("value") or "当前资料未提供可复算数据")
        state = str(match.get("state") or "待企业确认")
        source = str(match.get("source") or primary_source)
        if state == "企业提供待核" and re.search(r"(?:企查查|天眼查|\bQCC\b|\bTYC\b)", source, re.IGNORECASE):
            state = "第三方平台查询，待企业确认"
        gap = str(match.get("gap") or "数据和证据尚未闭合")
        action = str(match.get("action") or fixture["next_action"])
        if len(cells) == 5:
            cells[1].text = value
            cells[2].text = state
            cells[3].text = gap
            cells[4].text = action
        elif len(cells) >= 7:
            cells[1].text = value
            cells[2].text = source
            cells[3].text = state
            cells[4].text = str(match.get("fixed_score") or "不计分")
            cells[5].text = str(match.get("conditional_score") or "不计分")
            cells[6].text = gap


def _fill_path_table(table, fixture: dict[str, Any]) -> None:
    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) >= 5:
            cells[1].text = str(fixture["suggested_year"])
            cells[2].text = str(fixture["deadline"])
            cells[3].text = str(fixture.get("window_status") or "按报告基准日核验")
            cells[4].text = str(fixture["next_action"])


def _fill_strengthening_table(table, fixture: dict[str, Any]) -> None:
    defaults = {
        "国内产品技术水平评价咨询报告": (
            "尚未形成国内技术水平评价报告",
            "国内产品技术水平评价咨询报告",
        ),
        "国际产品技术水平评价咨询报告": (
            "尚未形成国际技术水平评价报告",
            "国际产品技术水平评价咨询报告",
        ),
        "技术查新或科技咨询": (
            "尚未提供技术查新或科技咨询成果",
            "技术查新报告或科技咨询意见",
        ),
        "任务指标检测方案和项目预算底稿": (
            "任务指标与项目预算底稿尚未闭合",
            "任务指标对照表、检测方案和项目预算底稿",
        ),
    }
    configured = [
        item
        for item in fixture.get("strengthening_tasks", [])
        if isinstance(item, dict) and str(item.get("task") or "").strip()
    ]
    if configured:
        template_row = table.rows[-1]._tr
        while len(table.rows) - 1 < len(configured):
            table._tbl.append(deepcopy(template_row))
        while len(table.rows) - 1 > len(configured):
            table._tbl.remove(table.rows[-1]._tr)
    suggested_year = str(fixture["suggested_year"])
    default_period = f"{suggested_year}申报前" if suggested_year.endswith("年") else f"{suggested_year}年申报前"
    for index, row in enumerate(table.rows[1:], start=1):
        cells = row.cells
        task = cells[1].text.strip()
        item = configured[index - 1] if configured else {}
        default_status, default_deliverable = defaults.get(
            task,
            ("尚未形成该项成果", "形成可核验的成果文件"),
        )
        cells[0].text = str(index)
        cells[1].text = str(item.get("task") or task)
        cells[2].text = str(item.get("status") or default_status)
        cells[3].text = str(item.get("target_period") or fixture.get("target_period") or default_period)
        cells[4].text = str(item.get("deliverable") or default_deliverable)


def _lock_table_pagination(table) -> None:
    if not table.rows:
        return
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        properties.append(cant_split)


def _fill_project_specific_paragraphs(document: Document, fixture: dict[str, Any]) -> None:
    replacements = {
        "申报主体：": f"申报主体：{fixture['enterprise']}",
        "主导产品或申报对象：": f"主导产品或申报对象：{fixture['project_object']}",
        "建议口径：": f"建议口径：{fixture.get('recommended_scope') or fixture['project_object']}",
    }
    for paragraph in document.paragraphs:
        for prefix, replacement in replacements.items():
            if paragraph.text.strip().startswith(prefix):
                _set_text(paragraph, replacement)


def _fill_table_by_header(table, fixture: dict[str, Any], feasibility: bool) -> None:
    if not table.rows:
        return
    header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
    if "政策层级" in header and ("现行文件" in header or "文件名称与文号" in header):
        _fill_policy_table(table, fixture, feasibility)
    elif "判断项" in header and "一句话依据" in header:
        _fill_summary_table(table, fixture, feasibility)
    elif "事项" in header and "依据及限制" in header:
        _fill_summary_table(table, fixture, feasibility)
    elif "现行条件 | 企业值" in header:
        _fill_condition_table(table, fixture, feasibility)
    elif "现行条件或评分项" in header:
        _fill_condition_table(table, fixture, feasibility)
    elif "项目 | 申报年度 | 申报截止日期" in header:
        _fill_path_table(table, fixture)
    elif "补强任务" in header and "成果证据" in header:
        _fill_strengthening_table(table, fixture)


def _generic_fill(document: Document, fixture: dict[str, Any]) -> None:
    source_name = fixture["_validated_sources"][0]["name"]
    for paragraph in _iter_paragraphs(document):
        if paragraph._p.findall(".//" + qn("w:fldChar")) or paragraph._p.findall(".//" + qn("w:instrText")):
            for run in paragraph.runs:
                value = _replace_training_text(
                    run.text,
                    fixture,
                    str(fixture.get("release_tag") or "V1.6.5.2"),
                )
                if value != run.text:
                    run.text = value
            continue
        original = paragraph.text
        value = _replace_training_text(original, fixture, str(fixture.get("release_tag") or "V1.6.5.2"))
        if "□可申报" in value:
            value = str(fixture["conclusion"])
        if "□建议申报" in value:
            value = str(fixture["conclusion"])
        value = value.replace("□已核验 □待补", "☐已核验 ☒待补")
        value = value.replace("□已有 □待确认 □待补强", "☐已有 ☒待确认 ☐待补强")
        value = value.replace("□通过 □不通过", "☐通过 ☒不通过")
        while True:
            match = PLACEHOLDER.search(value)
            if not match:
                break
            value = value[: match.start()] + _fallback_for_placeholder(match.group(), value, fixture) + value[match.end() :]
        value = value.replace("待核验", "待企业确认") if "政策" not in value else value
        if value != original:
            _set_text(paragraph, value)
        if "企业提供待核/待补强" in paragraph.text:
            paragraph.text = paragraph.text.replace("企业提供待核/待补强", "待企业确认")
        if "文件、页码或台账" in paragraph.text and PLACEHOLDER.search(paragraph.text):
            paragraph.text = PLACEHOLDER.sub(source_name, paragraph.text)


def _set_public_document_metadata(document: Document, fixture: dict[str, Any]) -> None:
    properties = document.core_properties
    properties.title = f"{fixture['enterprise']}_{fixture['project_object']}_项目前期评估报告"
    properties.subject = "政府项目申报前期评估"
    properties.author = "共创研究院"
    properties.last_modified_by = "共创研究院"
    properties.keywords = "政府项目申报,前期评估,共创研究院"
    properties.category = "政府项目咨询报告"
    properties.comments = ""


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell, value: str, *, centered: bool = False) -> None:
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(8)


def _add_hyperlink(paragraph, label: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), PORTABLE_CJK_FONT)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    properties.append(size)
    text = OxmlElement("w:t")
    text.text = label
    run.append(properties)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_source_ledger(document: Document, fixture: dict[str, Any]) -> None:
    document.add_heading("附录 C、数据来源", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["序号", "文件名称", "链接"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        _set_cell_shading(cell, "A51C30")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)
    for index, source in enumerate(fixture["_validated_sources"], start=1):
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(index), centered=True)
        _set_cell_text(cells[1], source["name"])
        reference = source["public_source"]
        if reference["url"] is None:
            _set_cell_text(cells[2], str(reference["label"]))
        else:
            _set_cell_text(cells[2], "")
            _add_hyperlink(cells[2].paragraphs[0], str(reference["label"]), str(reference["url"]))


def document_text(document: Document) -> str:
    return "\n".join(paragraph.text for paragraph in _iter_paragraphs(document))


def complete_report(
    *,
    template_path: Path,
    output_path: Path,
    fixture: dict[str, Any],
    report_type: str,
    release_tag: str,
    public_root: Path | None = None,
) -> dict[str, Any]:
    if report_type not in REPORT_TYPES:
        raise ValueError(f"未知报告类型:{report_type}")
    template_path = template_path.expanduser().resolve()
    output_path = output_path.expanduser().resolve()
    if not template_path.is_file():
        raise FileNotFoundError(template_path)
    if output_path.exists() and output_path != template_path:
        raise FileExistsError(f"拒绝覆盖现有文件:{output_path}")
    validated = validate_case_fixture(
        {**fixture, "release_tag": release_tag},
        expected_project_id=str(fixture["project_id"]),
        public_root=public_root,
    )
    document = Document(template_path)
    feasibility = report_type == "feasibility"
    for table in document.tables:
        _fill_table_by_header(table, validated, feasibility)
        _lock_table_pagination(table)
    _fill_project_specific_paragraphs(document, validated)
    _generic_fill(document, validated)
    _append_source_ledger(document, validated)
    _lock_table_pagination(document.tables[-1])
    _apply_portable_cjk_font(document)
    _set_public_document_metadata(document, validated)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    rendered = Document(output_path)
    text = document_text(rendered)
    errors: list[str] = []
    if PLACEHOLDER.search(text):
        errors.append("仍存在方括号填写占位符")
    for marker in TRAINING_MARKERS:
        if marker in text:
            errors.append(f"仍存在培训标记:{marker}")
    for marker in INTERNAL_DRAFT_MARKERS:
        if marker in text:
            errors.append(f"仍存在内部候选标记:{marker}")
    if str(validated["enterprise"]) not in text:
        errors.append("成稿缺企业名称")
    if str(validated["project_object"]) not in text:
        errors.append("成稿缺项目核心对象")
    if "数据来源" not in text:
        errors.append("成稿缺资料来源台账")
    if PUBLIC_DIGEST.search(text):
        errors.append("对外成稿不得展示内部文件校验值")
    result = {
        "schema": "gongchuang-completed-project-report/v1",
        "status": "pass" if not errors else "fail",
        "release_tag": release_tag,
        "project_id": validated["project_id"],
        "report_type": report_type,
        "enterprise": validated["enterprise"],
        "template_path": str(template_path),
        "template_sha256": sha256_file(template_path),
        "output_path": str(output_path),
        "output_sha256": sha256_file(output_path),
        "source_count": len(validated["_validated_sources"]),
        "source_receipts": [
            {
                "name": item["name"],
                "sha256": item["sha256"],
                "anchor_count": len(item["anchors"]),
            }
            for item in validated["_validated_sources"]
        ],
        "errors": errors,
    }
    receipt = output_path.with_suffix(".completion.json")
    receipt.write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    if errors:
        raise ValueError("成稿校验失败:" + "；".join(errors))
    return {**result, "receipt_path": str(receipt)}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--fixture", type=Path, required=True)
    parser.add_argument("--report-type", choices=sorted(REPORT_TYPES), required=True)
    parser.add_argument("--release-tag", required=True)
    parser.add_argument("--public-root", type=Path)
    args = parser.parse_args()
    fixture = json.loads(args.fixture.expanduser().resolve().read_text(encoding="utf-8"))
    result = complete_report(
        template_path=args.template,
        output_path=args.output,
        fixture=fixture,
        report_type=args.report_type,
        release_tag=args.release_tag,
        public_root=args.public_root.expanduser().resolve() if args.public_root else None,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
