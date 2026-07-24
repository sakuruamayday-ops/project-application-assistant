#!/usr/bin/env python3
"""从技能包内置母版初始化企业标准 Word，不修改母版本体。"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("缺少 python-docx，请使用宿主平台文档运行时执行。") from exc


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "gbt-1-1-enterprise-standard-template.docx"
STANDARD_TYPES = (
    "product",
    "test-method",
    "service",
    "management",
    "process",
    "terminology",
    "classification",
    "basic",
)
TYPE_NAMES = {
    "product": "产品标准",
    "test-method": "试验方法标准",
    "service": "服务标准",
    "management": "管理标准",
    "process": "过程标准",
    "terminology": "术语标准",
    "classification": "分类与编码标准",
    "basic": "基础通用标准",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="初始化GB/T 1.1企业标准Word母版字段")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    parser.add_argument("--company", required=True, help="企业完整名称")
    parser.add_argument("--standard-code", required=True, help="如 Q/ABC 001—2026")
    parser.add_argument("--title", required=True, help="标准中文名称")
    parser.add_argument("--english-title", default="【英文名称待确认】")
    parser.add_argument("--ics", default="【待确认】")
    parser.add_argument("--ccs", default="【待确认】")
    parser.add_argument("--publish-date", default="【发布日期待确认】")
    parser.add_argument("--implement-date", default="【实施日期待确认】")
    parser.add_argument("--standard-type", choices=STANDARD_TYPES, default="product")
    parser.add_argument("--template", help="可选母版；默认使用技能包内置母版")
    return parser.parse_args()


def iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def all_paragraphs(document):
    seen = set()
    for paragraph in iter_paragraphs(document):
        if id(paragraph._p) not in seen:
            seen.add(id(paragraph._p))
            yield paragraph
    for section in document.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for paragraph in iter_paragraphs(part):
                if id(paragraph._p) not in seen:
                    seen.add(id(paragraph._p))
                    yield paragraph


def replace_paragraph(paragraph, replacements):
    runs = paragraph.runs
    if not runs:
        return 0
    original = "".join(run.text for run in runs)
    updated = original
    for old, new in replacements:
        updated = updated.replace(old, new)
    if updated == original:
        return 0
    runs[0].text = updated
    for run in runs[1:]:
        run.text = ""
    return 1


def extract_text(document) -> str:
    return "\n".join(paragraph.text for paragraph in all_paragraphs(document))


def main() -> int:
    args = parse_args()
    output = Path(args.output).expanduser().resolve()
    template = Path(args.template).expanduser().resolve() if args.template else DEFAULT_TEMPLATE
    if template.suffix.lower() != ".docx" or not template.is_file():
        raise SystemExit(f"母版不存在或不是docx：{template}")
    if output.suffix.lower() != ".docx":
        raise SystemExit("输出文件必须为 .docx")
    if output == template:
        raise SystemExit("禁止直接覆盖技能包母版")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template, output)

    code_prefix = args.standard_code.split()[0]
    replacements = sorted(
        [
            ("ICS 【待确认】\nCCS 【待确认】",
             f"ICS {args.ics}\nCCS {args.ccs}"),
            ("【发布日期】 发布          【实施日期】 实施",
             f"{args.publish_date} 发布          {args.implement_date} 实施"),
            ("Q/【企业代号】 【顺序号】—【年代号】", args.standard_code),
            ("【企业完整名称】", args.company),
            ("【标准中文名称】", args.title),
            ("【英文名称】", args.english_title),
            ("Q/【企业代号】", code_prefix),
        ],
        key=lambda item: len(item[0]),
        reverse=True,
    )

    document = Document(output)
    changed = sum(replace_paragraph(paragraph, replacements)
                  for paragraph in all_paragraphs(document))
    document.save(output)

    verified = Document(output)
    text = extract_text(verified)
    residual = [
        token for token in (
            "【企业完整名称】",
            "Q/【企业代号】",
            "【顺序号】",
            "【年代号】",
            "【标准中文名称】",
            "【英文名称】",
            "【发布日期】",
            "【实施日期】",
        )
        if token in text
    ]

    result = {
        "status": "ok" if not residual else "failed",
        "output": str(output),
        "template": str(template),
        "standard_type": args.standard_type,
        "standard_type_name": TYPE_NAMES[args.standard_type],
        "paragraphs_changed": changed,
        "requires_body_restructure": args.standard_type != "product",
        "residual_template_tokens": residual,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if not residual else 2


if __name__ == "__main__":
    sys.exit(main())
