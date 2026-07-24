#!/usr/bin/env python3
"""审计企业标准DOCX的示例字段、章节类型和GB/T 1.1关键版式。"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError as exc:
    raise SystemExit("缺少 python-docx，请使用宿主平台文档运行时执行。") from exc


STANDARD_TYPES = (
    "product",
    "test-method",
    "service",
    "management",
    "process",
    "terminology",
    "classification",
    "basic",
)
EXAMPLE_TOKENS = (
    "【企业完整名称】",
    "Q/【企业代号】",
    "【顺序号】",
    "【年代号】",
    "【标准中文名称】",
    "【英文名称】",
    "【发布日期】",
    "【实施日期】",
)
REQUIRED_HEADING_GROUPS = {
    "product": [
        ("范围",), ("规范性引用文件",), ("术语和定义",),
        ("技术要求",), ("试验方法",), ("检验规则",),
    ],
    "test-method": [
        ("范围",), ("规范性引用文件",), ("术语和定义",), ("原理",),
        ("试剂", "材料"), ("仪器", "设备"), ("样品", "取样"),
        ("试验条件", "测试条件"), ("试验步骤", "试验程序"),
        ("结果", "计算"), ("试验报告", "测试报告"),
    ],
    "service": [
        ("范围",), ("规范性引用文件",), ("术语和定义",),
        ("服务原则",), ("服务条件",), ("服务流程",),
        ("服务要求",), ("评价",), ("投诉", "改进"),
    ],
    "management": [
        ("范围",), ("规范性引用文件",), ("术语和定义",),
        ("管理原则",), ("职责", "权限"), ("策划",),
        ("实施", "运行"), ("监视", "测量", "评价"), ("改进",),
    ],
    "process": [
        ("范围",), ("规范性引用文件",), ("术语和定义",),
        ("过程输入", "输入要求"), ("过程条件", "工艺条件"),
        ("过程步骤", "工艺步骤"), ("过程控制", "工艺控制"),
        ("验证", "检验"), ("安全", "环境"),
    ],
    "terminology": [
        ("范围",), ("规范性引用文件",), ("术语和定义", "术语条目"),
    ],
    "classification": [
        ("范围",), ("规范性引用文件",), ("分类原则",),
        ("代码结构", "编码结构"), ("编码规则",), ("维护", "扩展"),
    ],
    "basic": [
        ("范围",), ("规范性引用文件",), ("术语和定义",),
        ("总体原则", "总则"), ("通用要求",),
    ],
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="审计企业标准Word字段与GB/T 1.1版式")
    parser.add_argument("docx", help="待审计 .docx")
    parser.add_argument("--standard-type", choices=STANDARD_TYPES, required=True)
    parser.add_argument("--expected-company")
    parser.add_argument("--expected-code")
    parser.add_argument("--expected-title")
    parser.add_argument("--allow-term", action="append", default=[])
    parser.add_argument("--forbid", action="append", default=[])
    return parser.parse_args()


def cm(value):
    return round(value.cm, 2) if value is not None else 0.0


def pt(value):
    return round(value.pt, 1) if value is not None else 0.0


def iter_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def all_paragraphs(document):
    seen = set()
    for paragraph in iter_paragraphs(document):
        if id(paragraph._p) not in seen:
            seen.add(id(paragraph._p))
            yield paragraph
    for section in document.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            for paragraph in iter_paragraphs(part):
                if id(paragraph._p) not in seen:
                    seen.add(id(paragraph._p))
                    yield paragraph


def full_text(document) -> str:
    return "\n".join(paragraph.text for paragraph in all_paragraphs(document))


def contains_group(headings, group) -> bool:
    return any(any(keyword in heading for keyword in group) for heading in headings)


def has_terminal_line(document) -> bool:
    if not document.tables:
        return False
    table = document.tables[-1]
    if len(table.rows) != 1 or len(table.columns) != 1:
        return False
    borders = table.cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
    if borders is None:
        return False
    top = borders.find(qn("w:top"))
    if top is None:
        return False
    return top.get(qn("w:val")) not in (None, "nil", "none")


def main() -> int:
    args = parse_args()
    path = Path(args.docx).expanduser().resolve()
    if path.suffix.lower() != ".docx" or not path.is_file():
        raise SystemExit(f"待审计文件不存在或不是docx：{path}")
    document = Document(path)
    text = full_text(document)
    errors = []
    warnings = []

    expected_values = tuple(
        value for value in (args.expected_company, args.expected_code, args.expected_title)
        if value
    ) + tuple(args.allow_term)
    forbidden = list(EXAMPLE_TOKENS) + list(args.forbid)
    residual = sorted({
        token for token in forbidden
        if token in text and not any(token in expected for expected in expected_values)
    })
    if residual:
        errors.append("发现未清除的母版示例字段：" + "、".join(residual))

    for label, expected in (
        ("企业名称", args.expected_company),
        ("标准编号", args.expected_code),
        ("标准名称", args.expected_title),
    ):
        if expected and expected not in text:
            errors.append(f"未找到预期{label}：{expected}")

    headings = [
        re.sub(r"^\d+(?:\.\d+)*\s*", "", paragraph.text.strip())
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1" and paragraph.text.strip()
    ]
    missing_groups = [
        "/".join(group)
        for group in REQUIRED_HEADING_GROUPS[args.standard_type]
        if not contains_group(headings, group)
    ]
    if missing_groups:
        errors.append(
            f"{args.standard_type}章节模型缺少：" + "、".join(missing_groups)
        )

    for index, section in enumerate(document.sections, start=1):
        actual = (
            cm(section.top_margin), cm(section.bottom_margin),
            cm(section.left_margin), cm(section.right_margin),
        )
        expected = (3.3, 2.0, 2.5, 2.0)
        if actual != expected:
            errors.append(f"第{index}节页边距{actual}，预期{expected}")
        page_size = (cm(section.page_width), cm(section.page_height))
        if abs(page_size[0] - 21.0) > 0.05 or abs(page_size[1] - 29.7) > 0.05:
            errors.append(f"第{index}节不是A4页面：{page_size}")

    style_expectations = {
        "Normal": (10.5, 0.0, 0.0),
        "Heading 1": (10.5, 18.0, 18.0),
        "Heading 2": (10.5, 9.0, 9.0),
        "Note Text": (9.0, 0.0, 0.0),
        "Table Text": (9.0, 0.0, 0.0),
    }
    for name, expected in style_expectations.items():
        if name not in document.styles:
            errors.append(f"缺少Word样式：{name}")
            continue
        style = document.styles[name]
        actual = (
            pt(style.font.size),
            pt(style.paragraph_format.space_before),
            pt(style.paragraph_format.space_after),
        )
        if actual != expected:
            errors.append(f"样式{name}参数{actual}，预期{expected}")

    bad_body = []
    bad_lists = []
    in_main_text = False
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        if paragraph.style.name == "Heading 1":
            in_main_text = True
            continue
        if not in_main_text:
            continue
        fmt = paragraph.paragraph_format
        if re.match(r"^[a-z]\)", value):
            if (cm(fmt.left_indent), cm(fmt.first_line_indent)) != (1.48, -0.74):
                bad_lists.append(value[:24])
            continue
        skip = (
            paragraph.style.name != "Normal"
            or paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
            or cm(fmt.left_indent) != 0.0
            or value in {"前言", "附录 A", "（资料性）", "（规范性）"}
            or value.startswith(("表", "图", "Q/", "ICS ", "CCS "))
        )
        if not skip and len(value) >= 12 and cm(fmt.first_line_indent) not in (0.74,):
            bad_body.append(value[:24])
    if bad_body:
        errors.append("正文未按两个汉字首行缩进：" + "；".join(bad_body[:8]))
    if bad_lists:
        errors.append("列项悬挂缩进不符合要求：" + "；".join(bad_lists[:8]))

    if any(paragraph.style.name.startswith(("List Bullet", "List Number"))
           for paragraph in document.paragraphs):
        errors.append("发现Word默认项目符号或编号样式")
    if not has_terminal_line(document):
        errors.append("末页未检测到标准终结线")

    if "【待确认】" in text or "待确认" in text:
        warnings.append("文件仍含待确认字段，只能作为草案")

    result = {
        "status": "pass" if not errors else "fail",
        "docx": str(path),
        "standard_type": args.standard_type,
        "heading_1": headings,
        "errors": errors,
        "warnings": warnings,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if not errors else 2


if __name__ == "__main__":
    sys.exit(main())
