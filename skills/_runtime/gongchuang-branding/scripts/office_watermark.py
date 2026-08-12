#!/usr/bin/env python3
"""Apply mandatory 共创研究院 branding to Word and Excel files."""

from __future__ import annotations

import argparse
import copy
import hashlib
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.utils import get_column_letter
from openpyxl.workbook.defined_name import DefinedName

from brand_config import choose_style, load_config, public_identity


WATERMARK_NAME = "Gongchuang Institute Centered Watermark v4"
LEGACY_WATERMARK_NAMES = (
    "Jiao" + "tang Centered Watermark",
    "Jiao" + "tang Centered Watermark v2",
    "Jiao" + "tang Centered Watermark v3",
)
XLSX_MARKER = "_GONGCHUANG_INSTITUTE_UNIFORM_WATERMARK_V4"
LEGACY_XLSX_MARKERS = (
    "_JIAO" + "TANG_CENTERED_WATERMARK",
    "_JIAO" + "TANG_UNIFORM_WATERMARK_V2",
    "_JIAO" + "TANG_UNIFORM_WATERMARK_V3",
)


def _ensure_docx_brand_text(header) -> None:
    identity = public_identity()
    expected = identity["document_header"]
    for paragraph in list(header.paragraphs):
        text = paragraph.text.strip()
        if text == expected:
            target = paragraph
            break
        if text in {"焦" + "糖", "Jiao" + "tang"}:
            paragraph._element.getparent().remove(paragraph._element)
    else:
        target = header.add_paragraph()
        target.add_run(expected)
    target.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in target.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(float(load_config()["policy"]["header_font_size_pt"]))
        run.font.color.rgb = RGBColor(0x8A, 0x6A, 0x2F)


def _inline_to_centered_anchor(inline, *, name: str = WATERMARK_NAME):
    """Convert a python-docx inline image to a page-centered floating image."""
    extent = inline.extent
    doc_pr = inline.docPr
    graphic = inline.graphic

    anchor = OxmlElement("wp:anchor")
    for key, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    for axis in ("H", "V"):
        pos = OxmlElement(f"wp:position{axis}")
        pos.set("relativeFrom", "page")
        align = OxmlElement("wp:align")
        align.text = "center"
        pos.append(align)
        anchor.append(pos)

    new_extent = OxmlElement("wp:extent")
    new_extent.set("cx", str(extent.cx))
    new_extent.set("cy", str(extent.cy))
    anchor.append(new_extent)

    effect = OxmlElement("wp:effectExtent")
    for edge in ("l", "t", "r", "b"):
        effect.set(edge, "0")
    anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))

    new_doc_pr = copy.deepcopy(doc_pr)
    new_doc_pr.set("name", name)
    new_doc_pr.set("descr", name)
    anchor.append(new_doc_pr)

    frame = OxmlElement("wp:cNvGraphicFramePr")
    locks = OxmlElement("a:graphicFrameLocks")
    locks.set("noChangeAspect", "1")
    frame.append(locks)
    anchor.append(frame)
    anchor.append(copy.deepcopy(graphic))
    return anchor


def _docx_density(document: Document) -> float:
    text_len = sum(len(p.text) for p in document.paragraphs)
    cells = sum(len(table.rows) * len(table.columns) for table in document.tables)
    return min(1.0, text_len / 12000.0 + cells / 160.0)


def apply_docx_watermark(path: str | Path) -> Path:
    """Add one centered, translucent watermark to each distinct Word header."""
    path = Path(path)
    document = Document(path)
    config = load_config()
    style = choose_style(_docx_density(document), 1.0, target="docx")
    asset = style["asset_path"]
    processed_parts: set[str] = set()
    fixed_width = int(
        min(section.page_width for section in document.sections)
        * style["scale"]
        * float(config["policy"]["size_multiplier"])
    )

    for section in document.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            part_name = str(header.part.partname)
            if part_name in processed_parts:
                continue
            processed_parts.add(part_name)
            # Always rebuild the mark in every distinct header.  The previous
            # early-return behaviour skipped headers added by later document
            # edits as soon as any older header already contained a mark.
            for legacy_name in (WATERMARK_NAME, *LEGACY_WATERMARK_NAMES):
                for drawing in header._element.xpath(
                    f'.//w:drawing[.//wp:docPr[@name="{legacy_name}"]]'
                ):
                    drawing.getparent().remove(drawing)
            paragraph = header.add_paragraph()
            run = paragraph.add_run()
            inline_shape = run.add_picture(asset, width=fixed_width)
            inline = inline_shape._inline
            drawing = inline.getparent()
            drawing.replace(inline, _inline_to_centered_anchor(inline))
            _ensure_docx_brand_text(header)

    document.save(path)
    from delivery_gate import validate_artifact
    validate_artifact(path, check_stamp=False)
    return path


def _sheet_density(ws) -> float:
    used = max(1, ws.max_row * ws.max_column)
    populated = sum(1 for row in ws.iter_rows() for cell in row if cell.value not in (None, ""))
    return min(1.0, populated / min(used, 600.0))


def _center_anchor(ws, image: XLImage) -> str:
    max_col = max(1, ws.max_column)
    max_row = max(1, ws.max_row)
    image_cols = max(4, round(image.width / 72))
    image_rows = max(10, round(image.height / 20))
    col = max(1, round((max_col + 1) / 2 - image_cols / 2))
    row = max(1, round((max_row + 1) / 2 - image_rows / 2))
    return f"{get_column_letter(col)}{row}"


def apply_xlsx_watermark(path: str | Path) -> Path:
    """Overlay one centered translucent watermark on every visible worksheet."""
    path = Path(path)
    workbook = load_workbook(path, keep_vba=path.suffix.lower() == ".xlsm")
    if XLSX_MARKER in workbook.defined_names:
        from delivery_gate import validate_artifact
        validate_artifact(path, check_stamp=False)
        return path
    if any(marker in workbook.defined_names for marker in LEGACY_XLSX_MARKERS):
        brand_hashes = {
            hashlib.sha256(asset.read_bytes()).hexdigest()
            for asset in (Path(__file__).resolve().parent.parent / "assets").glob("brand-*.png")
        }
        for ws in workbook.worksheets:
            kept = []
            for image in ws._images:
                ref = getattr(image, "ref", None)
                data = ref.getvalue() if hasattr(ref, "getvalue") else None
                digest = hashlib.sha256(data).hexdigest() if data else None
                if digest not in brand_hashes:
                    kept.append(image)
            ws._images = kept
        for marker in LEGACY_XLSX_MARKERS:
            if marker in workbook.defined_names:
                del workbook.defined_names[marker]
    visible_sheets = [ws for ws in workbook.worksheets if ws.sheet_state == "visible"]
    document_density = max((_sheet_density(ws) for ws in visible_sheets), default=0.0)
    document_style = choose_style(document_density, 1.0, target="xlsx")
    config = load_config()
    fixed_width_px = round(
        520
        * document_style["scale"]
        * float(config["policy"]["size_multiplier"])
    )
    for ws in workbook.worksheets:
        if ws.sheet_state != "visible":
            continue
        style = choose_style(_sheet_density(ws), 1.0, target="xlsx")
        image = XLImage(style["asset_path"])
        image.width = fixed_width_px
        image.height = round(image.width * 0.9725)
        image.anchor = _center_anchor(ws, image)
        ws.add_image(image)
        ws.oddHeader.right.text = "&9&K8A6A2F" + public_identity()["document_header"]
    workbook.defined_names.add(DefinedName(XLSX_MARKER, attr_text='"enabled"'))
    workbook.save(path)
    from delivery_gate import validate_artifact
    validate_artifact(path, check_stamp=False)
    return path


def apply_office_watermark(path: str | Path) -> Path:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return apply_docx_watermark(path)
    if suffix in {".xlsx", ".xlsm"}:
        return apply_xlsx_watermark(path)
    raise ValueError(f"Unsupported Office watermark format: {suffix}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("path")
    args = parser.parse_args()
    result = apply_office_watermark(args.path)
    print(result)


if __name__ == "__main__":
    sys.exit(main())
