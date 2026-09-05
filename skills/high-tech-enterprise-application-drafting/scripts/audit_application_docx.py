#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


RD_PATTERN = re.compile(r"研发活动编号[:：]?\s*(RD\d{2})", re.IGNORECASE)
RD_TECH_PATTERN = re.compile(r"^[12]、.+技术：.+拟定技术指标为.+\d")
UNIT_OR_BOUNDARY = re.compile(
    r"(%|秒|分钟|小时|天|毫米|厘米|微米|千米|米|赫兹|分贝|摄氏度|℃|兆帕|MPa|"
    r"伏|安|瓦|帧|次|个|类|项|套|点|条|不低于|不高于|不大于|不小于|不超过|不少于)"
)
INNOVATION_SECTION_ORDER = (
    "知识产权对企业竞争力的作用",
    "科技成果转化情况",
    "研究开发与技术创新组织管理情况",
    "管理与科技人员情况",
)
PENDING_IP_TERMS = ("申请", "受理", "审中", "实质审查")
AUTHORIZED_IP_TERMS = ("已授权", "获得发明专利", "授权成果")
STATUS_BOUNDARY_TERMS = ("不计入", "不得", "不能", "尚未", "仅作为", "不作为")
UNSUPPORTED_RESULT_PHRASES = (
    "营业收入增长率",
    "用户高度认可",
    "大型工程",
    "多所高校",
    "科技项目资助",
    "税收优惠",
)


def physical_cells(row):
    """Return the row's real w:tc nodes instead of the grid-expanded proxy list."""
    yield from row._tr.findall(qn("w:tc"))


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in physical_cells(row):
                for paragraph in cell.iter(qn("w:p")):
                    yield Paragraph(paragraph, table)


def xml_text(node) -> str:
    return "".join(item.text or "" for item in node.iter(qn("w:t")))


def xml_cell_text(cell) -> str:
    paragraphs = []
    for paragraph in cell.findall(qn("w:p")):
        parts = []
        for item in paragraph.iter():
            if item.tag == qn("w:t"):
                parts.append(item.text or "")
            elif item.tag == qn("w:br"):
                parts.append("\n")
            elif item.tag == qn("w:tab"):
                parts.append("\t")
        paragraphs.append("".join(parts))
    return "\n".join(paragraphs)


def normalized(value: str) -> str:
    return "".join(value.replace("\u3000", " ").split())


def canonical_innovation_label(value: str) -> str | None:
    compact = normalized(value)
    for label in INNOVATION_SECTION_ORDER:
        if compact.startswith(normalized(label)):
            return label
    return None


def _evidence_list(evidence: dict | None, key: str) -> list[str]:
    if not evidence:
        return []
    value = evidence.get(key, [])
    return [item.strip() for item in value if isinstance(item, str) and item.strip()]


def audit_innovation_capability(document, evidence: dict | None = None) -> tuple[dict, list[dict]]:
    """Audit fixed section order, length, enterprise isolation and evidence boundaries."""
    result: dict[str, dict] = {}
    order: list[str] = []
    issues: list[dict] = []
    for table in document.tables:
        for row in table.rows:
            cells = list(physical_cells(row))
            if len(cells) < 2:
                continue
            label = canonical_innovation_label(xml_cell_text(cells[0]))
            if label is None:
                continue
            text = xml_cell_text(cells[1]).strip()
            length = len(normalized(text))
            order.append(label)
            result[label] = {"length": length, "text": text, "issues": []}

    if not order:
        return result, issues
    if order != list(INNOVATION_SECTION_ORDER):
        issues.append(
            {
                "section": "企业创新能力",
                "issue": f"四栏顺序或完整性错误：{order}",
            }
        )
    for label in INNOVATION_SECTION_ORDER:
        if label not in result:
            continue
        if result[label]["length"] > 400:
            issue = f"正文{result[label]['length']}字，超过400字上限"
            result[label]["issues"].append(issue)
            issues.append({"section": label, "issue": issue})

    if not evidence:
        return result, issues

    allowed_names = set(_evidence_list(evidence, "allowed_corporate_names"))
    company_pattern = re.compile(r"[\u4e00-\u9fffA-Za-z0-9·（）()]{2,50}有限公司")
    for label, item in result.items():
        foreign_names = sorted(
            {
                matched
                for matched in company_pattern.findall(item["text"])
                if not any(matched.endswith(name) for name in allowed_names)
            }
        )
        if foreign_names:
            issue = f"出现未在本企业证据中登记的企业名称：{foreign_names}"
            item["issues"].append(issue)
            issues.append({"section": label, "issue": issue})

    ip_text = result.get(INNOVATION_SECTION_ORDER[0], {}).get("text", "")
    for sentence in re.split(r"[。；;\n]", ip_text):
        if (
            any(term in sentence for term in PENDING_IP_TERMS)
            and any(term in sentence for term in AUTHORIZED_IP_TERMS)
            and not any(term in sentence for term in STATUS_BOUNDARY_TERMS)
        ):
            issue = f"申请、受理或审中知识产权疑似被写成授权成果：{sentence.strip()}"
            result[INNOVATION_SECTION_ORDER[0]]["issues"].append(issue)
            issues.append({"section": INNOVATION_SECTION_ORDER[0], "issue": issue})
    if not evidence.get("allow_financing_claims", False):
        financing = [term for term in ("融资", "估值", "投资者") if term in ip_text]
        if financing:
            issue = f"缺少融资作用证据但出现相关表述：{financing}"
            result[INNOVATION_SECTION_ORDER[0]]["issues"].append(issue)
            issues.append({"section": INNOVATION_SECTION_ORDER[0], "issue": issue})

    result_text = result.get(INNOVATION_SECTION_ORDER[1], {}).get("text", "")
    allowed_result_claims = _evidence_list(evidence, "allowed_result_claims")
    for phrase in UNSUPPORTED_RESULT_PHRASES:
        if phrase in result_text and not any(phrase in claim for claim in allowed_result_claims):
            issue = f"科技成果转化栏出现未登记的市场、合作或政策主张：{phrase}"
            result[INNOVATION_SECTION_ORDER[1]]["issues"].append(issue)
            issues.append({"section": INNOVATION_SECTION_ORDER[1], "issue": issue})

    management_text = result.get(INNOVATION_SECTION_ORDER[2], {}).get("text", "")
    policy_titles = set(_evidence_list(evidence, "policy_titles"))
    used_titles = set(re.findall(r"《([^》]+)》", management_text))
    unknown_titles = sorted(used_titles - policy_titles)
    if unknown_titles:
        issue = f"制度名称未逐字命中详细制度文件：{unknown_titles}"
        result[INNOVATION_SECTION_ORDER[2]]["issues"].append(issue)
        issues.append({"section": INNOVATION_SECTION_ORDER[2], "issue": issue})
    if policy_titles and not used_titles:
        issue = "已提供制度文件但组织管理栏未引用任何制度全称"
        result[INNOVATION_SECTION_ORDER[2]]["issues"].append(issue)
        issues.append({"section": INNOVATION_SECTION_ORDER[2], "issue": issue})

    personnel_text = result.get(INNOVATION_SECTION_ORDER[3], {}).get("text", "")
    allowed_personnel_claims = _evidence_list(evidence, "allowed_personnel_claims")
    for term in ("博士", "硕士", "留任率"):
        if term in personnel_text and not any(term in claim for claim in allowed_personnel_claims):
            issue = f"人员栏出现未登记的人员结构或稳定性主张：{term}"
            result[INNOVATION_SECTION_ORDER[3]]["issues"].append(issue)
            issues.append({"section": INNOVATION_SECTION_ORDER[3], "issue": issue})
    for value, unit in re.findall(r"(\d+(?:\.\d+)?)\s*(名|人|%)", personnel_text):
        token = f"{value}{unit}"
        if not any(token in claim.replace(" ", "") for claim in allowed_personnel_claims):
            issue = f"人员栏出现未登记的数量或比例：{token}"
            result[INNOVATION_SECTION_ORDER[3]]["issues"].append(issue)
            issues.append({"section": INNOVATION_SECTION_ORDER[3], "issue": issue})
    return result, issues


def find_xml_label_cell(table, label: str):
    for row in table.findall(qn("w:tr")):
        cells = row.findall(qn("w:tc"))
        for index, cell in enumerate(cells):
            if normalized(xml_text(cell)).startswith(label):
                return cells, index
    return None


def audit_rd_core_innovation(document) -> tuple[dict, list[dict]]:
    result = {}
    issues = []
    current_rd = None
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            match = RD_PATTERN.search(xml_text(child))
            if match:
                current_rd = match.group(1).upper()
            continue
        if child.tag != qn("w:tbl") or current_rd is None:
            continue
        field_row = find_xml_label_cell(child, "技术领域")
        core_row = find_xml_label_cell(child, "核心技术及创新点")
        title_row = find_xml_label_cell(child, "研发活动名称")
        if not field_row or not core_row or not title_row:
            continue
        field_cells, field_index = field_row
        core_cells, core_index = core_row
        if field_index + 1 >= len(field_cells) or core_index + 1 >= len(core_cells):
            issues.append({"rd_id": current_rd, "issue": "领域或核心技术数据单元格缺失"})
            continue
        field = xml_text(field_cells[field_index + 1]).strip()
        core_text = xml_cell_text(core_cells[core_index + 1]).strip()
        if not core_text:
            continue
        lines = [line.strip() for line in core_text.splitlines()]
        body_length = (
            len("".join(lines[index] for index in (2, 3, 5, 6)))
            if len(lines) == 7
            else 0
        )
        item_issues = []
        if field == "待企业确认四级领域":
            item_issues.append("四级技术领域待确认，不能作为已核定领域交付")
        if len(lines) != 7:
            item_issues.append(f"应为7行，实际为{len(lines)}行")
        if len(lines) >= 1 and lines[0] != f"所属技术领域：{field}。":
            item_issues.append("正文所属技术领域与本RD表头不一致")
        if len(lines) >= 2 and lines[1] != "核心技术：":
            item_issues.append("第2行必须为核心技术：")
        if len(lines) >= 5 and lines[4] != "创新点：":
            item_issues.append("第5行必须为创新点：")
        for index in (2, 3):
            if len(lines) <= index or not RD_TECH_PATTERN.search(lines[index]):
                item_issues.append(f"第{index + 1}行须为具名核心技术并含拟定技术指标与数值")
            elif not UNIT_OR_BOUNDARY.search(lines[index]):
                item_issues.append(f"第{index + 1}行技术指标缺少单位或阈值边界")
        for index, number in ((5, "1、"), (6, "2、")):
            if len(lines) <= index or not lines[index].startswith(number):
                item_issues.append(f"第{index + 1}行必须以{number}开头")
        if body_length > 400:
            item_issues.append(f"核心技术与创新点合计{body_length}字，超过400字上限")
        result[current_rd] = {
            "field": field,
            "line_count": len(lines),
            "body_length": body_length,
            "core_technology_count": sum(
                1 for line in lines if RD_TECH_PATTERN.search(line)
            ),
            "innovation_count": sum(
                1 for line in lines if line.startswith(("1、", "2、"))
            ) - sum(1 for line in lines if RD_TECH_PATTERN.search(line)),
            "issues": item_issues,
        }
        issues.extend({"rd_id": current_rd, "issue": issue} for issue in item_issues)
    return result, issues


def audit_rd_stage_results(document) -> tuple[dict, list[dict]]:
    result: dict[str, dict] = {}
    issues: list[dict] = []
    current_rd = None
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            match = RD_PATTERN.search(xml_text(child))
            if match:
                current_rd = match.group(1).upper()
            continue
        if child.tag != qn("w:tbl") or current_rd is None:
            continue
        title_row = find_xml_label_cell(child, "研发活动名称")
        result_row = find_xml_label_cell(child, "取得的阶段性成果")
        if not title_row or not result_row:
            continue
        cells, label_index = result_row
        if label_index + 1 >= len(cells):
            issues.append({"rd_id": current_rd, "issue": "阶段性成果数据单元格缺失"})
            continue
        text = xml_cell_text(cells[label_index + 1]).strip()
        if not text:
            continue
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        item_issues: list[str] = []
        if len(lines) != 4:
            item_issues.append(f"阶段性成果应为4条，实际为{len(lines)}条")
        for index, line in enumerate(lines[:4], 1):
            if not line.startswith(f"{index}、"):
                item_issues.append(f"阶段性成果第{index}条必须以{index}、开头")
        body_length = len(normalized(text))
        if body_length > 400:
            item_issues.append(f"阶段性成果合计{body_length}字，超过400字上限")
        result[current_rd] = {
            "line_count": len(lines),
            "body_length": body_length,
            "issues": item_issues,
        }
        issues.extend({"rd_id": current_rd, "issue": issue} for issue in item_issues)
    return result, issues


def audit_limited_fields(document) -> tuple[list[dict], list[dict]]:
    """Audit every populated table field whose physical label says “限400字”."""
    result: list[dict] = []
    issues: list[dict] = []
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            cells = list(physical_cells(row))
            for cell_index, label_cell in enumerate(cells):
                label = normalized(xml_cell_text(label_cell))
                if "限400字" not in label:
                    continue
                # RD专项检查按业务口径计数：核心正文不计领域和标题，
                # 阶段成果把四条合并计数，这里不重复审计。
                if label.startswith(("核心技术及创新点", "取得的阶段性成果")):
                    continue
                if cell_index + 1 >= len(cells):
                    issues.append({"field": label, "issue": "标签后缺少数据单元格"})
                    continue
                text = xml_cell_text(cells[cell_index + 1]).strip()
                if not text:
                    continue
                length = len(normalized(text))
                record = {
                    "table": table_index,
                    "row": row_index,
                    "field": label,
                    "length": length,
                }
                result.append(record)
                if length > 400:
                    issues.append({**record, "issue": f"正文{length}字，超过400字上限"})
    return result, issues


def main():
    parser = argparse.ArgumentParser(description="Audit a high-tech enterprise application DOCX.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--font", default="宋体")
    parser.add_argument("--size", type=float, default=12.0)
    parser.add_argument(
        "--innovation-evidence",
        type=Path,
        help="企业创新能力四栏证据边界JSON；正式回填时必须提供",
    )
    args = parser.parse_args()

    document = Document(args.docx)
    evidence = None
    if args.innovation_evidence:
        evidence = json.loads(args.innovation_evidence.read_text(encoding="utf-8"))
        if not isinstance(evidence, dict):
            raise SystemExit("--innovation-evidence必须指向JSON对象")
    font_issues = []
    placeholders = []

    for paragraph in iter_paragraphs(document):
        if "XXX" in paragraph.text or "待企业核定" in paragraph.text or "拟定指标" in paragraph.text:
            placeholders.append(paragraph.text.strip())
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            east_asia = None
            if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                east_asia = run._element.rPr.rFonts.get(qn("w:eastAsia"))
            size = run.font.size.pt if run.font.size is not None else None
            if args.font not in {run.font.name, east_asia} or size is None or abs(size - args.size) > 0.01:
                font_issues.append(
                    {"text": run.text[:40], "font": run.font.name, "eastAsia": east_asia, "size": size}
                )

    innovation_detail, innovation_issues = audit_innovation_capability(document, evidence)
    innovation_lengths = {
        label: item["length"] for label, item in innovation_detail.items()
    }

    rd_core_innovation, rd_core_issues = audit_rd_core_innovation(document)
    rd_stage_results, rd_stage_issues = audit_rd_stage_results(document)
    limited_fields, limited_field_issues = audit_limited_fields(document)

    result = {
        "file": str(args.docx),
        "tables": len(document.tables),
        "font_issue_count": len(font_issues),
        "font_issue_sample": font_issues[:20],
        "innovation_capability_lengths": innovation_lengths,
        "innovation_over_400": {
            label: length for label, length in innovation_lengths.items() if length > 400
        },
        "innovation_capability": innovation_detail,
        "innovation_issue_count": len(innovation_issues),
        "innovation_issues": innovation_issues,
        "rd_core_innovation": rd_core_innovation,
        "rd_core_issue_count": len(rd_core_issues),
        "rd_core_issues": rd_core_issues,
        "rd_stage_results": rd_stage_results,
        "rd_stage_issue_count": len(rd_stage_issues),
        "rd_stage_issues": rd_stage_issues,
        "limited_fields": limited_fields,
        "limited_field_issue_count": len(limited_field_issues),
        "limited_field_issues": limited_field_issues,
        "placeholder_count": len(placeholders),
        "placeholder_sample": placeholders[:20],
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    raise SystemExit(
        1
        if font_issues or innovation_issues or rd_core_issues or rd_stage_issues or limited_field_issues
        else 0
    )


if __name__ == "__main__":
    main()
