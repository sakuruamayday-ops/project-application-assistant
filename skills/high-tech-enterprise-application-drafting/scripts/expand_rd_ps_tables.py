#!/usr/bin/env python3
from __future__ import annotations

import argparse
import copy
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


RD_MARKER = "研发活动编号"
PS_MARKER = "编号：PS"

RD_LABELS = {
    "研发活动名称",
    "起止时间",
    "技术领域",
    "技术来源",
    "知识产权编号",
    "研发经费总预算（万元）",
    "研发经费近三年总支出（万元）",
    "其中",
    "第一年",
    "第二年",
    "第三年",
    "目的及组织实施方式（限400字）",
    "核心技术及创新点（限400字）",
    "取得的阶段性成果（限400字）",
}

PS_LABELS = {
    "产品（服务）名称",
    "技术领域",
    "技术来源",
    "上年度销售收入（万元）",
    "是否主要产品（服务）",
    "□是□否",
    "知识产权编号",
    "关键技术及主要技术指标（限400字）",
    "与同类产品（服务）的竞争优势（限400字）",
    "知识产权获得情况及其对产品（服务）在技术上发挥的支持作用（限400字）",
}


def normalize(text: str) -> str:
    return "".join(text.replace("\u3000", " ").split())


def unique_cells(row):
    seen = set()
    for cell in row.cells:
        key = cell._tc
        if key not in seen:
            seen.add(key)
            yield cell


def table_kind(table) -> str | None:
    text = normalize("".join(cell.text for row in table.rows for cell in unique_cells(row)))
    if "研发活动名称" in text and "目的及组织实施方式" in text and "阶段性成果" in text:
        return "rd"
    if "产品（服务）名称" in text and "关键技术及主要技术指标" in text and "竞争优势" in text:
        return "ps"
    return None


def paragraph_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def find_previous_marker(table_element, marker: str):
    current = table_element.getprevious()
    while current is not None:
        if current.tag == qn("w:p") and marker in normalize(paragraph_text(current)):
            return current
        current = current.getprevious()
    return None


def set_paragraph_text_preserve_format(paragraph_element, text: str) -> None:
    text_nodes = list(paragraph_element.iter(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ""
        return

    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph_element.append(run)


def clear_cell_preserve_format(cell) -> None:
    for paragraph in cell.paragraphs:
        text_nodes = list(paragraph._p.iter(qn("w:t")))
        for node in text_nodes:
            node.text = ""


def set_cell_text_preserve_format(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph()
    set_paragraph_text_preserve_format(cell.paragraphs[0]._p, text)
    for paragraph in cell.paragraphs[1:]:
        for node in paragraph._p.iter(qn("w:t")):
            node.text = ""


def clear_table_data(table, kind: str) -> None:
    labels = RD_LABELS if kind == "rd" else PS_LABELS
    for row in table.rows:
        for cell in unique_cells(row):
            if normalize(cell.text) not in labels:
                clear_cell_preserve_format(cell)
    if kind == "ps":
        for row in table.rows:
            cells = list(unique_cells(row))
            if cells and normalize(cells[0].text) == "是否主要产品（服务）" and len(cells) >= 2:
                set_cell_text_preserve_format(cells[1], "□是 □否")


def make_page_break_paragraph() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph.append(run)
    return paragraph


def is_page_break_only(paragraph_element) -> bool:
    if paragraph_element is None or paragraph_element.tag != qn("w:p"):
        return False
    if normalize(paragraph_text(paragraph_element)):
        return False
    breaks = list(paragraph_element.iter(qn("w:br")))
    if not breaks or any(node.get(qn("w:type")) != "page" for node in breaks):
        return False
    forbidden = ("w:drawing", "w:object", "w:pict", "w:tab")
    return not any(list(paragraph_element.iter(qn(tag))) for tag in forbidden)


def collect_units(document, kind: str):
    marker = RD_MARKER if kind == "rd" else PS_MARKER
    units = []
    for table in document.tables:
        if table_kind(table) != kind:
            continue
        marker_paragraph = find_previous_marker(table._tbl, marker)
        if marker_paragraph is None:
            raise ValueError(f"找到{kind.upper()}表，但未找到其编号段落")
        units.append((marker_paragraph, table))
    return units


def renumber_units(units, prefix: str) -> None:
    marker = "研发活动编号：" if prefix.upper() == "RD" else "编号："
    for index, (paragraph, _) in enumerate(units, 1):
        set_paragraph_text_preserve_format(paragraph, f"{marker}{prefix.upper()}{index:02d}")


def meaningful_table_values(table, kind: str):
    labels = RD_LABELS if kind == "rd" else PS_LABELS
    normalized_labels = {normalize(label) for label in labels}
    values = []
    seen_cells = set()
    for row_index, row in enumerate(table.rows, 1):
        for column_index, cell in enumerate(row.cells, 1):
            key = cell._tc
            if key in seen_cells:
                continue
            seen_cells.add(key)
            raw_value = cell.text.strip()
            normalized_value = normalize(raw_value)
            if normalized_value and normalized_value not in normalized_labels:
                values.append(
                    {
                        "row": row_index,
                        "column": column_index,
                        "value": raw_value[:120],
                    }
                )
    return values


def remove_unit(marker_paragraph, table) -> None:
    page_break = marker_paragraph.getprevious()
    table_parent = table._tbl.getparent()
    marker_parent = marker_paragraph.getparent()
    table_parent.remove(table._tbl)
    marker_parent.remove(marker_paragraph)
    if is_page_break_only(page_break) and page_break.getparent() is not None:
        page_break.getparent().remove(page_break)


def resize_kind(
    document,
    kind: str,
    target_count: int,
    prefix: str,
    trim_empty_tail: bool = False,
):
    units = collect_units(document, kind)
    original_count = len(units)
    if original_count == 0:
        raise ValueError(f"文档中未找到可复制的{kind.upper()}表")
    if target_count < original_count:
        if not trim_empty_tail:
            raise ValueError(
                f"{kind.upper()}目标数量{target_count}小于现有数量{original_count}；"
                "默认禁止缩表。如确认只删除末尾完全空白表，请增加 --trim-empty-tail"
            )

        trailing_units = units[target_count:]
        blocked = []
        for original_index, (_, table) in enumerate(
            trailing_units, target_count + 1
        ):
            values = meaningful_table_values(table, kind)
            if values:
                blocked.append(
                    {
                        "code": f"{prefix.upper()}{original_index:02d}",
                        "values": values,
                    }
                )
        if blocked:
            details = []
            for item in blocked:
                evidence = "；".join(
                    f"第{value['row']}行第{value['column']}列={value['value']!r}"
                    for value in item["values"][:3]
                )
                details.append(f"{item['code']}（{evidence}）")
            raise ValueError(
                f"{kind.upper()}缩表已阻断：待删除的末尾表中存在已填写内容："
                + "、".join(details)
                + "。未删除任何表格"
            )

        removed = []
        for index, (marker_paragraph, _) in enumerate(
            trailing_units, target_count + 1
        ):
            removed.append(
                {
                    "position": index,
                    "code": f"{prefix.upper()}{index:02d}",
                    "original_marker": paragraph_text(marker_paragraph).strip(),
                    "reason": "末尾完全空白表",
                }
            )
        for marker_paragraph, table in reversed(trailing_units):
            remove_unit(marker_paragraph, table)

        final_units = collect_units(document, kind)
        renumber_units(final_units, prefix)
        return {
            "before": original_count,
            "requested": target_count,
            "after": len(final_units),
            "operation": "trim",
            "added": [],
            "removed": removed,
            "renumbered_continuously": True,
        }

    renumber_units(units, prefix)
    added = []
    if target_count == original_count:
        return {
            "before": original_count,
            "requested": target_count,
            "after": original_count,
            "operation": "unchanged",
            "added": added,
            "removed": [],
            "renumbered_continuously": True,
        }

    source_paragraph, source_table = units[-1]
    cursor = source_table._tbl
    for index in range(original_count + 1, target_count + 1):
        paragraph_clone = copy.deepcopy(source_paragraph)
        table_clone = copy.deepcopy(source_table._tbl)
        page_break = make_page_break_paragraph()

        cursor.addnext(page_break)
        page_break.addnext(paragraph_clone)
        paragraph_clone.addnext(table_clone)

        table_proxy = next(table for table in document.tables if table._tbl is table_clone)
        clear_table_data(table_proxy, kind)

        code = f"{prefix.upper()}{index:02d}"
        marker = "研发活动编号：" if kind == "rd" else "编号："
        set_paragraph_text_preserve_format(paragraph_clone, f"{marker}{code}")

        cursor = table_clone
        added.append(code)

    final_units = collect_units(document, kind)
    renumber_units(final_units, prefix)
    return {
        "before": original_count,
        "requested": target_count,
        "after": len(final_units),
        "operation": "expand",
        "added": added,
        "removed": [],
        "renumbered_continuously": True,
    }


def parse_args():
    parser = argparse.ArgumentParser(
        description="按高企申请书原格式安全扩缩RD和PS表，并连续编号。"
    )
    parser.add_argument("input", type=Path, help="输入DOCX")
    parser.add_argument("output", type=Path, help="输出DOCX，不得与输入相同")
    parser.add_argument("--rd-count", type=int, help="最终RD表总数")
    parser.add_argument("--ps-count", type=int, help="最终PS表总数")
    parser.add_argument("--rd-prefix", default="RD")
    parser.add_argument("--ps-prefix", default="PS")
    parser.add_argument(
        "--trim-empty-tail",
        action="store_true",
        help="允许从末尾删除完全空白的RD/PS表；任一待删除表有内容即整体阻断",
    )
    parser.add_argument("--report", type=Path, help="可选JSON审计报告")
    parser.add_argument("--overwrite", action="store_true", help="允许覆盖已有输出文件")
    return parser.parse_args()


def main():
    args = parse_args()
    if args.rd_count is None and args.ps_count is None:
        raise SystemExit("必须至少提供 --rd-count 或 --ps-count")
    for label, value in (("RD", args.rd_count), ("PS", args.ps_count)):
        if value is not None and value < 1:
            raise SystemExit(f"{label}数量必须不小于1")

    input_path = args.input.resolve()
    output_path = args.output.resolve()
    if input_path == output_path:
        raise SystemExit("输出文件不得与输入文件相同")
    if not input_path.exists():
        raise SystemExit(f"输入文件不存在：{input_path}")
    if output_path.exists() and not args.overwrite:
        raise SystemExit(f"输出文件已存在；如需覆盖请增加 --overwrite：{output_path}")

    document = Document(input_path)
    report = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "input": str(input_path),
        "output": str(output_path),
        "trim_empty_tail_authorized": args.trim_empty_tail,
        "deletion_policy": (
            "仅允许删除末尾完全空白的RD、PS表；任一待删除表存在填写内容即阻断整个操作"
        ),
        "rd": None,
        "ps": None,
    }

    if args.rd_count is not None:
        report["rd"] = resize_kind(
            document,
            "rd",
            args.rd_count,
            args.rd_prefix,
            trim_empty_tail=args.trim_empty_tail,
        )
    if args.ps_count is not None:
        report["ps"] = resize_kind(
            document,
            "ps",
            args.ps_count,
            args.ps_prefix,
            trim_empty_tail=args.trim_empty_tail,
        )

    trim_performed = any(
        report[kind] is not None and report[kind]["operation"] == "trim"
        for kind in ("rd", "ps")
    )
    report_path = args.report.resolve() if args.report else None
    if trim_performed and report_path is None:
        report_path = output_path.with_suffix(output_path.suffix + ".audit.json")
    if report_path is not None:
        if report_path in (input_path, output_path):
            raise ValueError("审计报告路径不得与输入或输出DOCX相同")
        if report_path.exists() and not args.overwrite:
            raise ValueError(
                f"审计报告已存在；如需覆盖请增加 --overwrite：{report_path}"
            )
        report["audit_report"] = str(report_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    report_text = json.dumps(report, ensure_ascii=False, indent=2)
    if report_path:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(report_text, encoding="utf-8")
    print(report_text)


if __name__ == "__main__":
    try:
        main()
    except ValueError as exc:
        print(f"错误：{exc}", file=sys.stderr)
        raise SystemExit(2)
