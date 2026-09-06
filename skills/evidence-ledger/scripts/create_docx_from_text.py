#!/usr/bin/env python3
"""Create one native DOCX from chat-validated plain or Markdown text."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SCHEMA_VERSION = "gongchuang-docx-generation-operation/v1"
TABLE_DIVIDER = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
NUMBERED_HEADING = re.compile(
    r"^(?:第[一二三四五六七八九十百]+[章节篇部分]|[一二三四五六七八九十]+[、.]|\d+(?:\.\d+)*[、.])\s*(.+)$"
)
LIST_ITEM = re.compile(r"^\s*(?:[-*+]\s+|\d+[.)、]\s*)(.+)$")


def set_run_font(run, *, size: float, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def table_cells(line: str) -> list[str]:
    value = line.strip()
    if value.startswith("|"):
        value = value[1:]
    if value.endswith("|"):
        value = value[:-1]
    return [cell.strip() for cell in value.split("|")]


def add_table(document: Document, lines: list[str]) -> None:
    rows = [table_cells(line) for line in lines if not TABLE_DIVIDER.match(line)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            value = row[column_index] if column_index < len(row) else ""
            run = paragraph.add_run(value)
            set_run_font(run, size=10.5, bold=row_index == 0)


def add_paragraph(document: Document, line: str, *, first_content: bool) -> None:
    heading = HEADING.match(line)
    if heading:
        level = min(len(heading.group(1)), 3)
        paragraph = document.add_heading(level=level)
        run = paragraph.add_run(heading.group(2).strip())
        set_run_font(run, size={1: 16, 2: 14, 3: 12}[level], bold=True)
        return

    numbered = NUMBERED_HEADING.match(line)
    if first_content or numbered:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_content else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(line.strip())
        set_run_font(run, size=18 if first_content else 14, bold=True)
        return

    listed = LIST_ITEM.match(line)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    if listed:
        paragraph.paragraph_format.left_indent = Cm(0.74)
        text = listed.group(1).strip()
        prefix = "• "
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        text = line.strip()
        prefix = ""
    run = paragraph.add_run(prefix + text)
    set_run_font(run, size=11)


def build_document(content: str, output: Path) -> dict[str, object]:
    output = output.resolve()
    if output.exists():
        raise FileExistsError(f"输出文件已存在，拒绝覆盖：{output}")
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    document.core_properties.author = "共创研究院"
    document.core_properties.subject = "共创专业交付"

    lines = content.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    index = 0
    first_content = True
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if "|" in line and index + 1 < len(lines) and TABLE_DIVIDER.match(lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_lines.append(lines[index])
                index += 1
            add_table(document, table_lines)
            first_content = False
            continue
        add_paragraph(document, line, first_content=first_content)
        first_content = False
        index += 1

    document.save(output)
    # 保存后立即用同一原生解析器重开，避免返回一个损坏或空壳 OOXML 文件。
    reopened = Document(output)
    paragraph_count = len(reopened.paragraphs)
    table_count = len(reopened.tables)
    if paragraph_count == 0 and table_count == 0:
        raise ValueError("生成的 DOCX 不包含可解析正文或表格")
    return {
        "schema_version": SCHEMA_VERSION,
        "status": "passed",
        "artifact": str(output),
        "format": "docx",
        "characters": len(content),
        "paragraphs": paragraph_count,
        "tables": table_count,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("content")
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    print(json.dumps(build_document(args.content, args.output), ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
