from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCRIPT = Path(__file__).parents[1] / "scripts" / "create_docx_from_text.py"


def test_creates_a_native_docx_with_one_structured_stdout_record(tmp_path: Path) -> None:
    artifact = tmp_path / "quality-brand-draft.docx"
    content = """制造精品前期评估

# 总体结论
当前材料仅形成待补资料清单。

| 核验项 | 结论 |
|---|---|
| 质量管理 | 待补资料 |

- 请企业补充质量管理体系材料。
"""

    completed = subprocess.run(
        [sys.executable, str(SCRIPT), content, str(artifact)],
        check=True,
        capture_output=True,
        text=True,
    )

    # 受签名运行时只接受一条 JSON；警告或调试输出都必须让测试失败。
    assert completed.stdout.count("\n") == 1
    result = json.loads(completed.stdout)
    assert result["schema_version"] == "gongchuang-docx-generation-operation/v1"
    assert result["characters"] == len(content)
    assert result["tables"] == 1
    generated = Document(artifact)
    assert generated.paragraphs[0].text == "制造精品前期评估"
    assert generated.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert generated.sections[0].top_margin.cm == pytest.approx(2.0, abs=0.001)
    assert generated.sections[0].left_margin.cm == pytest.approx(2.2, abs=0.001)
    assert generated.styles["Normal"].font.size.pt == pytest.approx(10.5)
    assert generated.paragraphs[2].paragraph_format.line_spacing == 1.25
    assert generated.tables[0].cell(1, 0).text == "质量管理"
    table_run = generated.tables[0].cell(1, 0).paragraphs[0].runs[0]
    assert table_run.font.size.pt == pytest.approx(9.5)
    assert generated.tables[0].cell(1, 0).paragraphs[0].paragraph_format.space_after.pt == 0


def test_refuses_to_overwrite_an_existing_docx(tmp_path: Path) -> None:
    artifact = tmp_path / "existing.docx"
    Document().save(artifact)

    completed = subprocess.run(
        [sys.executable, str(SCRIPT), "不会覆盖", str(artifact)],
        capture_output=True,
        text=True,
    )

    assert completed.returncode != 0
    assert "拒绝覆盖" in completed.stderr
