#!/usr/bin/env python3
"""Build one traceable Chinese invention patent application DOCX from JSON."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover - installation failure is explicit
    raise SystemExit(
        "缺少python-docx；请先安装skills/patent-router/requirements.txt"
    ) from exc

import patent_case_manifest as case_manifest


INPUT_SCHEMA = "patent-application-input/v1"
AUDIT_SCHEMA = "patent-application-build-audit/v1"
DRAWING_SCHEMA = "patent-drawing-spec/v1"
PLACEHOLDER_PATTERNS = (
    re.compile(r"\{\{.+?\}\}"),
    re.compile(r"\b(?:TODO|TBD|FIXME)\b", re.IGNORECASE),
    re.compile(r"(?:待补|待填|占位|示例内容|虚构名称)"),
)


def sha256_file(path: Path) -> str:
    return case_manifest.sha256_file(path)


def load_input(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict) or value.get("schema_version") != INPUT_SCHEMA:
        raise ValueError(f"输入必须符合{INPUT_SCHEMA}")
    return value


def require_text(container: dict[str, Any], key: str, label: str) -> str:
    value = str(container.get(key) or "").strip()
    if not value:
        raise ValueError(f"缺少{label}")
    return value


def collect_text(payload: Any) -> str:
    if isinstance(payload, dict):
        return "\n".join(collect_text(value) for value in payload.values())
    if isinstance(payload, list):
        return "\n".join(collect_text(value) for value in payload)
    return str(payload or "")


def validate_input(payload: dict[str, Any]) -> None:
    fact_lock = payload.get("fact_lock")
    if not isinstance(fact_lock, dict) or fact_lock.get("status") != "confirmed":
        raise ValueError("事实锁未确认为confirmed，禁止生成申请文件")
    require_text(payload, "title", "发明名称")
    require_text(payload, "abstract", "摘要")
    claims = payload.get("claims")
    if not isinstance(claims, list) or not claims:
        raise ValueError("至少需要一项权利要求")
    if not all(str(item.get("text") if isinstance(item, dict) else item).strip() for item in claims):
        raise ValueError("权利要求不得为空")
    description = payload.get("description")
    if not isinstance(description, dict):
        raise ValueError("缺少说明书结构")
    for key, label in (
        ("technical_field", "技术领域"),
        ("background", "背景技术"),
        ("technical_problem", "要解决的技术问题"),
        ("technical_solution", "技术方案"),
        ("beneficial_effects", "有益效果"),
        ("drawing_description", "附图说明"),
        ("embodiments", "具体实施方式"),
    ):
        require_text(description, key, label)
    if not payload.get("anonymized_test_fixture"):
        text = collect_text(payload)
        for pattern in PLACEHOLDER_PATTERNS:
            match = pattern.search(text)
            if match:
                raise ValueError(f"正式案卷包含占位内容：{match.group(0)}")


def set_chinese_font(run: Any, size: int = 12, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_chinese_font(run, 16 if level == 1 else 14, True)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)


def add_body(document: Document, text: str) -> None:
    for block in [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(block)
        set_chinese_font(run)


def ensure_inside(root: Path, path_value: str) -> Path:
    candidate = Path(path_value)
    if candidate.is_absolute() or ".." in candidate.parts:
        raise ValueError("附图必须使用输入目录内的相对路径")
    resolved = (root / candidate).resolve()
    try:
        resolved.relative_to(root.resolve())
    except ValueError as exc:
        raise ValueError("附图越出输入目录") from exc
    if not resolved.is_file():
        raise ValueError(f"附图不存在：{path_value}")
    if resolved.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
        raise ValueError(f"不支持的附图格式：{path_value}")
    return resolved


def build_docx(payload: dict[str, Any], input_path: Path, output: Path) -> list[dict[str, Any]]:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title = require_text(payload, "title", "发明名称")
    add_heading(document, title)
    add_heading(document, "摘要")
    add_body(document, require_text(payload, "abstract", "摘要"))

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "权利要求书")
    for index, claim in enumerate(payload["claims"], start=1):
        text = str(claim.get("text") if isinstance(claim, dict) else claim).strip()
        if not re.match(r"^\d+[.．、]", text):
            text = f"{index}. {text}"
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(text)
        set_chinese_font(run)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "说明书")
    description = payload["description"]
    for heading, key in (
        ("技术领域", "technical_field"),
        ("背景技术", "background"),
    ):
        add_heading(document, heading, 2)
        add_body(document, require_text(description, key, heading))
    add_heading(document, "发明内容", 2)
    for heading, key in (
        ("要解决的技术问题", "technical_problem"),
        ("技术方案", "technical_solution"),
        ("有益效果", "beneficial_effects"),
        ("附图说明", "drawing_description"),
        ("具体实施方式", "embodiments"),
    ):
        add_heading(document, heading, 2)
        add_body(document, require_text(description, key, heading))

    drawing_records: list[dict[str, Any]] = []
    drawings = payload.get("drawings") or []
    if drawings:
        if not isinstance(drawings, list):
            raise ValueError("drawings必须是数组")
        document.add_section(WD_SECTION.NEW_PAGE)
        add_heading(document, "说明书附图")
        for index, item in enumerate(drawings, start=1):
            if not isinstance(item, dict):
                raise ValueError("每项附图必须是对象")
            source_value = require_text(item, "path", f"附图{index}路径")
            source = ensure_inside(input_path.parent, source_value)
            caption = str(item.get("caption") or f"图{index}").strip()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(source), width=Cm(14.5))
            caption_paragraph = document.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(caption)
            set_chinese_font(caption_run, 11)
            drawing_records.append(
                {
                    "index": index,
                    "path": source_value,
                    "caption": caption,
                    "sha256": sha256_file(source),
                }
            )

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.widow_control = True
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        suffix=".docx", dir=output.parent, delete=False
    ) as handle:
        temporary = Path(handle.name)
    document.save(temporary)
    temporary.replace(output)
    return drawing_records


def verify_docx(path: Path) -> dict[str, Any]:
    if not zipfile.is_zipfile(path):
        raise ValueError("生成结果不是有效DOCX ZIP")
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        required_parts = {"[Content_Types].xml", "word/document.xml"}
        missing = sorted(required_parts - names)
        if missing:
            raise ValueError(f"DOCX缺少必要部件：{', '.join(missing)}")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    headings = ("摘要", "权利要求书", "说明书", "技术领域", "背景技术", "具体实施方式")
    missing_headings = [heading for heading in headings if heading not in document_xml]
    if missing_headings:
        raise ValueError(f"DOCX缺少必要章节：{', '.join(missing_headings)}")
    return {
        "validation_gate": "patent-application-docx/v1",
        "zip_valid": True,
        "required_headings": list(headings),
        "required_headings_present": True,
    }


def write_json(path: Path, value: dict[str, Any]) -> None:
    case_manifest.atomic_write_json(path, value)


def resolve_case_relative(case_dir: Path, path: Path) -> str:
    try:
        return path.resolve().relative_to(case_dir.resolve()).as_posix()
    except ValueError as exc:
        raise ValueError("输入和输出必须位于案卷目录内") from exc


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--audit-json", type=Path)
    parser.add_argument("--drawing-spec", type=Path)
    parser.add_argument("--case-dir", type=Path)
    parser.add_argument("--contract", type=Path, default=case_manifest.DEFAULT_CONTRACT)
    args = parser.parse_args()
    try:
        input_path = args.input.resolve()
        output_path = args.output.resolve()
        payload = load_input(input_path)
        validate_input(payload)

        manifest: dict[str, Any] | None = None
        contract: dict[str, Any] | None = None
        if args.case_dir:
            case_dir = args.case_dir.resolve()
            manifest = case_manifest.load_manifest(case_dir)
            contract = case_manifest.load_contract(args.contract.resolve())
            if str(payload.get("case_id") or "") != str(manifest.get("case_id") or ""):
                raise ValueError("申请输入的case_id与全案唯一清单不一致")
            if int(payload.get("case_revision") or 0) != int(
                manifest.get("case_revision") or 0
            ):
                raise ValueError("申请输入的case_revision与全案唯一清单不一致")
            if bool(payload.get("anonymized_test_fixture")) != bool(
                manifest.get("anonymized_test_fixture")
            ):
                raise ValueError("输入的测试夹具标记与全案唯一清单不一致")
            input_relative = resolve_case_relative(case_dir, input_path)
            case_manifest.register_artifact(
                case_dir=case_dir,
                manifest=manifest,
                contract=contract,
                role="patent_application_input",
                path_value=input_relative,
                dependencies=None,
            )
            manifest = case_manifest.load_manifest(case_dir)

        drawing_records = build_docx(payload, input_path, output_path)
        validation = verify_docx(output_path)

        if args.drawing_spec:
            drawing_spec = {
                "schema_version": DRAWING_SCHEMA,
                "case_id": payload.get("case_id"),
                "case_revision": payload.get("case_revision"),
                "has_drawings": bool(drawing_records),
                "drawings": drawing_records,
                "validation_gate": "patent-drawing-spec/v1",
                "status": "pass",
            }
            write_json(args.drawing_spec.resolve(), drawing_spec)

        if args.case_dir and manifest is not None and contract is not None:
            case_dir = args.case_dir.resolve()
            output_relative = resolve_case_relative(case_dir, output_path)
            case_manifest.register_artifact(
                case_dir=case_dir,
                manifest=manifest,
                contract=contract,
                role="patent_application_docx",
                path_value=output_relative,
                dependencies=None,
            )
            manifest = case_manifest.load_manifest(case_dir)
            if args.drawing_spec:
                drawing_relative = resolve_case_relative(
                    case_dir, args.drawing_spec.resolve()
                )
                case_manifest.register_artifact(
                    case_dir=case_dir,
                    manifest=manifest,
                    contract=contract,
                    role="drawing_spec",
                    path_value=drawing_relative,
                    dependencies=None,
                )

        audit = {
            "schema_version": AUDIT_SCHEMA,
            "input_schema": INPUT_SCHEMA,
            "input_sha256": sha256_file(input_path),
            "output_sha256": sha256_file(output_path),
            "anonymized_test_fixture": bool(payload.get("anonymized_test_fixture")),
            "fact_lock_status": payload["fact_lock"]["status"],
            "claim_count": len(payload["claims"]),
            "drawing_count": len(drawing_records),
            **validation,
            "status": "pass",
        }
        if args.audit_json:
            write_json(args.audit_json.resolve(), audit)
        print(json.dumps(audit, ensure_ascii=False, indent=2))
        return 0
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        print(
            json.dumps(
                {
                    "schema_version": AUDIT_SCHEMA,
                    "status": "fail",
                    "completion_allowed": False,
                    "error": str(exc),
                },
                ensure_ascii=False,
                indent=2,
            )
        )
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
