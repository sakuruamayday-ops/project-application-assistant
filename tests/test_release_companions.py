from __future__ import annotations

import importlib.util
import json
from pathlib import Path

from docx import Document


SCRIPT = Path(__file__).parents[1] / "scripts" / "release_companions.py"
SPEC = importlib.util.spec_from_file_location("release_companions", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(MODULE)


def fixture_root(tmp_path: Path) -> Path:
    root = tmp_path / "repository"
    (root / "skills").mkdir(parents=True)
    (root / "docs/user-guide").mkdir(parents=True)
    manifest = {
        "product_name": "企业全生命周期助手",
        "release": {
            "tag": "V2.0.1",
            "version": "2.0.1",
            "summary": "修复接入流程。",
            "changes": ["新增手工接入兜底。"],
            "installation": {
                "generic": "导入通用包。",
                "workbuddy": "确认后安装插件包。",
                "manual_fallback": "复制手工配置。",
                "verification": "完成五步验收。",
            },
            "compatibility": [
                {"target": "macOS", "status": "已验证", "note": "实机通过。"}
            ],
            "rollback": ["恢复上一正式版本。"],
        },
        "skills": ["a", "b"],
        "release_companions": {
            "manual_template": "docs/user-guide/manual.docx",
            "manual_filename": "manual-{tag}.docx",
            "companion_filename": "companions-{tag}.json",
            "delivery_directory": "Desktop/焦糖待处理",
            "word_manual_only": True,
            "require_branding": True,
            "require_render_qa": True,
        },
    }
    (root / "skills/suite-manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False),
        encoding="utf-8",
    )
    document = Document()
    document.add_paragraph("企业全生命周期助手用户使用手册")
    document.add_paragraph("适用版本：V1.9")
    document.add_paragraph("本手册由9个平台无关Skills组成。")
    document.add_paragraph("下载V1.9通用包或WorkBuddy插件包。")
    document.save(root / "docs/user-guide/manual.docx")
    return root


def test_generate_uses_manifest_as_the_only_release_source(tmp_path: Path) -> None:
    root = fixture_root(tmp_path)
    output = tmp_path / "output"

    result = MODULE.generate(root, output, apply_brand=False, render=False)

    manual = Path(result["manual"])
    text = MODULE.extracted_text(manual)
    assert "适用版本：V2.0.1" in text
    assert "由2个平台无关Skills" in text
    assert "修复接入流程。" in text
    assert "新增手工接入兜底。" in text
    assert "导入通用包。" in text
    assert "macOS：已验证" in text
    companion = json.loads(Path(result["companion"]).read_text(encoding="utf-8"))
    assert companion["release_tag"] == "V2.0.1"
    assert companion["skill_count"] == 2
    assert companion["source_of_truth"] == "skills/suite-manifest.json"


def test_validate_manual_fails_when_manifest_fact_is_missing(tmp_path: Path) -> None:
    root = fixture_root(tmp_path)
    spec = MODULE.release_spec(MODULE.load_manifest(root / "skills/suite-manifest.json"))
    stale = root / "docs/user-guide/manual.docx"

    try:
        MODULE.validate_manual_content(stale, spec)
    except ValueError as error:
        assert "Word 手册缺少清单事实" in str(error)
    else:
        raise AssertionError("stale manual should be rejected")


def test_deliver_copies_only_word_manual_and_machine_audit(tmp_path: Path) -> None:
    root = fixture_root(tmp_path)
    output = tmp_path / "output"
    MODULE.generate(root, output, apply_brand=False, render=False)
    destination = tmp_path / "delivery"

    result = MODULE.deliver(root, output, destination)

    assert result["status"] == "pass"
    assert sorted(path.name for path in destination.iterdir()) == [
        "companions-V2.0.1.json",
        "manual-V2.0.1.docx",
    ]


def test_generated_word_manual_is_byte_deterministic(tmp_path: Path) -> None:
    root = fixture_root(tmp_path)
    first = MODULE.generate(
        root,
        tmp_path / "first",
        apply_brand=False,
        render=False,
    )
    second = MODULE.generate(
        root,
        tmp_path / "second",
        apply_brand=False,
        render=False,
    )

    assert MODULE.sha256(Path(first["manual"])) == MODULE.sha256(
        Path(second["manual"])
    )
