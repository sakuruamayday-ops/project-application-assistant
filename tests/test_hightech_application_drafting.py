from __future__ import annotations

import importlib.util
import copy
import json
import struct
import subprocess
import sys
import zipfile
import zlib
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SKILL_ROOT = ROOT / "skills" / "high-tech-enterprise-application-drafting"
SCRIPT = SKILL_ROOT / "scripts" / "expand_rd_ps_tables.py"
FILL_SCRIPT = SKILL_ROOT / "scripts" / "fill_rd_core_innovation.py"
AUDIT_SCRIPT = SKILL_ROOT / "scripts" / "audit_application_docx.py"
INPUT_PLAN_SCRIPT = SKILL_ROOT / "scripts" / "plan_hightech_inputs.py"
XLS_FIXTURE = ROOT / "skills" / "project-application-assistant" / "tests" / "fixtures" / "document-extraction-sample.xls"
DELIVERY_GATE = SKILL_ROOT / "scripts" / "hightech_delivery_gate.py"
TEMPLATE = SKILL_ROOT / "assets" / "高新技术企业认定申请书空白模板.docx"
SPEC = importlib.util.spec_from_file_location("expand_rd_ps_tables", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(MODULE)
FILL_SPEC = importlib.util.spec_from_file_location("fill_rd_core_innovation", FILL_SCRIPT)
FILL_MODULE = importlib.util.module_from_spec(FILL_SPEC)
assert FILL_SPEC and FILL_SPEC.loader
FILL_SPEC.loader.exec_module(FILL_MODULE)
AUDIT_SPEC = importlib.util.spec_from_file_location("audit_application_docx", AUDIT_SCRIPT)
AUDIT_MODULE = importlib.util.module_from_spec(AUDIT_SPEC)
assert AUDIT_SPEC and AUDIT_SPEC.loader
AUDIT_SPEC.loader.exec_module(AUDIT_MODULE)


FOUR_LEVEL_FIELD = "五、高技术服务—一、研发与设计服务—2、涉及服务—专业设计技术"


def set_first_rd_field(document: Document) -> None:
    MODULE.resize_kind(document, "rd", 1, "RD")
    for table in document.tables:
        labels = {"".join(cell.text.split()) for row in table.rows for cell in row.cells}
        if not all(any(value.startswith(label) for value in labels) for label in (
            "研发活动名称",
            "技术领域",
            "核心技术及创新点",
        )):
            continue
        for row in table.rows:
            cells = list(MODULE.unique_cells(row))
            for index, cell in enumerate(cells):
                if "".join(cell.text.split()).startswith("技术领域") and index + 1 < len(cells):
                    cells[index + 1].text = FOUR_LEVEL_FIELD
                    for paragraph in cells[index + 1].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "宋体"
                            run.font.size = Pt(12)
                            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
                    return
    raise AssertionError("空白模板中未找到RD技术领域单元格")


def realistic_spec() -> dict:
    return {
        "schema_version": 1,
        "enterprise_profile": {
            "name": "某传统制造中小企业",
            "scale_summary": "中小型传统制造企业，以通用加工设备和人工辅助装配为主",
            "main_business": "精密五金部件及配套治具的设计、加工与交付",
            "patent_summary": "当前资料显示企业专利集中在治具结构、定位机构和防护组件",
            "scale_evidence": ["企业提供的员工及设备清单，核验日期2026-08-11"],
            "patent_evidence": ["企业提供的知识产权汇总表，核验日期2026-08-11"],
            "verified_advanced_terms": [],
        },
        "items": [
            {
                "rd_id": "RD01",
                "core_technologies": [
                    {
                        "name": "弹性浮动夹持与重复定位技术",
                        "description": "采用分区弹片、限位台阶和可调预紧结构形成柔性夹持，以底板定位孔和设备定位柱控制装夹基准，减少薄壁零件翘曲与边缘压痕",
                        "indicators": "重复装夹位置偏差控制在±0.20毫米以内，单件装夹时间不超过45秒，边缘可见压痕发生率不高于1%",
                    },
                    {
                        "name": "模块化遮挡防护与换型技术",
                        "description": "将通用载具与可拆遮挡片分体设计，用卡扣和定位销快速安装，通过孔位防错、首件确认和末件复核闭合换型记录",
                        "indicators": "不同型号换型时间不超过30分钟，遮挡边界偏差不大于0.30毫米，通用载具复用率不低于70%",
                    },
                ],
                "innovations": [
                    "将弹性夹持、基准定位和双面作业集成，在不改造主体设备的前提下减少翻面二次找正和薄壁局部受力",
                    "将遮挡件与载具主体分离，通过更换低成本薄片完成换型，并用防错与首件记录减少型号混用和遮挡不到位",
                ],
                "stage_results": [
                    "项目已完成结构方案与试制验证",
                    "已形成治具图纸、工艺参数表和样件记录",
                    "知识产权情况按企业提供汇总表填写",
                    "已具备后续产品导入和小批试用基础",
                ],
            }
        ],
    }


def test_rd_tables_expand_and_empty_tail_trim() -> None:
    document = Document(TEMPLATE)
    expansion = MODULE.resize_kind(document, "rd", 20, "RD")
    assert expansion["operation"] == "expand"
    assert expansion["before"] == 1
    assert expansion["after"] == 20
    assert expansion["renumbered_continuously"] is True

    trimming = MODULE.resize_kind(
        document,
        "rd",
        13,
        "RD",
        trim_empty_tail=True,
    )
    assert trimming["operation"] == "trim"
    assert trimming["before"] == 20
    assert trimming["after"] == 13
    assert [item["code"] for item in trimming["removed"]] == [
        "RD14",
        "RD15",
        "RD16",
        "RD17",
        "RD18",
        "RD19",
        "RD20",
    ]
    assert trimming["renumbered_continuously"] is True


def test_trim_blocks_when_trailing_table_contains_content() -> None:
    document = Document(TEMPLATE)
    MODULE.resize_kind(document, "rd", 3, "RD")
    units = MODULE.collect_units(document, "rd")
    assert len(units) == 3
    units[-1][1].rows[0].cells[-1].text = "已填写项目名称"

    with pytest.raises(ValueError, match="缩表已阻断"):
        MODULE.resize_kind(
            document,
            "rd",
            2,
            "RD",
            trim_empty_tail=True,
        )

    assert len(MODULE.collect_units(document, "rd")) == 3


def test_cli_trim_generates_default_audit_report(tmp_path: Path) -> None:
    expanded = tmp_path / "expanded.docx"
    output = tmp_path / "trimmed.docx"
    expansion = subprocess.run(
        [
            sys.executable,
            str(SCRIPT),
            str(TEMPLATE),
            str(expanded),
            "--rd-count",
            "18",
            "--ps-count",
            "3",
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert expansion.returncode == 0, expansion.stderr
    process = subprocess.run(
        [
            sys.executable,
            str(SCRIPT),
            str(expanded),
            str(output),
            "--rd-count",
            "17",
            "--ps-count",
            "2",
            "--trim-empty-tail",
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert process.returncode == 0, process.stderr
    audit = output.with_suffix(output.suffix + ".audit.json")
    assert output.is_file()
    assert audit.is_file()
    report = json.loads(audit.read_text(encoding="utf-8"))
    assert report["trim_empty_tail_authorized"] is True
    assert report["rd"]["after"] == 17
    assert report["ps"]["after"] == 2


def test_fill_rd_core_innovation_uses_exact_field_and_seven_line_format(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "filled.docx"
    report_path = tmp_path / "filled.audit.json"
    spec_path = tmp_path / "spec.json"
    document = Document(TEMPLATE)
    set_first_rd_field(document)
    document.save(source)
    spec_path.write_text(json.dumps(realistic_spec(), ensure_ascii=False, indent=2), encoding="utf-8")

    process = subprocess.run(
        [sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output), "--report", str(report_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert process.returncode == 0, process.stderr
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["status"] == "pass"
    assert report["rd"]["RD01"]["body_length"] <= 400
    assert report["rd"]["RD01"]["core_technology_count"] == 2
    assert report["rd"]["RD01"]["innovation_count"] == 2
    assert report["rd"]["RD01"]["stage_result_count"] == 4
    assert report["rd"]["RD01"]["stage_result_length"] <= 400
    assert report["postcondition_verified"] is True
    assert report["enterprise"]["scale_evidence_count"] == 1
    assert report["enterprise"]["patent_evidence_count"] == 1

    rd_result, issues = AUDIT_MODULE.audit_rd_core_innovation(Document(output))
    assert issues == []
    assert rd_result["RD01"]["field"] == FOUR_LEVEL_FIELD
    assert rd_result["RD01"]["line_count"] == 7
    assert rd_result["RD01"]["body_length"] <= 400
    assert rd_result["RD01"]["core_technology_count"] == 2
    assert rd_result["RD01"]["innovation_count"] == 2
    stage_result, stage_issues = AUDIT_MODULE.audit_rd_stage_results(Document(output))
    assert stage_issues == []
    assert stage_result["RD01"]["line_count"] == 4


def test_fill_basic_fields_handles_three_merged_rd_tables_without_touching_later_tables(tmp_path: Path) -> None:
    source, output, spec_path = (tmp_path / name for name in ("source.docx", "filled.docx", "spec.json"))
    document = Document(TEMPLATE)
    MODULE.resize_kind(document, "rd", 3, "RD")
    document.save(source)
    data = realistic_spec()
    seed = data["items"][0]
    data["items"] = []
    for number in range(1, 4):
        item = copy.deepcopy(seed)
        item["rd_id"] = f"RD{number:02}"
        item["basic_fields"] = {
            "研发活动名称": f"合成项目{number}", "起止时间": "2025.05.01-2025.8.31",
            "技术领域": FOUR_LEVEL_FIELD, "技术来源": "企业自述研发",
            "知识产权编号": f"TEST-P{number:02}", "研发经费总预算": str(number * 10),
            "目的及组织实施方式": "研究现有资料记载的技术。组织方式待企业核定。",
        }
        data["items"].append(item)
    spec_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    process = subprocess.run([sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output)], capture_output=True, text=True)
    assert process.returncode == 0, process.stderr
    report = json.loads(process.stdout)
    assert report["status"] == "pass"
    assert report["rd_count"] == 3
    with zipfile.ZipFile(source) as before, zipfile.ZipFile(output) as after:
        for name in before.namelist():
            if name != "word/document.xml":
                assert after.read(name) == before.read(name)
        original = FILL_MODULE.etree.fromstring(before.read("word/document.xml"))
        filled = FILL_MODULE.etree.fromstring(after.read("word/document.xml"))
    targets = FILL_MODULE.collect_rd_targets(filled)
    assert list(targets) == ["RD01", "RD02", "RD03"]
    for number, (rd_id, target) in enumerate(targets.items(), 1):
        cells = target["basic_cells"]
        assert FILL_MODULE.element_text(cells["研发活动名称"]) == f"合成项目{number}"
        assert FILL_MODULE.element_text(cells["起止时间"]) == "2025.05.01-2025.08.31"
        assert FILL_MODULE.element_text(cells["研发经费总预算"]) == str(number * 10)
        assert report["rd"][rd_id]["basic_fields"]["知识产权编号"] == f"TEST-P{number:02}"
    # 后续PS、创新能力及汇总表保持逐节点一致；RD表格属性和合并关系也不改变。
    old_tables = original.xpath(".//w:body/w:tbl", namespaces=FILL_MODULE.NS)
    new_tables = filled.xpath(".//w:body/w:tbl", namespaces=FILL_MODULE.NS)
    assert len(old_tables) == len(new_tables)
    for old, new in zip(old_tables, new_tables):
        rows = FILL_MODULE.table_rows(old)
        if not FILL_MODULE.find_label_row(rows, "研发活动名称"):
            assert FILL_MODULE.etree.tostring(old) == FILL_MODULE.etree.tostring(new)
        else:
            properties = ".//w:tblPr | .//w:tblGrid | .//w:trPr | .//w:tcPr"
            assert [FILL_MODULE.etree.tostring(n) for n in old.xpath(properties, namespaces=FILL_MODULE.NS)] == [FILL_MODULE.etree.tostring(n) for n in new.xpath(properties, namespaces=FILL_MODULE.NS)]


def test_fill_explicit_unknowns_produces_draft_and_formal_audit_keeps_gaps(tmp_path: Path) -> None:
    spec_path, output = tmp_path / "spec.json", tmp_path / "draft.docx"
    source = tmp_path / "source.docx"
    document = Document(TEMPLATE)
    MODULE.resize_kind(document, "rd", 1, "RD")
    document.save(source)
    data = realistic_spec()
    item = data["items"][0]
    item["basic_fields"] = {"技术领域": "待企业确认四级领域"}
    for technology in item["core_technologies"]:
        technology["indicators"] = None
    spec_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    process = subprocess.run([sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output)], capture_output=True, text=True)
    assert process.returncode == 0, process.stderr
    report = json.loads(process.stdout)
    assert report["status"] == "draft"
    assert report["rd"]["RD01"]["pending_indicators"] == 2
    assert report["rd"]["RD01"]["pending_field"] is True
    results, issues = AUDIT_MODULE.audit_rd_core_innovation(Document(output))
    assert len(issues) == 3
    assert any("四级技术领域待确认" in issue["issue"] for issue in issues)
    assert results["RD01"]["line_count"] == 7


@pytest.mark.parametrize("fields", [
    {"起止时间": "2025.09.31-2025.10.01"},
    {"起止时间": "2025.09.01-2025.08.31"},
    {"未定义栏目": "不得写入"},
])
def test_fill_rejects_invalid_basic_fields_before_creating_output(tmp_path: Path, fields: dict) -> None:
    spec_path, output = tmp_path / "spec.json", tmp_path / "filled.docx"
    source = tmp_path / "source.docx"
    document = Document(TEMPLATE)
    MODULE.resize_kind(document, "rd", 1, "RD")
    document.save(source)
    data = realistic_spec()
    data["items"][0]["basic_fields"] = fields
    spec_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    process = subprocess.run([sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output)], capture_output=True, text=True)
    assert process.returncode != 0
    assert "文档中未找到目标RD表" not in process.stderr
    assert not output.exists()


def test_fill_blocks_unverified_advanced_terms_before_output(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "blocked.docx"
    spec_path = tmp_path / "spec.json"
    document = Document(TEMPLATE)
    set_first_rd_field(document)
    document.save(source)
    data = realistic_spec()
    data["items"][0]["innovations"][0] += "项目同步建设数字孪生系统。"
    spec_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

    process = subprocess.run(
        [sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert process.returncode != 0
    assert "缺少企业能力证据的高阶技术词" in process.stderr
    assert not output.exists()


def test_fill_rejects_core_body_over_400_without_creating_output(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "blocked.docx"
    spec_path = tmp_path / "spec.json"
    document = Document(TEMPLATE)
    set_first_rd_field(document)
    document.save(source)
    data = realistic_spec()
    data["items"][0]["innovations"][1] += "超" * 401
    spec_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

    process = subprocess.run(
        [sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output)],
        check=False,
        capture_output=True,
        text=True,
    )

    assert process.returncode != 0
    assert "超过400字上限" in process.stderr
    assert not output.exists()


def test_fill_requires_scale_and_patent_evidence(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "blocked.docx"
    spec_path = tmp_path / "spec.json"
    document = Document(TEMPLATE)
    set_first_rd_field(document)
    document.save(source)
    data = realistic_spec()
    data["enterprise_profile"]["patent_evidence"] = []
    spec_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")

    process = subprocess.run(
        [sys.executable, str(FILL_SCRIPT), str(source), str(spec_path), str(output)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert process.returncode != 0
    assert "patent_evidence必须是非空字符串数组" in process.stderr
    assert not output.exists()


@pytest.mark.parametrize(("length", "passes"), [(399, True), (400, True), (401, False)])
def test_limit_400_is_a_maximum_boundary(length: int, passes: bool) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "目的及组织实施方式（限400字）"
    table.cell(0, 2).text = "甲" * length

    fields, issues = AUDIT_MODULE.audit_limited_fields(document)

    assert fields[0]["length"] == length
    assert (issues == []) is passes
    if not passes:
        assert "超过400字上限" in issues[0]["issue"]


def test_physical_merged_cells_preserve_semantic_label_and_value_pairing() -> None:
    document = Document()
    labels = list(AUDIT_MODULE.INNOVATION_SECTION_ORDER)
    for label in labels:
        table = document.add_table(rows=1, cols=3)
        table.cell(0, 0).merge(table.cell(0, 1)).text = f"{label}（限400字）"
        table.cell(0, 2).text = "已核定事实"

    result, issues = AUDIT_MODULE.audit_innovation_capability(document)

    assert issues == []
    assert list(result) == labels
    assert all(item["text"] == "已核定事实" for item in result.values())


def run_input_plan(manifest: Path, state: Path, report: Path) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [
            sys.executable,
            str(INPUT_PLAN_SCRIPT),
            str(manifest),
            "--state",
            str(state),
            "--report",
            str(report),
        ],
        check=False,
        capture_output=True,
        text=True,
    )


def write_input_manifest(path: Path, documents: list[dict[str, str]]) -> None:
    path.write_text(
        json.dumps({"schema_version": 1, "documents": documents}, ensure_ascii=False),
        encoding="utf-8",
    )


def write_plan_docx(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '</Types>',
        )
        archive.writestr("word/document.xml", "<document/>")


def write_plan_ole_doc(path: Path) -> None:
    data = bytearray(XLS_FIXTURE.read_bytes())
    needle = "Workbook\0".encode("utf-16le")
    position = data.find(needle)
    assert position >= 0
    entry_start = position - (position % 128)
    data[entry_start : entry_start + 64] = b"\0" * 64
    name = "WordDocument\0".encode("utf-16le")
    data[entry_start : entry_start + len(name)] = name
    struct.pack_into("<H", data, entry_start + 64, len(name))
    path.write_bytes(data)


def test_input_plan_allows_one_signed_extraction_attempt_for_legacy_doc(tmp_path: Path) -> None:
    legacy = tmp_path / "支撑专利.doc"
    write_plan_ole_doc(legacy)
    application = tmp_path / "申请书.docx"
    write_plan_docx(application)
    original = legacy.read_bytes()
    manifest = tmp_path / "inputs.json"
    state = tmp_path / "state.json"
    report = tmp_path / "report.json"
    write_input_manifest(
        manifest,
        [
            {"path": legacy.name, "role": "supporting"},
            {"path": application.name, "role": "essential"},
        ],
    )

    process = run_input_plan(manifest, state, report)

    assert process.returncode == 0, process.stderr
    payload = json.loads(report.read_text(encoding="utf-8"))
    assert payload["status"] == "ready"
    assert payload["supported_count"] == 2
    assert payload["documents"][0]["status"] == "legacy_doc_probe"
    assert payload["documents"][0]["action"] == "extract_once_then_convert_if_required"
    assert "COM" in payload["documents"][0]["reason"]
    assert payload["originals_modified"] is False
    assert legacy.read_bytes() == original

    repeated = run_input_plan(manifest, state, report)
    repeated_payload = json.loads(report.read_text(encoding="utf-8"))
    assert repeated.returncode == 3
    assert repeated_payload["status"] == "stopped_no_progress"
    assert json.loads(state.read_text(encoding="utf-8"))["retry_allowed"] is False


def test_input_plan_stops_second_unchanged_proprietary_wps_failure(tmp_path: Path) -> None:
    legacy = tmp_path / "必需资料.wps"
    legacy.write_bytes(b"\x01WPS-PROPRIETARY\x00\xff")
    original = legacy.read_bytes()
    manifest = tmp_path / "inputs.json"
    state = tmp_path / "state.json"
    report = tmp_path / "report.json"
    write_input_manifest(manifest, [{"path": legacy.name, "role": "essential"}])

    first = run_input_plan(manifest, state, report)
    first_payload = json.loads(report.read_text(encoding="utf-8"))
    second = run_input_plan(manifest, state, report)
    second_payload = json.loads(report.read_text(encoding="utf-8"))

    assert first.returncode == 2
    assert first_payload["status"] == "blocked_conversion_required"
    assert second.returncode == 3
    assert second_payload["status"] == "stopped_no_progress"
    assert json.loads(state.read_text(encoding="utf-8"))["retry_allowed"] is False
    assert legacy.read_bytes() == original


def test_hightech_input_plan_declares_shared_detector_packaging_dependency() -> None:
    manifest = json.loads((ROOT / "skills" / "suite-manifest.json").read_text(encoding="utf-8"))
    dependency = manifest["dependencies"]["high-tech-enterprise-application-drafting"]
    assert "project-application-assistant" in dependency["required_skills"]
    assert "project-application-assistant/scripts/document_format_detection.py" in dependency["required_paths"]


def make_innovation_document(texts: list[str]) -> Document:
    document = Document()
    table = document.add_table(rows=4, cols=2)
    labels = [
        "知识产权对企业\n竞争力的作用（限400字）",
        "科技成果转化情况\n（限400字）",
        "研究开发与技术创新组织管理情况（限400字）",
        "管理与科技人员情况\n（限400字）",
    ]
    for index, (label, value) in enumerate(zip(labels, texts, strict=True)):
        table.rows[index].cells[0].text = label
        table.rows[index].cells[1].text = value
    return document


def pad_innovation_text(seed: str) -> str:
    supplement = (
        "公司围绕现有产品和生产环节建立研发、成果、知识产权与人员资料的对应关系，"
        "各项主张以本企业提供的汇总表、制度原件、项目记录和人员材料为核验边界，"
        "不能确认的经营成效和执行结果留待企业补充记录后再写入正式申请材料。"
    )
    return seed + supplement * 3


def innovation_evidence() -> dict:
    return {
        "allowed_corporate_names": ["某制造有限公司"],
        "policy_titles": [
            "关于成立企业研发中心的通知",
            "科研项目研究开发组织管理制度",
            "研发经费投入核算制度",
        ],
        "allowed_personnel_claims": ["现有科技人员8名"],
        "allowed_result_claims": ["近三年形成10项科技成果转化记录"],
        "allow_financing_claims": False,
    }


def test_innovation_capability_accepts_fixed_order_and_evidence_boundaries() -> None:
    texts = [
        pad_innovation_text(
            "某制造有限公司围绕主营产品保护关键结构和工艺成果，已登记成果与RD及PS建立对应关系。"
            "另有专利申请处于受理阶段，仅作为审中成果列示，不计入已授权数量。"
        ),
        pad_innovation_text(
            "近三年形成10项科技成果转化记录，成果通过企业自主投资实施转化并导入现有生产环节。"
        ),
        pad_innovation_text(
            "公司依据《关于成立企业研发中心的通知》《科研项目研究开发组织管理制度》"
            "和《研发经费投入核算制度》明确研发机构、项目及经费管理规则。"
        ),
        pad_innovation_text(
            "人员材料能够独立核定现有科技人员8名，岗位覆盖研发组织、设备、工艺、测试和辅助环节。"
        ),
    ]
    result, issues = AUDIT_MODULE.audit_innovation_capability(
        make_innovation_document(texts), innovation_evidence()
    )
    assert issues == []
    assert list(result) == list(AUDIT_MODULE.INNOVATION_SECTION_ORDER)
    assert all(item["length"] <= 400 for item in result.values())


def test_innovation_capability_blocks_cross_enterprise_and_unverified_claims() -> None:
    texts = [
        pad_innovation_text(
            "另一示例科技有限公司拥有申请受理专利并已形成授权成果，相关知识产权提高了企业融资估值。"
        ),
        pad_innovation_text(
            "成果已在10余个大型工程中应用，用户高度认可，营业收入增长率超过30%，并获得科技项目资助。"
        ),
        pad_innovation_text(
            "公司依据《通用科研管理制度》开展项目管理，但当前企业详细制度文件未提供该名称。"
        ),
        pad_innovation_text(
            "公司拥有9名科技人员，包括博士和硕士，核心成员留任率较高，能够持续支持研发工作。"
        ),
    ]
    _, issues = AUDIT_MODULE.audit_innovation_capability(
        make_innovation_document(texts), innovation_evidence()
    )
    joined = "\n".join(item["issue"] for item in issues)
    assert "未在本企业证据中登记的企业名称" in joined
    assert "写成授权成果" in joined
    assert "融资作用证据" in joined
    assert "未登记的市场、合作或政策主张" in joined
    assert "制度名称未逐字命中" in joined
    assert "未登记的人员结构或稳定性主张" in joined
    assert "未登记的数量或比例：9名" in joined


def run_gate(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(DELIVERY_GATE), *args],
        check=False,
        capture_output=True,
        text=True,
    )


def make_ps_summary(path: Path, *, prohibited: bool = False) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=9)
    headers = [
        "编号",
        "高新技术产品（服务）名称",
        "高新收入",
        "关键技术",
        "所属高新领域",
        "对应的RD",
        "知识产权名称",
        "技术指标",
        "证明材料",
    ]
    values = [
        "1",
        "精密部件表面处理服务",
        "" if not prohibited else "XXX万元，待核定",
        "异常联锁技术\n清洗除尘技术",
        "八、先进制造与自动化/（四）先进制造工艺与装备/4、特种加工技术",
        "RD01\nRD02",
        "IP01\nIP02",
        "报警响应≤3秒\n颗粒去除率≥95%",
        "知识产权\n检测报告\n销售合同",
    ]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for index, value in enumerate(values):
        table.rows[1].cells[index].text = value
    document.save(path)


def make_result_summary(path: Path, *, split_tables: bool = False) -> None:
    document = Document()
    headers = [
        "序号",
        "成果名称",
        "成果来源",
        "转化方式",
        "转化目标产品",
        "转化时间",
        "涉及关键技术",
        "转化所取得成效",
        "关联项目RD编号",
        "关联专利IP编号",
        "成果转化证明材料",
    ]
    rows = [
        ["1", "清洗设备", "自有技术", "自行投资实施转化", "新技术应用", "2025", "多阶段清洗技术", "形成一体化清洗设备。", "RD01", "IP01", "知识产权\n销售合同"],
        ["2", "除尘装置", "自有技术", "自行投资实施转化", "新技术应用", "2025", "离子风除尘技术", "形成进炉前除尘装置。", "RD02", "IP02", "知识产权\n检测记录"],
    ]
    groups = [[item] for item in rows] if split_tables else [rows]
    for group in groups:
        table = document.add_table(rows=1 + len(group), cols=len(headers))
        for index, value in enumerate(headers):
            table.rows[0].cells[index].text = value
        for row_index, row in enumerate(group, 1):
            for column_index, value in enumerate(row):
                table.rows[row_index].cells[column_index].text = value
    document.save(path)


def make_ip_summary(path: Path, *, bad_number: bool = False) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=6)
    headers = ["知识产权编号", "知识产权名称", "类别", "授权日期", "授权号", "获得方式"]
    rows = [
        ["IP01", "一种定位治具", "实用新型", "2025-03-01", "ZL2024XXXXXX.X", "自主研发"],
        ["IP03" if bad_number else "IP02", "工艺参数管理软件V1.0", "软件著作权", "2025-05-01", "2025SRXXXXXX", "自主研发"],
    ]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row_index, row in enumerate(rows, 1):
        for column_index, value in enumerate(row):
            table.rows[row_index].cells[column_index].text = value
    document.save(path)


def test_summary_lint_accepts_compact_ps_and_continuous_results(tmp_path: Path) -> None:
    ps = tmp_path / "ps.docx"
    results = tmp_path / "results.docx"
    ip = tmp_path / "ip.docx"
    make_ps_summary(ps)
    make_result_summary(results)
    make_ip_summary(ip)
    for source in (ps, results, ip):
        report = source.with_suffix(".audit.json")
        process = run_gate("summary-lint", str(source), "--report", str(report))
        assert process.returncode == 0, process.stdout + process.stderr
        payload = json.loads(report.read_text(encoding="utf-8"))
        assert payload["status"] == "pass"
        assert payload["tables"]


def test_summary_lint_blocks_placeholder_and_split_result_tables(tmp_path: Path) -> None:
    ps = tmp_path / "bad-ps.docx"
    results = tmp_path / "split-results.docx"
    make_ps_summary(ps, prohibited=True)
    make_result_summary(results, split_tables=True)

    ps_process = run_gate("summary-lint", str(ps), "--report", str(tmp_path / "bad-ps.json"))
    assert ps_process.returncode != 0
    assert "禁用占位语" in ps_process.stdout

    results_process = run_gate(
        "summary-lint", str(results), "--report", str(tmp_path / "split-results.json")
    )
    assert results_process.returncode != 0
    assert "科技成果转化汇总表被拆成" in results_process.stdout

    ip = tmp_path / "bad-ip.docx"
    make_ip_summary(ip, bad_number=True)
    ip_process = run_gate("summary-lint", str(ip), "--report", str(tmp_path / "bad-ip.json"))
    assert ip_process.returncode != 0
    assert "IP编号不连续" in ip_process.stdout


def test_template_copy_requires_byte_identical_initial_copy(tmp_path: Path) -> None:
    copied = tmp_path / "copy.docx"
    copied.write_bytes(TEMPLATE.read_bytes())
    report = tmp_path / "copy.json"
    passing = run_gate(
        "template-copy", str(TEMPLATE), str(copied), "--report", str(report)
    )
    assert passing.returncode == 0, passing.stdout + passing.stderr
    copied.write_bytes(copied.read_bytes() + b"changed")
    failing = run_gate(
        "template-copy", str(TEMPLATE), str(copied), "--report", str(report)
    )
    assert failing.returncode != 0
    assert "字节级复制" in failing.stdout


def wps_checklist(path: Path, screenshots: list[Path]) -> None:
    checks = {
        "header_footer": "pass",
        "table_boundaries": "pass",
        "repeated_header": "pass",
        "overflow": "pass",
        "overlap": "pass",
        "blank_page": "pass",
        "continuous_numbering": "pass",
        "missing_fields": "pass",
    }
    path.write_text(
        json.dumps(
            {
                "engine": "WPS Office",
                "reviewer": "Codex",
                "page_count": len(screenshots),
                "pages": [
                    {"page": index, "screenshot": screenshot.name, "checks": checks}
                    for index, screenshot in enumerate(screenshots, 1)
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def write_png(path: Path, width: int = 800, height: int = 600, color: bytes = b"\xff\xff\xff") -> None:
    def chunk(kind: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + kind + data + struct.pack(">I", zlib.crc32(kind + data))

    scanline = b"\x00" + color * width
    pixels = scanline * height
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(pixels))
        + chunk(b"IEND", b"")
    )


def test_wps_review_and_finalize_are_bound_to_final_docx_hash(tmp_path: Path) -> None:
    template = tmp_path / "summary-template.docx"
    document = tmp_path / "summary.docx"
    make_ps_summary(template)
    document.write_bytes(template.read_bytes())
    screenshot = tmp_path / "page-1.png"
    write_png(screenshot)
    checklist = tmp_path / "wps-checklist.json"
    wps_checklist(checklist, [screenshot])
    wps_receipt = tmp_path / "wps-receipt.json"
    summary_receipt = tmp_path / "summary-receipt.json"
    copy_receipt = tmp_path / "copy-receipt.json"
    assert run_gate(
        "template-copy", str(template), str(document), "--report", str(copy_receipt)
    ).returncode == 0
    assert run_gate(
        "summary-lint", str(document), "--report", str(summary_receipt)
    ).returncode == 0
    assert run_gate(
        "record-wps-review", str(document), str(checklist), "--report", str(wps_receipt)
    ).returncode == 0

    final_receipt = tmp_path / "final.json"
    brand_receipt = tmp_path / "brand.json"
    brand_receipt.write_text(
        json.dumps(
            {
                "status": "passed",
                "artifacts": [{"format": "docx", "status": "passed", "path": str(document)}],
            }
        ),
        encoding="utf-8",
    )
    passing = run_gate(
        "finalize",
        str(document),
        "--copy-receipt",
        str(copy_receipt),
        "--summary-receipt",
        str(summary_receipt),
        "--wps-receipt",
        str(wps_receipt),
        "--brand-receipt",
        str(brand_receipt),
        "--report",
        str(final_receipt),
    )
    assert passing.returncode == 0, passing.stdout + passing.stderr

    changed_document = Document(document)
    changed_document.add_paragraph("后续保存变更")
    changed_document.save(document)
    failing = run_gate(
        "finalize",
        str(document),
        "--copy-receipt",
        str(copy_receipt),
        "--summary-receipt",
        str(summary_receipt),
        "--wps-receipt",
        str(wps_receipt),
        "--brand-receipt",
        str(brand_receipt),
        "--report",
        str(final_receipt),
    )
    assert failing.returncode != 0
    assert "当前终稿不一致" in failing.stdout


def test_wps_review_rejects_non_wps_engine_and_incomplete_pages(tmp_path: Path) -> None:
    document = tmp_path / "summary.docx"
    make_ps_summary(document)
    screenshot = tmp_path / "page-1.png"
    write_png(screenshot)
    checklist = tmp_path / "bad-checklist.json"
    wps_checklist(checklist, [screenshot])
    data = json.loads(checklist.read_text(encoding="utf-8"))
    data["engine"] = "LibreOffice"
    data["page_count"] = 2
    checklist.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    process = run_gate(
        "record-wps-review",
        str(document),
        str(checklist),
        "--report",
        str(tmp_path / "bad-wps.json"),
    )
    assert process.returncode != 0
    assert '必须精确为\\"WPS Office\\"' in process.stdout
    assert "逐页清单不连续或不完整" in process.stdout
