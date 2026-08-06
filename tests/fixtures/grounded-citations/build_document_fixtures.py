#!/usr/bin/env python3
"""Build real DOCX fixtures from the shared grounded rendering contract."""

from __future__ import annotations

import importlib.util
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "real"
ENGINE_PATH = ROOT / "skills" / "evidence-ledger" / "scripts" / "grounded_evidence.py"

SPEC = importlib.util.spec_from_file_location("fixture_grounded_evidence", ENGINE_PATH)
ENGINE = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(ENGINE)


def set_font(run, *, name: str = "Arial Unicode MS", size: float = 11, bold: bool = False, color: str = "000000") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document, *, standard_override: bool = False) -> None:
    section = doc.sections[0]
    if standard_override:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    doc.core_properties.comments = (
        "Preset: standard_business_brief. Named override: Arial Unicode MS for cross-renderer CJK coverage. "
        + ("Named override: A4 page for Chinese standard fixture." if standard_override else "No page-geometry override.")
    )
    doc.core_properties.author = "Jiaotang QA"
    doc.core_properties.last_modified_by = "Jiaotang QA"
    doc.core_properties.created = datetime(2026, 8, 5, 0, 0, tzinfo=timezone.utc)
    doc.core_properties.modified = datetime(2026, 8, 5, 0, 0, tzinfo=timezone.utc)
    header = section.header.paragraphs[0]
    header.text = "Grounded Citations | Real round-trip fixture"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_font(run, size=8.5, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.text = "Automated acceptance fixture only"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_font(run, size=8.5, color="6B7280")


def title(doc: Document, text: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    set_font(paragraph.add_run(text), size=24, bold=True, color="0B2545")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    set_font(paragraph.add_run(subtitle), size=10.5, color="6B7280")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    set_font(paragraph.add_run(text), size=11)


def build_report() -> Path:
    payload = json.loads((ROOT / "skills" / "evidence-ledger" / "examples" / "normal-grounded-report.json").read_text(encoding="utf-8"))
    payload["sources"][0].update({"title": "Current application notice", "publisher": "Example Authority"})
    payload["sources"][1]["file_name"] = "enterprise-product-sales-workpaper.xlsx"
    payload["records"][0].update({"subject": "Example program", "claim": "The notice requires a market statement", "evidence_excerpt": "Applicants shall submit a market statement for the main product."})
    payload["records"][1].update({"subject": "Example enterprise", "claim": "A product-level sales workpaper is available", "evidence_excerpt": "The workpaper aggregates annual sales by main product."})
    payload["records"][2].update({"subject": "Material readiness", "claim": "Drafting can begin", "limits": "This does not verify market share."})
    payload["document"]["blocks"][0].update({"heading": "Conclusion", "text": "The available materials support drafting the main-product market statement, while market share still requires separate verification."})
    bundle = ENGINE.render_profile_bundle(payload, profile="analysis-report", artifact="docx")
    doc = Document()
    configure(doc)
    title(doc, "Grounded Source Analysis Report", "Shared citation numbering across Word and PDF")
    for block in bundle["blocks"]:
        if block.get("heading"):
            doc.add_heading(block["heading"], level=1)
        markers = "".join(f"[{number}]" for number in block["source_numbers"])
        add_body(doc, block["text"] + markers)
    doc.add_heading("Data Sources", level=1)
    for entry in bundle["source_entries"]:
        add_body(doc, entry)
    path = OUTPUT / "grounded-analysis-report.docx"
    doc.save(path)
    return path


def build_standard() -> tuple[Path, Path]:
    payload = json.loads((HERE / "standard-ledger.json").read_text(encoding="utf-8"))
    payload["sources"][0].update({"title": "Drafting rules for standardization documents", "publisher": "Example Standards Authority"})
    payload["sources"][1]["file_name"] = "product-test-data.xlsx"
    payload["records"][0].update({"subject": "Normative references", "claim": "The standard uses GB/T 1.1-2020", "evidence_excerpt": "The structure and drafting rules follow the applicable drafting standard."})
    payload["records"][1].update({"subject": "Durability requirement", "claim": "The sample remains functional after 1,000 cycles", "evidence_excerpt": "The sample remained functional after 1,000 cycles."})
    payload["document"]["blocks"][0].update({"heading": "Normative references", "text": "GB/T 1.1-2020 Directives for standardization - Part 1: Rules for the structure and drafting of standardizing documents."})
    payload["document"]["blocks"][1].update({"heading": "Technical requirements", "text": "The sample shall remain functional after 1,000 cycles."})
    bundle = ENGINE.render_profile_bundle(payload, profile="standard-native", artifact="docx")

    doc = Document()
    configure(doc, standard_override=True)
    title(doc, "Q/JT 001-2026", "Grounded citation validation standard")
    doc.add_heading("1 Scope", level=1)
    add_body(doc, "This document specifies validation requirements for separating evidence records from the standard body.")
    for block in bundle["blocks"]:
        number = "2" if block["heading"] == "Normative references" else "3"
        doc.add_heading(f"{number} {block['heading']}", level=1)
        add_body(doc, block["text"])
    doc.add_heading("4 Test method", level=1)
    add_body(doc, "Run the sample for the specified number of cycles and record its functional state.")
    standard_path = OUTPUT / "Q-JT-001-2026-grounded-standard.docx"
    doc.save(standard_path)

    memo = Document()
    configure(memo, standard_override=True)
    title(memo, "Standard Source Explanation", "Companion to Q/JT 001-2026; not part of the standard body")
    memo.add_heading("Claim-to-source mapping", level=1)
    for block in bundle["blocks"]:
        markers = "".join(f"[{number}]" for number in block["source_numbers"])
        add_body(memo, f"{block['heading']}：{block['text']} {markers}")
    memo.add_heading("Source register", level=1)
    for entry in bundle["source_entries"]:
        add_body(memo, entry)
    memo_path = OUTPUT / "Q-JT-001-2026-grounded-standard-source-explanation.docx"
    memo.save(memo_path)
    return standard_path, memo_path


def verify(paths: list[Path]) -> None:
    report_text = "\n".join(p.text for p in Document(paths[0]).paragraphs)
    standard_text = "\n".join(p.text for p in Document(paths[1]).paragraphs)
    memo_text = "\n".join(p.text for p in Document(paths[2]).paragraphs)
    assert report_text.rfind("Data Sources") > report_text.find("Conclusion")
    assert "https://example.gov.cn/policy/current" in report_text
    assert "client-dossier" not in report_text
    assert "Normative references" in standard_text
    assert "数据来源" not in standard_text
    assert "Data Sources" not in standard_text
    assert "[1]" not in standard_text
    assert "Standard Source Explanation" in memo_text
    assert "product-test-data.xlsx" in memo_text
    assert "knowledge/internal" not in memo_text


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    report = build_report()
    standard, memo = build_standard()
    verify([report, standard, memo])
    print(json.dumps({"status": "pass", "files": [str(report), str(standard), str(memo)]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
