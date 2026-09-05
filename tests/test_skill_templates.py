from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "validate_skill_templates.py"
SPEC = importlib.util.spec_from_file_location("validate_skill_templates", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(MODULE)


def test_current_suite_templates_complete_initial_baseline() -> None:
    result = MODULE.validate_templates(
        ROOT,
        expected_office_count=29,
        expected_source_count=4,
    )

    assert result["status"] == "pass"
    assert result["office_template_count"] == 29
    assert result["source_template_count"] == 4
    assert len(result["templates"]) == 33


def test_discovery_rejects_legacy_template_even_when_other_template_is_valid(
    tmp_path: Path,
) -> None:
    assets = tmp_path / "skills" / "demo" / "assets"
    assets.mkdir(parents=True)
    Document().save(assets / "formal.docx")
    (assets / "legacy.doc").write_bytes(b"legacy")

    with pytest.raises(MODULE.TemplateValidationError, match="旧式"):
        MODULE.discover_templates(tmp_path)


def test_validation_rejects_corrupt_ooxml_template(tmp_path: Path) -> None:
    assets = tmp_path / "skills" / "demo" / "assets"
    assets.mkdir(parents=True)
    (assets / "broken.docx").write_bytes(b"not-a-docx")

    with pytest.raises(Exception):
        MODULE.validate_templates(tmp_path, expected_office_count=1)


def test_hightech_workbook_prints_each_table_without_empty_horizontal_pages() -> None:
    workbook = load_workbook(ROOT / 'skills/high-tech-enterprise-preassessment/assets/高企预评估双年度采集表.xlsx')
    for sheet in workbook:
        # 19 列台账曾被默认打印配置横向拆成空白页；宽度一页、高度自动，不能把百行表压成一页。
        assert sheet.print_area, sheet.title
        assert sheet.page_setup.orientation == 'landscape', sheet.title
        assert sheet.page_setup.fitToWidth == 1, sheet.title
        assert sheet.page_setup.fitToHeight == 0, sheet.title
        assert sheet.sheet_properties.pageSetUpPr.fitToPage is True, sheet.title
        assert sheet.print_title_rows == '$4:$4', sheet.title


def test_hightech_blank_ip_rows_do_not_duplicate_empty_paragraphs_and_can_expand() -> None:
    document = Document(ROOT / 'skills/high-tech-enterprise-application-drafting/assets/高新技术企业认定申请书空白模板.docx')
    tables = [table for table in document.tables if [cell.text.replace('\n', '') for cell in table.rows[0].cells] ==
              ['知识产权编号', '知识产权名称', '类别', '授权日期', '授权号', '获得方式']]
    assert len(tables) == 1
    assert len(tables[0].rows) == 9
    for row in tables[0].rows[1:]:
        height = row._tr.find('./' + qn('w:trPr') + '/' + qn('w:trHeight'))
        assert height is not None and height.get(qn('w:hRule')) == 'atLeast'
        for cell in row.cells:
            # 重复空段落叠加文档网格曾挤出空页；保留填写行、原行距和可扩展行高。
            if not cell.text.strip():
                assert len(cell.paragraphs) == 1
                assert cell.paragraphs[0].paragraph_format.line_spacing == 1.25
