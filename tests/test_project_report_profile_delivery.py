import hashlib
import importlib.util
import json
import os
from pathlib import Path

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = (
    ROOT
    / "skills/project-feasibility/scripts/validate_report_profile_delivery.py"
)
SPEC = importlib.util.spec_from_file_location("report_profile_validator", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(MODULE)


def test_report_profile_requires_dual_format_sections_tables_and_branding(tmp_path):
    document = Document()
    document.add_heading("项目版本与窗口", level=1)
    document.add_heading("总体结论", level=1)
    target = tmp_path / "incomplete.docx"
    document.save(target)
    result = MODULE.validate_profile(
        plugin_root=ROOT,
        profile_id="project-feasibility-analysis-report",
        artifacts=[target],
    )
    assert result["status"] == "fail"
    assert "缺少要求格式:pdf" in result["errors"]
    assert any("缺少必备章节" in item for item in result["errors"])
    assert any("缺少必备表格" in item for item in result["errors"])
    assert any("品牌校验失败" in item for item in result["errors"])


def test_stop_hook_rejects_tampered_profile_receipt(tmp_path):
    manager_scripts = Path(
        os.environ.get(
            "JIAOTANG_RELEASE_MANAGER_SCRIPTS",
            Path.home() / ".codex/skills/skill-release-manager/scripts",
        )
    )
    hook_path = manager_scripts / "workbuddy_behavior_hook.py"
    if not hook_path.is_file():
        pytest.skip("requires the separately installed skill-release-manager host integration")
    spec = importlib.util.spec_from_file_location("hotfix_behavior", hook_path)
    hook = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(hook)
    if not hasattr(hook, "load_profile_validator_receipts"):
        pytest.skip("installed skill-release-manager predates report-profile receipts")
    artifact = tmp_path / "report.docx"
    artifact.write_bytes(b"initial")
    turn_id = "turn-profile-test"
    receipt_dir = tmp_path / "validator-receipts" / turn_id
    receipt_dir.mkdir(parents=True)
    receipt = {
        "validator_id": MODULE.VALIDATOR_ID,
        "status": "pass",
        "turn_id": turn_id,
        "profile_id": "project-feasibility-analysis-report",
        "artifacts": [
            {
                "path": str(artifact),
                "sha256": hashlib.sha256(artifact.read_bytes()).hexdigest(),
            }
        ],
    }
    (receipt_dir / "profile.json").write_text(
        json.dumps(receipt), encoding="utf-8"
    )
    assert len(hook.load_profile_validator_receipts(tmp_path, turn_id)) == 1
    artifact.write_bytes(b"tampered")
    assert hook.load_profile_validator_receipts(tmp_path, turn_id) == []


def test_pdf_section_matching_tolerates_glyph_fragmentation_but_not_missing_text():
    fragmented = "\u4e8c\u3001\u7533\n\u62a5\n\u6761\u4ef6\n\u5bf9\n\u7167"
    assert MODULE._compact_text("申报条件对照") in MODULE._compact_text(fragmented)
    assert MODULE._compact_text("科技咨询补强") not in MODULE._compact_text(fragmented)


def test_pdf_text_extraction_uses_visual_sort_order(monkeypatch, tmp_path):
    calls = []

    class Page:
        def get_text(self, mode, *, sort=False):
            calls.append((mode, sort))
            return "申报条件对照"

    class Pdf:
        def __iter__(self):
            return iter([Page()])

        def close(self):
            return None

    import fitz

    monkeypatch.setattr(fitz, "open", lambda _: Pdf())
    assert MODULE._pdf_text(tmp_path / "fixture.pdf") == "申报条件对照"
    assert calls == [("text", True)]


def test_pdf_cjk_font_gate_requires_embedded_portable_font(monkeypatch, tmp_path):
    class Page:
        def get_fonts(self, *, full=False):
            assert full
            return [(7, "ttf", "TrueType", "ABCDEF+NotoSansSC-Regular", "F1", "", 0)]

        def get_text(self, mode):
            assert mode == "dict"
            return {"blocks": [{"lines": [{"spans": [{"text": "中文正文", "font": "NotoSansSC-Regular"}]}]}]}

    class Pdf:
        def __iter__(self):
            return iter([Page()])

        def extract_font(self, xref):
            assert xref == 7
            return ("NotoSansSC-Regular", "ttf", "TrueType", b"embedded subset")

        def close(self):
            return None

    import fitz

    monkeypatch.setattr(fitz, "open", lambda _: Pdf())
    assert MODULE._pdf_cjk_font_errors(tmp_path / "fixture.pdf") == []


def test_pdf_cjk_font_gate_accepts_self_contained_type3_glyphs(monkeypatch, tmp_path):
    class Page:
        def get_fonts(self, *, full=False):
            assert full
            return [(7, "n/a", "Type3", "NotoSansSC-Thin", "F1", "", 0)]

        def get_text(self, mode):
            assert mode == "dict"
            return {"blocks": [{"lines": [{"spans": [{"text": "中文正文", "font": "NotoSansSC-Thin"}]}]}]}

    class Pdf:
        def __iter__(self):
            return iter([Page()])

        def extract_font(self, _xref):
            return ("NotoSansSC-Thin", "n/a", "Type3", b"")

        def xref_get_key(self, xref, key):
            assert (xref, key) == (7, "CharProcs")
            return ("dict", "<</gid1 11 0 R/gid2 12 0 R>>")

        def xref_is_stream(self, xref):
            return xref in {11, 12}

        def xref_stream(self, xref):
            return b"glyph" if xref in {11, 12} else b""

        def close(self):
            return None

    import fitz

    monkeypatch.setattr(fitz, "open", lambda _: Pdf())
    assert MODULE._pdf_cjk_font_errors(tmp_path / "fixture.pdf") == []


def test_pdf_cjk_font_gate_rejects_host_fallback(monkeypatch, tmp_path):
    class Page:
        def get_fonts(self, *, full=False):
            return [(7, "ttf", "TrueType", "ABCDEF+YuGothicUI", "F1", "", 0)]

        def get_text(self, mode):
            return {"blocks": [{"lines": [{"spans": [{"text": "中文正文", "font": "YuGothicUI"}]}]}]}

    class Pdf:
        def __iter__(self):
            return iter([Page()])

        def extract_font(self, _xref):
            return ("YuGothicUI", "ttf", "TrueType", b"")

        def close(self):
            return None

    import fitz

    monkeypatch.setattr(fitz, "open", lambda _: Pdf())
    errors = MODULE._pdf_cjk_font_errors(tmp_path / "fixture.pdf")
    assert any("未嵌入" in item for item in errors)
