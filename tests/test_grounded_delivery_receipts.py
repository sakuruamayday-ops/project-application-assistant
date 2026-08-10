import importlib.util
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "skills" / "evidence-ledger" / "scripts" / "grounded_evidence.py"
REAL = ROOT / "tests" / "fixtures" / "grounded-citations" / "real"
SPEC = importlib.util.spec_from_file_location("grounded_delivery", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(MODULE)


def test_dependency_free_xlsx_dump_reads_values_formulas_and_sheet_order():
    result = MODULE.dump_xlsx(REAL / "xlsx" / "grounded-market-share.xlsx")
    assert [sheet["name"] for sheet in result["sheets"]] == ["分析结果", "计算底稿", "数据来源"]
    cells = [cell for sheet in result["sheets"] for row in sheet["rows"] for cell in row]
    assert any(cell["formula"] for cell in cells)
    assert any("市场占有率" in str(cell["value"]) for cell in cells)


def test_delivery_receipt_is_bound_to_current_turn_artifact_and_ledger_hash(tmp_path: Path):
    state_root = tmp_path / "state"
    state_root.mkdir()
    (state_root / "current-turn.json").write_text(
        json.dumps({"turn_id": "turn-grounded-001"}),
        encoding="utf-8",
    )
    ledger = ROOT / "skills" / "evidence-ledger" / "examples" / "normal-grounded-report.json"
    artifact = REAL / "grounded-analysis-report.docx"
    result = MODULE.write_delivery_receipt(
        ledger,
        artifact,
        profile="analysis-report",
        state_root=state_root,
    )
    receipt_path = Path(result["receipt_path"])
    assert result["status"] == "pass"
    assert result["turn_id"] == "turn-grounded-001"
    assert result["artifact"]["sha256"] == MODULE.sha256_file(artifact)
    assert receipt_path.parent == state_root / "validator-receipts" / "turn-grounded-001"
    assert json.loads(receipt_path.read_text(encoding="utf-8"))["validator_id"] == "grounded-delivery/v1"


def test_delivery_receipt_keeps_canonical_copy_and_exports_user_visible_copy(tmp_path: Path):
    state_root = tmp_path / "state"
    state_root.mkdir()
    (state_root / "current-turn.json").write_text(
        json.dumps({"turn_id": "turn-grounded-export"}),
        encoding="utf-8",
    )
    export_root = tmp_path / "artifacts" / "validator-receipts"
    ledger = ROOT / "skills" / "evidence-ledger" / "examples" / "normal-grounded-report.json"
    artifact = REAL / "grounded-analysis-report.docx"

    result = MODULE.write_delivery_receipt(
        ledger,
        artifact,
        profile="analysis-report",
        state_root=state_root,
        receipt_export_dir=export_root,
    )

    canonical = Path(result["receipt_path"])
    exported = Path(result["receipt_export_path"])
    assert canonical.parent == state_root / "validator-receipts" / "turn-grounded-export"
    assert exported.parent == export_root / "turn-grounded-export"
    assert canonical.read_bytes() == exported.read_bytes()
    assert json.loads(exported.read_text(encoding="utf-8"))["turn_id"] == "turn-grounded-export"


def test_macos_marketplace_state_root_matches_hook_data_directory(tmp_path: Path, monkeypatch):
    plugin_root = (
        tmp_path
        / ".workbuddy"
        / "plugins"
        / "marketplaces"
        / "jiaotang"
        / "plugins"
        / "jiaotang-workbuddy-skills"
    )
    (plugin_root / ".codebuddy-plugin").mkdir(parents=True)
    (plugin_root / ".codebuddy-plugin" / "plugin.json").write_text("{}", encoding="utf-8")
    monkeypatch.delenv("JIAOTANG_BEHAVIOR_STATE_ROOT", raising=False)
    monkeypatch.delenv("CODEBUDDY_PLUGIN_DATA", raising=False)
    monkeypatch.setenv("CODEBUDDY_PLUGIN_ROOT", str(plugin_root))

    assert MODULE._default_state_root() == (
        tmp_path
        / ".workbuddy"
        / "plugins"
        / "data"
        / "jiaotang-workbuddy-skills-jiaotang"
        / "behavior-hook"
    )


def test_windows_state_root_ignores_transient_plugin_data(tmp_path: Path, monkeypatch):
    explicit_home = tmp_path / "profile"
    transient_plugin_data = tmp_path / "host-plugin-data"
    monkeypatch.delenv("JIAOTANG_BEHAVIOR_STATE_ROOT", raising=False)
    monkeypatch.setenv("CODEBUDDY_PLUGIN_DATA", str(transient_plugin_data))
    monkeypatch.setattr(MODULE.os, "name", "nt")
    monkeypatch.setattr(MODULE.Path, "home", classmethod(lambda cls: explicit_home))

    assert MODULE._default_state_root() == (
        explicit_home / ".workbuddy" / "state" / "jiaotang-behavior"
    )


def test_docx_validator_rejects_user_observed_provenance_and_structure_defects(tmp_path: Path):
    path = tmp_path / "bad-report.docx"
    document = Document()
    document.core_properties.author = "python-docx"
    document.add_paragraph("市场占有率分析报告 [1]")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "结果"
    document.add_paragraph("数据来源")
    document.add_paragraph("[1] 测试台账.xlsx")
    document.save(path)

    result = MODULE.validate_artifact(path, "analysis-report")
    joined = " | ".join(result["errors"])
    assert result["status"] == "fail"
    assert "真实标题样式" in joined
    assert "重复表头" in joined
    assert "暴露生成库" in joined
    assert "默认时间" in joined
    assert "东亚字体" in joined


def test_docx_validator_rejects_default_modified_time_and_generator_description(tmp_path: Path):
    path = tmp_path / "metadata-leak.docx"
    document = Document()
    document.core_properties.author = "焦糖证据链台账"
    document.core_properties.last_modified_by = "焦糖证据链台账"
    document.core_properties.comments = "generated by python-docx"
    document.add_heading("市场占有率证据分析报告", level=1)
    document.add_paragraph("结论 [1]")
    document.add_heading("数据来源", level=1)
    document.add_paragraph("[1] 测试台账.xlsx")
    document.save(path)

    import zipfile
    from xml.etree import ElementTree as ET

    rewritten = tmp_path / "metadata-leak-rewritten.docx"
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(rewritten, "w") as target:
        core = ET.fromstring(source.read("docProps/core.xml"))
        modified = core.find("{http://purl.org/dc/terms/}modified")
        assert modified is not None
        modified.text = "2013-12-23T23:15:00Z"
        for member in source.infolist():
            content = source.read(member.filename)
            if member.filename == "docProps/core.xml":
                content = ET.tostring(core, encoding="utf-8", xml_declaration=True)
            target.writestr(member, content)
    path.write_bytes(rewritten.read_bytes())

    result = MODULE.validate_artifact(path, "analysis-report")
    joined = " | ".join(result["errors"])
    assert result["status"] == "fail"
    assert "修改时间缺失或仍为模板默认时间" in joined
    assert "核心元数据暴露文档生成器" in joined


def test_delivery_rejects_reference_only_sources_rendered_as_retrieved(tmp_path: Path):
    state_root = tmp_path / "state"
    state_root.mkdir()
    (state_root / "current-turn.json").write_text(
        json.dumps({"turn_id": "turn-reference-only-001"}),
        encoding="utf-8",
    )
    ledger = ROOT / "skills" / "evidence-ledger" / "examples" / "reference-only-market-share.json"
    path = tmp_path / "bad-source-disclosure.docx"
    document = Document()
    current = datetime.now(timezone.utc)
    document.core_properties.author = "焦糖证据链台账"
    document.core_properties.last_modified_by = "焦糖证据链台账"
    document.core_properties.created = current
    document.core_properties.modified = current
    document.core_properties.comments = ""
    for style_name in ("Normal", "Title", "Heading 1"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")
    document.add_heading("市场占有率证据分析报告", level=1)
    document.add_paragraph("受限复现10%，不得对外使用。[1][2][3]")
    document.add_heading("数据来源", level=1)
    document.add_paragraph("[1] 用户文件：分产品销售台账.xlsx（原件未提供）")
    document.add_paragraph(
        "[2] 示例行业研究机构；《全国上位市场规模报告》；"
        "https://example.org/report/market-2025；检索日期 2026-08-06（原文未提供）"
    )
    document.add_paragraph("[3] 企业陈述：主导产品应用场景拆分说明（原件未提供）")
    document.save(path)

    result = MODULE.write_delivery_receipt(
        ledger,
        path,
        profile="analysis-report",
        state_root=state_root,
    )
    joined = " | ".join(result["errors"])
    assert result["status"] == "fail"
    assert "未标明工作簿登记属性" in joined
    assert "不得声称检索日期" in joined
    assert "未披露登记载体" in joined


def test_delivery_accepts_reference_only_sources_with_explicit_access_status(tmp_path: Path):
    state_root = tmp_path / "state"
    state_root.mkdir()
    (state_root / "current-turn.json").write_text(
        json.dumps({"turn_id": "turn-reference-only-002"}),
        encoding="utf-8",
    )
    ledger = ROOT / "skills" / "evidence-ledger" / "examples" / "reference-only-market-share.json"
    payload = json.loads(ledger.read_text(encoding="utf-8"))
    bundle = MODULE.render_profile_bundle(payload, profile="analysis-report", artifact="docx")
    path = tmp_path / "good-source-disclosure.docx"
    document = Document()
    current = datetime.now(timezone.utc)
    document.core_properties.author = "焦糖证据链台账"
    document.core_properties.last_modified_by = "焦糖证据链台账"
    document.core_properties.created = current
    document.core_properties.modified = current
    document.core_properties.comments = ""
    for style_name in ("Normal", "Title", "Heading 1"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")
    document.add_heading("市场占有率证据分析报告", level=1)
    document.add_paragraph("受限复现10%，不得对外使用。[1][2][3][4]")
    document.add_paragraph("证据等级为D级，排名未核验且不作排名结论。[1][2][3][4]")
    document.add_paragraph("底层原件未取得，补齐后应重新校验与分级。[1][2][3][4]")
    document.add_heading("数据来源", level=1)
    for entry in bundle["source_entries"]:
        document.add_paragraph(entry)
    document.save(path)

    result = MODULE.write_delivery_receipt(
        ledger,
        path,
        profile="analysis-report",
        state_root=state_root,
    )
    assert result["status"] == "pass", result["errors"]
    assert "source-disclosure-access-status" in result["checks"]


def test_standard_delivery_requires_and_hash_binds_separate_source_memo(tmp_path: Path):
    state_root = tmp_path / "state"
    state_root.mkdir()
    (state_root / "current-turn.json").write_text(
        json.dumps({"turn_id": "turn-standard-001"}),
        encoding="utf-8",
    )
    ledger = ROOT / "tests" / "fixtures" / "grounded-citations" / "standard-ledger.json"
    standard = REAL / "Q-JT-001-2026-grounded-standard.docx"
    memo = REAL / "Q-JT-001-2026-grounded-standard-source-explanation.docx"
    missing = MODULE.write_delivery_receipt(
        ledger,
        standard,
        profile="standard-native",
        state_root=state_root,
    )
    assert missing["status"] == "fail"
    assert any("标准数据来源说明" in item for item in missing["errors"])

    passed = MODULE.write_delivery_receipt(
        ledger,
        standard,
        profile="standard-native",
        source_memo_path=memo,
        state_root=state_root,
    )
    assert passed["status"] == "pass"
    assert passed["sidecars"] == [
        {
            "name": memo.name,
            "type": "docx",
            "path": str(memo.resolve()),
            "sha256": MODULE.sha256_file(memo),
        }
    ]
