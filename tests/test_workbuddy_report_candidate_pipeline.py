from __future__ import annotations

import importlib.util
import json
import re
import zipfile
from pathlib import Path
from uuid import UUID
from xml.etree import ElementTree as ET

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SKILLS = ROOT / "skills"
REGISTRY = SKILLS / "project-feasibility/references/report-template-registry.json"
SELECTOR_PATH = SKILLS / "project-feasibility/scripts/select_report_template.py"
FILLER_PATH = SKILLS / "project-feasibility/scripts/fill_report_template.py"
PROFILE_VALIDATOR_PATH = SKILLS / "project-feasibility/scripts/validate_report_profile_delivery.py"
PIPELINE_PATH = ROOT / "scripts/run_workbuddy_report_candidate_pipeline.py"
VISUAL_FINALIZER_PATH = ROOT / "scripts/record_workbuddy_report_visual_review.py"


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


SELECTOR = load_module("candidate_test_selector", SELECTOR_PATH)
FILLER = load_module("candidate_test_filler", FILLER_PATH)
PROFILE_VALIDATOR = load_module("candidate_test_profile_validator", PROFILE_VALIDATOR_PATH)
PIPELINE = load_module("candidate_test_pipeline", PIPELINE_PATH)
VISUAL_FINALIZER = load_module("candidate_test_visual_finalizer", VISUAL_FINALIZER_PATH)


def test_report_font_is_embedded_and_document_remains_editable(tmp_path: Path):
    document = Document()
    document.add_paragraph("GC-QA 字体兼容 17 + 31 = 48")
    document.sections[0].first_page_header.paragraphs[0].text = "首页页眉"
    FILLER._apply_portable_cjk_font(document)
    # 字体处理不能换掉 Document 持有的 XML 根，否则后续修改会悄悄不保存。
    document.add_paragraph("字体处理后新增段落")
    output = tmp_path / "embedded.docx"
    document.save(output)
    reopened = Document(output)
    assert "字体处理后新增段落" in [paragraph.text for paragraph in reopened.paragraphs]
    with zipfile.ZipFile(output) as archive:
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        table = ET.fromstring(archive.read("word/fontTable.xml"))
        fonts = table.findall("w:font", namespace)
        assert len(fonts) == 1
        assert fonts[0].get(f"{{{namespace['w']}}}name") == "Noto Sans SC"
        regular = fonts[0].find("w:embedRegular", namespace)
        assert regular is not None
        key = UUID(regular.get(f"{{{namespace['w']}}}fontKey")).bytes[::-1]
        data = bytearray(archive.read("word/fonts/gongchuang-noto-sans-sc.odttf"))
        for index in range(32):
            data[index] ^= key[index % 16]
        assert bytes(data) == FILLER.PORTABLE_CJK_FONT_PATH.read_bytes()
        assert b"Microsoft YaHei" not in archive.read("word/styles.xml")
        assert "首页页眉" in archive.read("word/header1.xml").decode("utf-8")


def test_table_spacers_do_not_consume_body_lines_or_remove_chapter_breaks():
    document = Document()
    document.add_table(rows=1, cols=1)
    spacer = document.add_paragraph()
    document.add_table(rows=1, cols=1)
    chapter_break = document.add_paragraph()
    chapter_break.paragraph_format.page_break_before = True
    document.add_table(rows=1, cols=1)
    content = document.add_paragraph("仍需保留的说明")
    heading = document.add_heading("6.2 任务排序", level=2)
    first = document.add_paragraph("一级：完成关键证据。", style="List Bullet")
    last = document.add_paragraph("二级：补齐一般材料。", style="List Bullet")
    FILLER._preserve_report_pagination(document)
    assert spacer.paragraph_format.line_spacing.pt == 1
    assert spacer.paragraph_format.keep_with_next is True
    assert chapter_break.paragraph_format.page_break_before is True
    assert chapter_break.paragraph_format.line_spacing is None
    assert content.text == "仍需保留的说明"
    assert content.paragraph_format.line_spacing is None
    assert heading.text == "6.2 任务排序"
    assert first.paragraph_format.keep_with_next is True
    assert last.paragraph_format.keep_with_next is None


def client_source(path: Path, *, name: str = "某高端装备有限公司") -> Path:
    document = Document()
    document.add_heading("企业项目材料", level=1)
    document.add_paragraph(f"企业名称：{name}")
    document.add_paragraph("申报对象：工程机械智能控制系统")
    document.add_paragraph("已形成自主研发、版本测试和客户应用资料。")
    document.save(path)
    return path


def case_fixture(source: Path, project_id: str = "first-equipment") -> dict:
    return {
        "project_id": project_id,
        "enterprise": "某高端装备有限公司",
        "project_object": "工程机械智能控制系统",
        "suggested_year": "2027年",
        "deadline": "待当年通知",
        "conclusion": "培育后申报",
        "conclusion_basis": "已有产品和研发材料，检测、查新和首次应用证据待补齐。",
        "primary_gap": "性能验证与首次应用证据尚未闭合",
        "next_action": "锁定产品版本，完成第三方测试、查新和客户验收台账。",
        "materials": [
            {
                "path": str(source),
                "role": "企业技术材料",
                "anchors": ["某高端装备有限公司", "工程机械智能控制系统"],
                "source_type": "knowledge-base",
            }
        ],
        "policies": [
            {
                "title": "当期项目申报通知待发布",
                "locator": "规划参考，待当年官方通知",
                "status": "规划基线",
            }
        ],
        "conditions": [
            {
                "match": "核心技术拥有自主知识产权",
                "value": "当前材料已列明自主研发，权利清单待核",
                "state": "企业提供待核",
                "gap": "权利状态与产品映射未闭合",
                "action": "建立知识产权与产品技术对应表",
            }
        ],
    }


@pytest.mark.parametrize("report_type", ["preassessment", "feasibility"])
def test_real_source_anchor_fill_produces_complete_editable_report(tmp_path: Path, report_type: str):
    source = client_source(tmp_path / "client.docx")
    selection = SELECTOR.resolve_template("first-equipment", report_type, registry_path=REGISTRY)
    copied = SELECTOR.materialize(selection, tmp_path / "masters", enterprise="某高端装备有限公司")
    output = tmp_path / f"completed-{report_type}.docx"
    result = FILLER.complete_report(
        template_path=Path(copied["output_path"]),
        output_path=output,
        fixture=case_fixture(source),
        report_type=report_type,
        release_tag="V1.6.5.2",
        public_root=ROOT,
    )
    assert result["status"] == "pass"
    assert result["source_count"] == 1
    document = Document(output)
    expected_label = "项目前期评估报告" if report_type == "preassessment" else "项目申报可行性分析报告"
    assert document.core_properties.title.endswith(expected_label)
    assert document.core_properties.subject == expected_label
    text = FILLER.document_text(document)
    assert "［填写" not in text
    assert "培训模板" not in text
    assert "某高端装备有限公司" in text
    assert "数据来源" in text
    assert "序号" in text and "文件名称" in text and "链接" in text
    assert "来源共创知识库" in text
    assert "SHA-256" not in text
    assert "真实客户资料候选验收稿" not in text
    assert "V1.6.5.1" not in text
    assert re.search(r"(?<![0-9a-f])[0-9a-f]{64}(?![0-9a-f])", text, re.IGNORECASE) is None
    visible_runs = [
        run
        for paragraph in list(document.paragraphs)
        + [p for table in document.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]
        for run in paragraph.runs
        if run.text.strip()
    ]
    assert visible_runs
    assert all(run.font.name == FILLER.PORTABLE_CJK_FONT for run in visible_runs)
    profile_id = (
        "project-presale-assessment-report"
        if report_type == "preassessment"
        else "project-feasibility-analysis-report"
    )
    validation = PROFILE_VALIDATOR.validate_profile(
        plugin_root=ROOT, profile_id=profile_id, artifacts=[output]
    )
    assert validation["status"] == "pass", validation["errors"]


def test_report_maps_policy_roles_and_strengthening_tasks_without_repeated_headline(tmp_path: Path):
    source = client_source(tmp_path / "client.docx")
    fixture = case_fixture(source, project_id="science-plan")
    fixture["conclusion"] = "待资料"
    fixture["headline_fact"] = "不得重复进补强任务表的工商与专利概况"
    fixture["policies"] = [
        {
            "title": "关于组织申报2027年度某项目的通知",
            "locator": "https://example.test/notice",
            "status": "已核验基线",
        },
        {
            "title": "关于推进项目实施的若干措施",
            "locator": "https://example.test/measure",
            "status": "已核验基线",
        },
    ]
    selection = SELECTOR.resolve_template("science-plan", "preassessment", registry_path=REGISTRY)
    copied = SELECTOR.materialize(selection, tmp_path / "masters", enterprise="某高端装备有限公司")
    output = tmp_path / "mapped.docx"
    result = FILLER.complete_report(
        template_path=Path(copied["output_path"]),
        output_path=output,
        fixture=fixture,
        report_type="preassessment",
        release_tag="V1.6.6",
        public_root=ROOT,
    )
    assert result["status"] == "pass"
    document = Document(output)
    policy_table = next(
        table for table in document.tables
        if "政策层级" in " | ".join(cell.text for cell in table.rows[0].cells)
    )
    policies = {row.cells[0].text.strip(): row for row in policy_table.rows[1:]}
    assert policies["管理办法或评价办法"].cells[1].text == "关于推进项目实施的若干措施"
    assert policies["申报年度通知"].cells[1].text == "关于组织申报2027年度某项目的通知"
    assert policies["申报指南与评分附件"].cells[1].text == "当前未取得独立文件"
    assert policies["截止日期"].cells[1].text == "关于组织申报2027年度某项目的通知"

    strengthening = next(
        table for table in document.tables
        if "补强任务" in " | ".join(cell.text for cell in table.rows[0].cells)
    )
    visible = "\n".join(cell.text for row in strengthening.rows for cell in row.cells)
    assert fixture["headline_fact"] not in visible
    assert "尚未形成国内技术水平评价报告" in visible
    assert "任务指标对照表、检测方案和项目预算底稿" in visible
    assert "暂无法判断" in FILLER.document_text(document)


def test_report_accepts_ordered_custom_strengthening_tasks_and_preserves_page_fields(tmp_path: Path):
    source = client_source(tmp_path / "client.docx")
    fixture = case_fixture(source, project_id="science-plan")
    fixture["strengthening_tasks"] = [
        {
            "task": "核定关键技术指标",
            "status": "指标口径尚待企业确认",
            "target_period": "2026年9月",
            "deliverable": "关键技术指标确认表",
        },
        {
            "task": "完成项目预算测算",
            "status": "设备与研发投入尚未归集",
            "target_period": "2026年10月",
            "deliverable": "项目预算及资金来源表",
        },
        {
            "task": "形成申报时间表",
            "status": "等待当期通知确认截止时间",
            "target_period": "通知发布后两日内",
            "deliverable": "申报倒排计划",
        },
    ]
    selection = SELECTOR.resolve_template("science-plan", "preassessment", registry_path=REGISTRY)
    copied = SELECTOR.materialize(selection, tmp_path / "masters", enterprise="某高端装备有限公司")
    output = tmp_path / "custom-tasks.docx"
    result = FILLER.complete_report(
        template_path=Path(copied["output_path"]),
        output_path=output,
        fixture=fixture,
        report_type="preassessment",
        release_tag="V1.6.6",
        public_root=ROOT,
    )
    assert result["status"] == "pass"
    document = Document(output)
    strengthening = next(
        table for table in document.tables
        if "补强任务" in " | ".join(cell.text for cell in table.rows[0].cells)
    )
    assert len(strengthening.rows) == 4
    assert [row.cells[1].text for row in strengthening.rows[1:]] == [
        "核定关键技术指标",
        "完成项目预算测算",
        "形成申报时间表",
    ]
    assert strengthening.rows[2].cells[3].text == "2026年10月"
    assert strengthening.rows[3].cells[4].text == "申报倒排计划"
    assert document.core_properties.title == "某高端装备有限公司_工程机械智能控制系统_项目前期评估报告"
    assert document.core_properties.author == "共创研究院"
    with zipfile.ZipFile(output) as archive:
        footers = "\n".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )
        metadata = archive.read("docProps/core.xml").decode("utf-8")
    assert "PAGE" in footers
    assert "培训模板" not in footers
    assert "V1.6.5.1" not in footers
    assert "培训模板" not in metadata
    assert "候选验收稿" not in metadata


def test_source_anchor_must_match_real_material(tmp_path: Path):
    source = client_source(tmp_path / "client.docx")
    fixture = case_fixture(source)
    fixture["materials"][0]["anchors"] = ["原文完全不包含的假断言"]
    with pytest.raises(ValueError, match="原文锚点未命中"):
        FILLER.validate_case_fixture(fixture, public_root=ROOT)


def test_private_customer_material_is_rejected_inside_public_skill_tree(tmp_path: Path):
    source = client_source(tmp_path / "_private-client-fixture.docx")
    with pytest.raises(ValueError, match="不得放入公共技能树"):
        FILLER.validate_case_fixture(case_fixture(source), public_root=tmp_path)


def test_fixture_manifest_requires_exact_twelve_project_order(tmp_path: Path):
    cases = [{"project_id": project_id} for project_id in PIPELINE.PROJECT_IDS]
    manifest = tmp_path / "fixtures.json"
    manifest.write_text(
        json.dumps(
            {
                "schema": "gongchuang-private-real-client-report-fixtures/v1",
                "release_tag": "V1.6.5.2",
                "cases": cases,
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    assert [item["project_id"] for item in PIPELINE.load_fixture_manifest(manifest, release_tag="V1.6.5.2")] == list(
        PIPELINE.PROJECT_IDS
    )
    payload = json.loads(manifest.read_text(encoding="utf-8"))
    payload["cases"] = list(reversed(payload["cases"]))
    manifest.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(ValueError, match="12类受控项目顺序"):
        PIPELINE.load_fixture_manifest(manifest, release_tag="V1.6.5.2")


def test_zip_privacy_scan_blocks_customer_name_and_path(tmp_path: Path):
    source = client_source(tmp_path / "client.docx")
    fixture = case_fixture(source)
    clean = tmp_path / "clean.zip"
    with zipfile.ZipFile(clean, "w") as archive:
        archive.writestr("skills/example/SKILL.md", "---\nname: example\n---\n")
    assert PIPELINE.scan_archive_privacy(clean, [fixture])["status"] == "pass"
    dirty = tmp_path / "dirty.zip"
    with zipfile.ZipFile(dirty, "w") as archive:
        archive.writestr("leak.txt", fixture["enterprise"])
    with pytest.raises(RuntimeError, match="混入真实客户信息"):
        PIPELINE.scan_archive_privacy(dirty, [fixture])
    dirty_name = tmp_path / "dirty-name.zip"
    with zipfile.ZipFile(dirty_name, "w") as archive:
        archive.writestr(f"skills/{fixture['enterprise']}/SKILL.md", "clean body")
    with pytest.raises(RuntimeError, match="混入真实客户信息"):
        PIPELINE.scan_archive_privacy(dirty_name, [fixture])


def test_candidate_pipeline_source_excludes_generic_and_zcode_packaging():
    text = PIPELINE_PATH.read_text(encoding="utf-8").casefold()
    assert "package_workbuddy_suite.py" in text
    assert "package_skill_suite.py" not in text
    assert "run_release_gates" in text
    assert "run_post_package_gates" in text
    assert "installed_filler" in text
    assert "zcode" in text
    assert "not-built-not-tested" in text
    assert PIPELINE.PLATFORMS == ("macos", "windows")


def test_safe_extract_rejects_path_traversal(tmp_path: Path):
    archive = tmp_path / "unsafe.zip"
    with zipfile.ZipFile(archive, "w") as bundle:
        bundle.writestr("../escape.txt", "bad")
    with pytest.raises(RuntimeError, match="不安全条目"):
        PIPELINE.safe_extract_zip(archive, tmp_path / "out")


def visual_fixture(tmp_path: Path) -> tuple[Path, Path, Path]:
    sheet = tmp_path / "contact.png"
    sheet.write_bytes(b"contact-sheet")
    digest = PIPELINE.sha256_file(sheet)
    receipt = tmp_path / "pipeline.json"
    receipt.write_text(
        json.dumps(
            {
                "schema": "gongchuang-workbuddy-report-candidate-pipeline/v1",
                "status": "pending",
                "automated_gate_status": "pass",
                "candidate_state": "pending-visual-review",
                "release_tag": "V1.6.5.2",
                "source_commit": "abc123",
                "contact_sheet": {"path": str(sheet), "sha256": digest, "sample_count": 24},
                "real_host_acceptance": {"macos": "pending", "windows": "pending"},
                "zcode": {"status": "not-built-not-tested"},
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    checklist = tmp_path / "checklist.json"
    checklist.write_text(
        json.dumps(
            {
                "schema": "gongchuang-report-visual-review/v1",
                "release_tag": "V1.6.5.2",
                "contact_sheet_sha256": digest,
                "reviewer": "test-reviewer",
                "review_method": "test-visual-review",
                "reviewed_at": "2026-08-13T21:00:00+08:00",
                "items": [
                    {
                        "project_id": project_id,
                        "report_type": report_type,
                        "status": "pass",
                        "checks": {name: True for name in VISUAL_FINALIZER.REQUIRED_CHECKS},
                    }
                    for project_id in PIPELINE.PROJECT_IDS
                    for report_type in PIPELINE.REPORT_TYPES
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return receipt, checklist, tmp_path / "visual-receipt.json"


def test_visual_finalizer_requires_exact_24_passed_samples(tmp_path: Path):
    receipt, checklist, output = visual_fixture(tmp_path)
    result = VISUAL_FINALIZER.finalize_visual_review(
        repo_root=ROOT,
        pipeline_receipt=receipt,
        checklist_path=checklist,
        output_path=output,
    )
    assert result["status"] == "pass"
    assert result["candidate_state"] == "ready-for-real-host-testing"
    assert result["formal_release_eligible"] is False
    payload = json.loads(checklist.read_text(encoding="utf-8"))
    payload["items"][0]["checks"]["no_missing_glyphs"] = False
    bad = tmp_path / "bad.json"
    bad.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(ValueError, match="视觉抽检未通过"):
        VISUAL_FINALIZER.finalize_visual_review(
            repo_root=ROOT,
            pipeline_receipt=receipt,
            checklist_path=bad,
            output_path=tmp_path / "bad-receipt.json",
        )
