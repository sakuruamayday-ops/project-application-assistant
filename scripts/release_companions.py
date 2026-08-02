#!/usr/bin/env python3
"""Generate and verify release companions from the suite manifest."""

from __future__ import annotations

import argparse
from contextlib import contextmanager
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
MANIFEST_PATH = ROOT / "skills" / "suite-manifest.json"
RELEASE_PREFIXES = (
    "本版本更新：",
    "能力变化：",
    "通用安装：",
    "WorkBuddy 安装：",
    "手工接入：",
    "安装验收：",
    "兼容性：",
    "回滚：",
)


def _recovery_root() -> Path:
    configured = os.environ.get("JIAOTANG_RELEASE_WORK_ROOT", "").strip()
    if configured:
        return Path(configured).expanduser()
    if sys.platform == "darwin":
        return Path.home() / ".Trash" / "jiaotang-release-workspaces"
    if os.name == "nt":
        return Path(tempfile.gettempdir()) / "jiaotang-release-workspaces"
    return (
        Path(os.environ.get("XDG_DATA_HOME", Path.home() / ".local/share"))
        / "Trash"
        / "files"
        / "jiaotang-release-workspaces"
    )


@contextmanager
def recoverable_workspace(prefix: str) -> Iterable[Path]:
    """Create scratch space in a recoverable location and never hard-delete it."""
    root = _recovery_root()
    root.mkdir(parents=True, exist_ok=True)
    directory = Path(
        tempfile.mkdtemp(prefix=f"{prefix}{int(time.time())}-", dir=root)
    )
    yield directory


def move_to_recovery(path: Path) -> None:
    """Move an unexpected leftover into recoverable storage."""
    if not path.exists():
        return
    root = _recovery_root()
    root.mkdir(parents=True, exist_ok=True)
    destination = root / f"{path.name}.{time.time_ns()}"
    shutil.move(os.fspath(path), os.fspath(destination))


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def canonicalize_docx(path: Path) -> None:
    """Normalize ZIP metadata so stage/promote regeneration has one hash."""
    with tempfile.NamedTemporaryFile(
        prefix=f".{path.stem}-",
        suffix=".docx",
        dir=path.parent,
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(path) as source, zipfile.ZipFile(
            temporary,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target:
            for name in sorted(source.namelist()):
                original = source.getinfo(name)
                info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = original.external_attr
                info.create_system = original.create_system
                target.writestr(info, source.read(name))
        temporary.replace(path)
    finally:
        if temporary.exists():
            move_to_recovery(temporary)


def load_manifest(path: Path = MANIFEST_PATH) -> dict[str, Any]:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise ValueError("suite-manifest.json 必须是 JSON 对象")
    return value


def release_spec(manifest: dict[str, Any]) -> dict[str, Any]:
    release = manifest.get("release")
    companion = manifest.get("release_companions")
    if not isinstance(release, dict) or not isinstance(companion, dict):
        raise ValueError("套件清单缺少 release 或 release_companions")
    required_release = (
        "tag",
        "version",
        "summary",
        "changes",
        "installation",
        "compatibility",
        "rollback",
    )
    missing = [key for key in required_release if not release.get(key)]
    if missing:
        raise ValueError("release 缺少伴随物事实：" + "、".join(missing))
    skills = manifest.get("skills")
    if not isinstance(skills, list) or not skills:
        raise ValueError("skills 必须是非空数组")
    return {
        "product_name": str(manifest.get("product_name") or ""),
        "tag": str(release["tag"]),
        "version": str(release["version"]),
        "skill_count": len(skills),
        "summary": str(release["summary"]),
        "changes": [str(item) for item in release["changes"]],
        "installation": {
            str(key): str(value)
            for key, value in dict(release["installation"]).items()
        },
        "compatibility": [
            {
                "target": str(item["target"]),
                "status": str(item["status"]),
                "note": str(item["note"]),
            }
            for item in release["compatibility"]
        ],
        "rollback": [str(item) for item in release["rollback"]],
        "manual_template": str(companion["manual_template"]),
        "manual_filename": str(companion["manual_filename"]).format(
            tag=release["tag"],
            version=release["version"],
        ),
        "companion_filename": str(companion["companion_filename"]).format(
            tag=release["tag"],
            version=release["version"],
        ),
        "delivery_directory": str(companion["delivery_directory"]),
        "word_manual_only": bool(companion.get("word_manual_only")),
        "require_branding": bool(companion.get("require_branding")),
        "require_render_qa": bool(companion.get("require_render_qa")),
    }


def _all_paragraphs(document: Document) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_paragraph_text(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _insert_after(paragraph: Any, text: str) -> Any:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    created = paragraph._parent.add_paragraph()
    created._p.getparent().remove(created._p)
    new_element.addnext(created._p)
    created._p.getparent().remove(created._p)
    new_element.getparent().replace(new_element, created._p)
    created.add_run(text)
    return created


def _set_east_asia_font(document: Document, font_name: str) -> None:
    def apply_to_run_properties(run_properties: Any) -> None:
        fonts = run_properties.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            run_properties.insert(0, fonts)
        for theme_attribute in (
            "w:asciiTheme",
            "w:hAnsiTheme",
            "w:eastAsiaTheme",
            "w:cstheme",
        ):
            fonts.attrib.pop(qn(theme_attribute), None)
        fonts.set(qn("w:ascii"), font_name)
        fonts.set(qn("w:hAnsi"), font_name)
        fonts.set(qn("w:eastAsia"), font_name)
        fonts.set(qn("w:cs"), font_name)
        language = run_properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            run_properties.append(language)
        language.set(qn("w:eastAsia"), "zh-CN")

    for style in document.styles:
        if not hasattr(style._element, "get_or_add_rPr"):
            continue
        apply_to_run_properties(style._element.get_or_add_rPr())

    paragraphs = list(_all_paragraphs(document))
    for section in document.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            paragraphs.extend(part.paragraphs)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            apply_to_run_properties(run._element.get_or_add_rPr())


def _remove_existing_brand_watermarks(document: Document) -> None:
    """Remove template watermarks before the branding skill applies one size."""

    for part in document.part.package.parts:
        if not str(part.partname).startswith("/word/header"):
            continue
        root = getattr(part, "element", None)
        if root is None:
            continue
        for parent in root.iter():
            for child in list(parent):
                if child.tag != qn("w:r"):
                    continue
                if any(
                    "Jiaotang Centered Watermark" in str(value)
                    for node in child.iter()
                    for value in node.attrib.values()
                ):
                    parent.remove(child)




def release_lines(spec: dict[str, Any]) -> list[str]:
    def clean(value: str) -> str:
        return value.rstrip("。；; ")

    compatibility = "；".join(
        f"{item['target']}：{item['status']}，{clean(item['note'])}"
        for item in spec["compatibility"]
    ) + "。"
    manual_access = spec["installation"].get(
        "non_workbuddy_manual_mcp",
        spec["installation"].get("manual_fallback", ""),
    )
    return [
        f"本版本更新：{spec['summary']}",
        "能力变化：" + "；".join(clean(item) for item in spec["changes"]) + "。",
        f"通用安装：{spec['installation']['generic']}",
        f"WorkBuddy 安装：{spec['installation']['workbuddy']}",
        f"手工接入：{manual_access}",
        f"安装验收：{spec['installation']['verification']}",
        f"兼容性：{compatibility}",
        "回滚：" + "；".join(clean(item) for item in spec["rollback"]) + "。",
    ]


def update_manual(template: Path, output: Path, spec: dict[str, Any]) -> None:
    document = Document(template)
    paragraphs = list(_all_paragraphs(document))
    version_paragraph = next(
        (
            paragraph
            for paragraph in paragraphs
            if paragraph.text.strip().startswith("适用版本：")
        ),
        None,
    )
    if version_paragraph is None:
        raise ValueError("Word 手册缺少可识别的 Skills 版本段落")
    old_tag_match = re.search(r"V\d+(?:\.\d+){1,3}", version_paragraph.text)
    if old_tag_match is None:
        raise ValueError("Word 手册无法识别旧版本号")
    old_tag = old_tag_match.group(0)
    old_count_match = re.search(r"共\s*(\d+)\s*项", version_paragraph.text)
    old_count = old_count_match.group(1) if old_count_match else None

    for paragraph in paragraphs:
        text = paragraph.text
        updated = text.replace(old_tag, spec["tag"])
        if re.search(r"由\s*\d+\s*个[^。]*Skills", updated, flags=re.I):
            updated = re.sub(
                r"由\s*\d+\s*个",
                f"由{spec['skill_count']}个",
                updated,
                count=1,
            )
        if old_count is not None:
            updated = re.sub(
                rf"(?<!\d){re.escape(old_count)}\s*项",
                f"{spec['skill_count']} 项",
                updated,
            )
            updated = re.sub(
                rf"共\s*{re.escape(old_count)}\s*项",
                f"共 {spec['skill_count']} 项",
                updated,
            )
        if updated != text:
            _replace_paragraph_text(paragraph, updated)

    existing = {
        prefix: next(
            (
                paragraph
                for paragraph in document.paragraphs
                if paragraph.text.strip().startswith(prefix)
            ),
            None,
        )
        for prefix in RELEASE_PREFIXES
    }
    lines = release_lines(spec)
    if any(existing.values()):
        for line in lines:
            prefix = next(prefix for prefix in RELEASE_PREFIXES if line.startswith(prefix))
            paragraph = existing[prefix]
            if paragraph is None:
                raise ValueError("Word 手册的自动发布信息区块不完整")
            _replace_paragraph_text(paragraph, line)
    else:
        anchor = next(
            (
                paragraph
                for paragraph in document.paragraphs
                if f"由{spec['skill_count']}个" in paragraph.text
            ),
            version_paragraph,
        )
        for line in lines:
            anchor = _insert_after(anchor, line)

    _set_east_asia_font(document, "Hiragino Sans GB")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    canonicalize_docx(output)


def extracted_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in _all_paragraphs(document))


def validate_manual_content(path: Path, spec: dict[str, Any]) -> dict[str, Any]:
    text = extracted_text(path)
    required = [
        f"适用版本：{spec['tag']}",
        f"由{spec['skill_count']}个",
        spec["summary"],
        *spec["changes"],
        *spec["installation"].values(),
        *(item["target"] for item in spec["compatibility"]),
        *spec["rollback"],
    ]
    missing = [item for item in required if item.rstrip("。；; ") not in text]
    if missing:
        raise ValueError("Word 手册缺少清单事实：" + "、".join(missing))
    return {
        "status": "pass",
        "version": spec["tag"],
        "skill_count": spec["skill_count"],
        "required_facts": len(required),
        "sha256": sha256(path),
    }


def _branding_root() -> Path:
    configured = os.environ.get("JIAOTANG_BRANDING_ROOT")
    return (
        Path(configured).expanduser()
        if configured
        else Path.home() / ".agents" / "skills" / "jiaotang-branding"
    )


def apply_branding(path: Path) -> dict[str, Any]:
    root = _branding_root()
    watermark = root / "scripts" / "office_watermark.py"
    gate = root / "scripts" / "delivery_gate.py"
    if not watermark.is_file() or not gate.is_file():
        raise RuntimeError(f"焦糖品牌运行时不完整：{root}")
    subprocess.run(
        [sys.executable, os.fspath(watermark), os.fspath(path)],
        check=True,
    )
    canonicalize_docx(path)
    completed = subprocess.run(
        [sys.executable, os.fspath(gate), os.fspath(path)],
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(completed.stdout)


def render_qa(path: Path, output_dir: Path) -> dict[str, Any]:
    soffice = shutil.which("soffice")
    pdftoppm = shutil.which("pdftoppm")
    pdfinfo = shutil.which("pdfinfo")
    if not soffice or not pdftoppm or not pdfinfo:
        raise RuntimeError("逐页渲染需要 soffice、pdftoppm 和 pdfinfo")
    output_dir.mkdir(parents=True, exist_ok=True)
    with recoverable_workspace("manual-render-") as profile:
        render_environment = os.environ.copy()
        if sys.platform == "darwin":
            font_cache = profile / "fontconfig-cache"
            font_cache.mkdir(parents=True, exist_ok=True)
            fontconfig = profile / "fontconfig.conf"
            fontconfig.write_text(
                "\n".join(
                    [
                        '<?xml version="1.0"?>',
                        '<!DOCTYPE fontconfig SYSTEM "urn:fontconfig:fonts.dtd">',
                        "<fontconfig>",
                        "  <dir>/System/Library/Fonts</dir>",
                        "  <dir>/System/Library/Fonts/Supplemental</dir>",
                        "  <dir>/Library/Fonts</dir>",
                        f"  <dir>{Path.home() / 'Library/Fonts'}</dir>",
                        f"  <cachedir>{font_cache}</cachedir>",
                        "</fontconfig>",
                    ]
                )
                + "\n",
                encoding="utf-8",
            )
            render_environment["FONTCONFIG_FILE"] = os.fspath(fontconfig)
        subprocess.run(
            [
                soffice,
                f"-env:UserInstallation=file://{profile}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                os.fspath(output_dir),
                os.fspath(path),
            ],
            check=True,
            capture_output=True,
            text=True,
            env=render_environment,
        )
    pdf = output_dir / f"{path.stem}.pdf"
    if not pdf.is_file():
        raise RuntimeError("Word 手册未生成 PDF 渲染稿")
    subprocess.run(
        [pdftoppm, "-png", "-r", "150", os.fspath(pdf), os.fspath(output_dir / "page")],
        check=True,
        capture_output=True,
        text=True,
    )
    info = subprocess.run(
        [pdfinfo, os.fspath(pdf)],
        check=True,
        capture_output=True,
        text=True,
    ).stdout
    pages_match = re.search(r"^Pages:\s+(\d+)$", info, flags=re.M)
    pages = int(pages_match.group(1)) if pages_match else 0
    images = sorted(output_dir.glob("page-*.png"))
    if not pages or len(images) != pages:
        raise RuntimeError("Word 手册逐页 PNG 数量与 PDF 页数不一致")
    return {
        "status": "pass",
        "page_count": pages,
        "pdf_sha256": sha256(pdf),
        "pages": [
            {"file": image.name, "sha256": sha256(image)}
            for image in images
        ],
    }


def write_companion(
    output: Path,
    manual: Path,
    spec: dict[str, Any],
    content_audit: dict[str, Any],
    branding_audit: dict[str, Any] | None,
    render_audit: dict[str, Any] | None,
) -> dict[str, Any]:
    def portable_audit(value: Any) -> Any:
        if isinstance(value, dict):
            result = {}
            for key, item in value.items():
                if key == "pdf_sha256":
                    continue
                result[key] = (
                    Path(str(item)).name
                    if key == "path" and item
                    else portable_audit(item)
                )
            return result
        if isinstance(value, list):
            return [portable_audit(item) for item in value]
        return value

    payload = {
        "schema_version": 1,
        "product_name": spec["product_name"],
        "release_tag": spec["tag"],
        "release_version": spec["version"],
        "skill_count": spec["skill_count"],
        "summary": spec["summary"],
        "changes": spec["changes"],
        "installation": spec["installation"],
        "compatibility": spec["compatibility"],
        "rollback": spec["rollback"],
        "manual": {
            "file": manual.name,
            "sha256": sha256(manual),
            "content_audit": content_audit,
            "branding_audit": portable_audit(branding_audit),
            "render_audit": portable_audit(render_audit),
        },
        "source_of_truth": "skills/suite-manifest.json",
    }
    output.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return payload


def generate(
    root: Path,
    output_dir: Path,
    *,
    apply_brand: bool,
    render: bool,
) -> dict[str, Any]:
    manifest = load_manifest(root / "skills" / "suite-manifest.json")
    spec = release_spec(manifest)
    template = root / spec["manual_template"]
    manual = output_dir / spec["manual_filename"]
    companion = output_dir / spec["companion_filename"]
    update_manual(template, manual, spec)
    content_audit = validate_manual_content(manual, spec)
    branding_audit = apply_branding(manual) if apply_brand else None
    render_audit = (
        render_qa(manual, output_dir / f"{spec['tag']}-render")
        if render
        else None
    )
    payload = write_companion(
        companion,
        manual,
        spec,
        content_audit,
        branding_audit,
        render_audit,
    )
    return {
        "status": "pass",
        "manual": str(manual),
        "companion": str(companion),
        "payload": payload,
    }


def deliver(
    root: Path,
    generated_dir: Path,
    delivery_dir: Path | None = None,
) -> dict[str, Any]:
    spec = release_spec(load_manifest(root / "skills" / "suite-manifest.json"))
    destination = (
        delivery_dir
        if delivery_dir is not None
        else Path.home() / spec["delivery_directory"]
    )
    destination.mkdir(parents=True, exist_ok=True)
    copied: list[str] = []
    for filename in (spec["manual_filename"], spec["companion_filename"]):
        source = generated_dir / filename
        if not source.is_file():
            raise FileNotFoundError(source)
        target = destination / filename
        shutil.copy2(source, target)
        copied.append(str(target))
    return {"status": "pass", "delivery_directory": str(destination), "files": copied}


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--root", type=Path, default=ROOT)
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument(
        "--check",
        action="store_true",
        help="只验证清单和Word模板可生成，不执行水印、渲染或交付。",
    )
    parser.add_argument("--deliver", action="store_true")
    parser.add_argument("--delivery-dir", type=Path)
    parser.add_argument("--skip-branding", action="store_true")
    parser.add_argument("--skip-render", action="store_true")
    args = parser.parse_args()
    try:
        if args.check:
            with recoverable_workspace("release-companion-contract-") as directory:
                result = generate(
                    args.root.resolve(),
                    directory,
                    apply_brand=False,
                    render=False,
                )
                result = {
                    "status": result["status"],
                    "release_tag": result["payload"]["release_tag"],
                    "skill_count": result["payload"]["skill_count"],
                    "source_of_truth": result["payload"]["source_of_truth"],
                }
        else:
            if args.output_dir is None:
                parser.error("--output-dir is required unless --check is used")
            result = generate(
                args.root.resolve(),
                args.output_dir.resolve(),
                apply_brand=not args.skip_branding,
                render=not args.skip_render,
            )
            if args.deliver:
                result["delivery"] = deliver(
                    args.root.resolve(),
                    args.output_dir.resolve(),
                    args.delivery_dir.resolve() if args.delivery_dir else None,
                )
    except Exception as exc:
        print(json.dumps({"status": "blocked", "error": str(exc)}, ensure_ascii=False))
        return 2
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
