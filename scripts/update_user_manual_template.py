#!/usr/bin/env python3
"""Keep the Word manual aligned with the simplified WorkBuddy install flow."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

from release_companions import canonicalize_docx


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MANUAL = ROOT / "docs/user-guide/企业全生命周期助手用户使用手册.docx"

FORBIDDEN_LEGACY_TEXT = (
    "jiaotang_kb_setup",
    "bootstrap_url",
    "设备私钥",
    "设备密钥",
    "更换绑定设备",
    "已绑定设备通过门户升级事务",
    "签名插件根目录 .mcp.json",
)


def _paragraph(document: Document, exact_text: str):
    matches = [item for item in document.paragraphs if item.text.strip() == exact_text]
    if len(matches) != 1:
        raise RuntimeError(f"Word 手册段落定位异常：{exact_text}，命中 {len(matches)} 项")
    return matches[0]


def _replace_section(
    document: Document,
    start_heading: str,
    end_heading: str,
    entries: tuple[tuple[str, str], ...],
) -> None:
    start = _paragraph(document, start_heading)
    end = _paragraph(document, end_heading)
    nodes = list(document._body._body)
    start_index = nodes.index(start._p)
    end_index = nodes.index(end._p)
    if start_index >= end_index:
        raise RuntimeError(f"Word 手册章节顺序异常：{start_heading} / {end_heading}")
    for node in nodes[start_index + 1 : end_index]:
        document._body._body.remove(node)
    for style, text in entries:
        end.insert_paragraph_before(text, style=style)


def update(path: Path) -> None:
    document = Document(path)
    package_boundary_text = (
        "可安装的 WorkBuddy 正式包只包含插件清单、49 项 Skills、最小行为 Hook、必要参考资料和业务脚本，"
        "不包含本地 MCP 服务、启动器、个人 Token 或插件根目录 .mcp.json。安装后必须通过真实 "
        "tools/list 枚举和 knowledge_service_status 调用；只有 49 项 Skills 可识别、三个核心知识工具齐全、"
        "状态返回 connected: true 且其他 MCP 未被覆盖，才算完成。"
    )
    package_boundary_matches = [
        item
        for item in document.paragraphs
        if item.text.strip().startswith("可安装的 WorkBuddy 正式包")
    ]
    if len(package_boundary_matches) != 1:
        raise RuntimeError(
            "Word 手册正式包边界段落定位异常："
            f"命中 {len(package_boundary_matches)} 项"
        )
    package_boundary_matches[0].text = package_boundary_text
    _replace_section(
        document,
        "连接团队知识库",
        "跨版本升级",
        (
            (
                "Normal",
                "WorkBuddy 使用上述一键安装指令同时完成 Skills 安装和远程 MCP 连接，不需要分开执行安装与绑定。",
            ),
            (
                "Normal",
                "其他支持 Streamable HTTP MCP 的 Agent 可进入“手工配置 MCP”页面。网站会自动复用当前账号的有效个人 Token；没有有效 Token 时才生成新值，并直接填入完整 MCP JSON。用户点击“复制完整配置”即可，不再到其他页面查找 Token。",
            ),
            (
                "Normal",
                "Token 只在当前登录用户的手工配置页和一键安装指令中出现，不写入公共插件包或公共代码。页面使用 Cache-Control: private, no-store；服务端普通日志和 Agent 最终回复不得记录或复述完整 Token。Token 会明文保存在当前用户的 WorkBuddy MCP 配置中；怀疑泄露时撤销旧 Token，下次打开配置页自动生成新值。",
            ),
        ),
    )
    _replace_section(
        document,
        "跨版本升级",
        "个人偏好与跨设备同步",
        (
            ("Normal", "1. 已安装旧版的用户不要手工删除插件目录。"),
            (
                "Normal",
                "2. 仍使用门户的同一段一键安装指令。WorkBuddy 会更新或替换旧插件，移除旧本地 jiaotang-kb 连接方式，并且不再读取旧的系统凭据。",
            ),
            (
                "Normal",
                "3. 升级只替换 mcpServers.jiaotang-kb，保留其他 MCP 配置，并只重载一次。",
            ),
            (
                "Normal",
                "4. 升级前备份插件目录和用户 MCP 配置。任一步失败时先保全现场，再恢复升级前备份，不覆盖个人项目资料、其他 MCP 条目或个人偏好。",
            ),
        ),
    )
    replacements = {
        "WorkBuddy 已绑定设备通过门户升级事务复用原设备身份和系统凭据；验签或只读工具调用失败时恢复升级前插件目录。": (
            "WorkBuddy 旧版用户仍使用同一段一键安装指令覆盖升级；升级前保留插件目录和用户 MCP 配置备份，失败时恢复。"
        ),
        "网站保存不可变发布记录和升级结果；访问令牌、设备私钥及一次性引导地址不进入普通日志。": (
            "网站保存不可变发布记录；个人 Token 不进入公共包、公共代码或普通日志，怀疑泄露时在门户撤销并重新生成。"
        ),
    }
    for old, new in replacements.items():
        old_matches = [
            item for item in document.paragraphs if item.text.strip() == old
        ]
        new_matches = [
            item for item in document.paragraphs if item.text.strip() == new
        ]
        if old_matches:
            if len(old_matches) != 1:
                raise RuntimeError(f"Word 手册旧段落定位异常：{old}")
            old_matches[0].text = new
        elif len(new_matches) != 1:
            raise RuntimeError(f"Word 手册更新段落定位异常：{new}")
    document.save(path)
    canonicalize_docx(path)


def check(path: Path) -> None:
    text = "\n".join(item.text for item in Document(path).paragraphs)
    hits = [item for item in FORBIDDEN_LEGACY_TEXT if item in text]
    if hits:
        raise RuntimeError("Word 手册仍包含旧安装机制：" + "、".join(hits))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manual", type=Path, default=DEFAULT_MANUAL)
    parser.add_argument("--check", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    path = args.manual.resolve()
    if not args.check:
        update(path)
    check(path)
    print(f"Word 手册简化安装边界通过：{path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
