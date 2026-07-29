#!/usr/bin/env python3
"""Rebuild the concise Skills Manager 0.2.0 Word user manual."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[3]
OUTPUT = ROOT / "docs" / "user-guide" / "企业全生命周期助手用户使用手册.docx"

# LibreOffice on macOS does not consistently honor the DOCX eastAsia fallback
# when a run's primary font lacks Chinese glyphs. A single installed Unicode
# family keeps both the source DOCX and the canonical PDF render deterministic.
LATIN_FONT = "Arial Unicode MS"
CJK_FONT = "Arial Unicode MS"
CODE_FONT = "Arial Unicode MS"

NAVY = "132B3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "66717C"
GOLD = "B98532"
INK = "17212B"
WHITE = "FFFFFF"
GRID = "D8E0E7"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LIGHT_GOLD = "FFF8E8"
LIGHT_RED = "FFF1F0"
RED = "9B1C1C"
GREEN = "315C45"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def hex_color(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width(parent, tag: str, width_dxa: int) -> None:
    width = ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def set_run_font(
    run,
    *,
    latin: str = LATIN_FONT,
    cjk: str = CJK_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr()
    fonts = ensure_child(run._element.rPr, "w:rFonts")
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), cjk)
    fonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = hex_color(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = hex_color(color)
    style.font.bold = bold
    style._element.get_or_add_rPr()
    fonts = ensure_child(style._element.rPr, "w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    fonts.set(qn("w:cs"), LATIN_FONT)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.05
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.08
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.1
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True


def set_cell_margins(cell, margins: dict[str, int] = CELL_MARGINS_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = ensure_child(tc_pr, "w:tcMar")
    for side in ("top", "bottom", "start", "end"):
        margin = ensure_child(tc_mar, f"w:{side}")
        margin.set(qn("w:w"), str(margins[side]))
        margin.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    set_width(tbl_pr, "w:tblW", CONTENT_WIDTH_DXA)
    indent = ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        row.height = None
        cant_split = ensure_child(row._tr.get_or_add_trPr(), "w:cantSplit")
        cant_split.set(qn("w:val"), "1")
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[index])
            set_cell_margins(cell)


def set_table_borders(table, *, color: str = GRID, size: int = 6) -> None:
    borders = ensure_child(table._tbl.tblPr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = ensure_child(borders, f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_fill(cell, color: str) -> None:
    shading = ensure_child(cell._tc.get_or_add_tcPr(), "w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), color)


def set_repeat_header(row) -> None:
    header = ensure_child(row._tr.get_or_add_trPr(), "w:tblHeader")
    header.set(qn("w:val"), "true")


def set_paragraph_spacing(paragraph, *, after: float = 0, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.widow_control = True


def add_field(paragraph, code: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {code} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_running_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, after=0, line=1)
    run = paragraph.add_run("企业全生命周期助手  ·  Skills 管理器 0.2.0")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=0, line=1)
    run = paragraph.add_run("第 ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_running_header_footer(section)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def create_numbering(document: Document, *, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(value)
        for item in numbering.findall(qn("w:abstractNum"))
        if (value := item.get(qn("w:abstractNumId"))) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    marker = "%1." if kind == "decimal" else ("☐" if kind == "check" else "•")
    level_text.set(qn("w:val"), marker)
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num_ids = [
        int(value)
        for item in numbering.findall(qn("w:num"))
        if (value := item.get(qn("w:numId"))) is not None
    ]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = ensure_child(p_pr, "w:numPr")
    level = ensure_child(num_pr, "w:ilvl")
    level.set(qn("w:val"), "0")
    number = ensure_child(num_pr, "w:numId")
    number.set(qn("w:val"), str(num_id))


def add_numbered_steps(document: Document, items: list[tuple[str, str]]) -> None:
    num_id = create_numbering(document, kind="decimal")
    for label, detail in items:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.left_indent = Twips(540)
        paragraph.paragraph_format.first_line_indent = Twips(-270)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(label)
        set_run_font(run, size=11, color=INK, bold=True)
        run = paragraph.add_run(detail)
        set_run_font(run, size=11, color=INK)


def add_bullets(document: Document, items: list[str], *, check: bool = False) -> None:
    num_id = create_numbering(document, kind="check" if check else "bullet")
    for text in items:
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.left_indent = Twips(540)
        paragraph.paragraph_format.first_line_indent = Twips(-270)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(text)
        set_run_font(run, size=10.8 if check else 11, color=INK)


def add_heading(
    document: Document,
    text: str,
    level: int,
    *,
    page_break_before: bool = False,
    libreoffice_section_anchor: bool = True,
) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    if page_break_before:
        paragraph.paragraph_format.page_break_before = True
    elif libreoffice_section_anchor:
        # LibreOffice needs an explicit false value on the first heading after
        # a section break to retain that section's page margins and running
        # header. A heading that naturally straddles a page must opt out.
        paragraph.paragraph_format.page_break_before = False
    for run in paragraph.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}[level],
            color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
            bold=True,
        )


def add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead:
        run = paragraph.add_run(bold_lead)
        set_run_font(run, size=11, color=INK, bold=True)
        text = text.removeprefix(bold_lead)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_callout(
    document: Document,
    title: str,
    body: str,
    *,
    fill: str = LIGHT_GRAY,
    title_color: str = DARK_BLUE,
) -> None:
    table = document.add_table(rows=1, cols=1)
    apply_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=fill, size=4)
    set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=2, line=1.12)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.7, color=title_color, bold=True)
    paragraph = cell.add_paragraph()
    set_paragraph_spacing(paragraph, after=0, line=1.18)
    run = paragraph.add_run(body)
    set_run_font(run, size=10.2, color=INK)
def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    font_size: float = 9.7,
    centered_columns: set[int] | None = None,
) -> None:
    centered_columns = centered_columns or set()
    table = document.add_table(rows=1, cols=len(headers))
    for row in rows:
        table.add_row()
    apply_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_header(table.rows[0])

    for column, value in enumerate(headers):
        cell = table.rows[0].cells[column]
        set_cell_fill(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if column in centered_columns
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        set_paragraph_spacing(paragraph, after=0, line=1.1)
        run = paragraph.add_run(value)
        set_run_font(run, size=font_size, color=NAVY, bold=True)

    for row_index, values in enumerate(rows, start=1):
        for column, value in enumerate(values):
            cell = table.rows[row_index].cells[column]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if column in centered_columns
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_spacing(paragraph, after=0, line=1.12)
            run = paragraph.add_run(value)
            set_run_font(run, size=font_size, color=INK)
def add_code_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.right_indent = Inches(0.22)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    shading = ensure_child(paragraph._p.get_or_add_pPr(), "w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    run = paragraph.add_run(text)
    set_run_font(run, latin=CODE_FONT, cjk=CJK_FONT, size=9.4, color=INK)


def start_new_page_section(document: Document) -> None:
    """Start a linked section on a clean page with the normal header/footer."""
    previous = document.sections[-1]
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = previous.page_width
    section.page_height = previous.page_height
    section.top_margin = previous.top_margin
    section.right_margin = previous.right_margin
    section.bottom_margin = previous.bottom_margin
    section.left_margin = previous.left_margin
    section.header_distance = previous.header_distance
    section.footer_distance = previous.footer_distance
    section.different_first_page_header_footer = False
    configure_running_header_footer(section)


def build_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(50)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, after=14, line=1)
    run = kicker.add_run("FOCUS GUIDE  /  0.2.0")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=8, line=1.02)
    run = title.add_run("企业全生命周期助手")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=8, line=1.08)
    run = subtitle.add_run("Skills 管理器 0.2.0 用户手册")
    set_run_font(run, size=18, color=BLUE, bold=True)

    descriptor = document.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(descriptor, after=30, line=1.15)
    run = descriptor.add_run("macOS 与 Windows · Agent 扫描 · 一键导入 · 可恢复更新")
    set_run_font(run, size=11, color=MUTED)

    add_callout(
        document,
        "先记住三件事",
        "管理器版本与 Skills 内容版本相互独立。当前管理器为 0.2.0，正式 Skills 为 V1.3.1.3，共 49 项。"
        "桌面客户端采用未签名本机授权分发；应用身份与 Skills 内容验真是两条独立信任链。"
        "系统或企业策略不允许未签名应用时，继续使用 HTTPS PWA，不关闭系统防护。",
        fill=LIGHT_GOLD,
        title_color=GOLD,
    )

    flow = document.add_paragraph()
    flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow.paragraph_format.space_before = Pt(18)
    flow.paragraph_format.space_after = Pt(10)
    flow.paragraph_format.line_spacing = 1.15
    run = flow.add_run("下载  ·  校验  ·  本机授权  ·  扫描  ·  导入  ·  验收")
    set_run_font(run, size=11.5, color=NAVY, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(meta, after=0, line=1.15)
    run = meta.add_run("适用：macOS arm64 / x64，Windows x64  ·  更新日期：2026 年 7 月")
    set_run_font(run, size=9.5, color=MUTED)
    start_new_page_section(document)


def build_install_page(document: Document) -> None:
    add_heading(document, "1  安装桌面客户端", 1)
    add_callout(
        document,
        "本地授权版的含义",
        "安装包不含 Apple Developer ID、公证或 Windows Authenticode 签名。"
        "用户只能在确认下载来源、文件名、SHA-256 与发布审计后，按本机系统允许的入口继续。"
        "管理员权限只负责本机提权，不会把未知发布者变成可信发布者。",
        fill=LIGHT_RED,
        title_color=RED,
    )

    add_heading(document, "下载前先核对", 2)
    add_numbered_steps(
        document,
        [
            (
                "选择本机架构。 ",
                "Apple Silicon 选择 macOS arm64；Intel Mac 选择 macOS x64；Windows 选择 x64。",
            ),
            (
                "确认文件名。 ",
                "正式本地授权包名称必须包含 unsigned-local，版本必须为 0.2.0。",
            ),
            (
                "核对完整性。 ",
                "按发布页说明比对 SHA256SUMS.txt 与 release-audit.json；任一不一致立即停止。",
            ),
        ],
    )

    add_heading(document, "macOS 与 Windows 操作差异", 2)
    add_table(
        document,
        ["环节", "macOS", "Windows"],
        [
            [
                "首次打开",
                "双击 DMG，将应用放入“应用程序”。若被阻止，到“系统设置 → 隐私与安全性”选择“仍要打开”，再完成本机认证。",
                "运行 EXE。系统策略允许时，可在 SmartScreen 中选择“更多信息 → 仍要运行”。",
            ],
            [
                "不要做",
                "不要删除隔离属性，不要关闭 Gatekeeper，不要运行来源不明的终端命令。",
                "不要关闭 Smart App Control、杀毒软件或组织策略，不要把 UAC 提示当成发布者证明。",
            ],
            [
                "完全阻止时",
                "继续使用门户 HTTPS PWA，或等待取得 Developer ID 的签名版。",
                "继续使用门户 HTTPS PWA；企业设备应联系管理员确认策略，不尝试绕过。",
            ],
        ],
        [1300, 4030, 4030],
        font_size=9.25,
        centered_columns={0},
    )

    add_callout(
        document,
        "现有用户不受影响",
        "安装管理器 0.2.0 不会强制升级 49 项 Skills，不会改变已有 WorkBuddy 插件、设备绑定、MCP 凭据或旧下载地址。"
        "同一台 Mac 或 Windows 已绑定团队知识库时，正常情况下继续复用现有设备凭据。",
    )
    start_new_page_section(document)


def build_scan_page(document: Document) -> None:
    add_heading(document, "2  扫描 Agent 与一键导入", 1)
    add_callout(
        document,
        "扫描由用户主动触发",
        "只有点击“扫描本机 Agent”后才开始。扫描范围限于系统应用目录、常见用户安装目录和已配置命令位置；"
        "不读取项目正文、客户资料、聊天记录或用户文档。多个平台命中同一托管目录时只写入一次并登记对应平台。",
        fill=LIGHT_BLUE,
    )

    add_heading(document, "标准操作", 2)
    add_numbered_steps(
        document,
        [
            ("连接发布门户。 ", "使用已有设备凭据连接；管理员调试时才使用管理员令牌。"),
            ("点击扫描。 ", "总览或平台页均可重新扫描，结果会显示命中路径、耗时与适配器 revision。"),
            ("选择平台。 ", "已安装平台显示对应导入按钮；未安装平台仍保留在主流平台清单中。"),
            ("预览计划。 ", "管理器列出新增、替换和冲突。发现未登记的同名目录时自动阻断覆盖。"),
            ("确认导入。 ", "只向已确认的用户级 Skills 目录写入已验签内容，不请求修改未知平台数据库。"),
            ("回到 Agent 验收。 ", "新建会话，确认 49 项目录可见，并抽样调用一个文档或知识库能力。"),
        ],
    )

    add_heading(document, "六个平台的当前动作", 2)
    add_table(
        document,
        ["平台", "识别入口", "管理器动作", "用户最终动作"],
        [
            ["WorkBuddy", "应用安装位置", "准备并验签通用插件市场包", "在 WorkBuddy 应用内添加市场、安装并启用"],
            ["TRAE", "应用与用户目录", "向稳定用户级 Skills 目录一键导入", "新建会话并复验"],
            ["Kimi Code", "命令位置", "向 ~/.agents/skills 一键导入", "新建会话并复验"],
            ["通义灵码", "应用安装位置", "准备已验证通用包", "按平台官方界面导入"],
            ["Qoder", "应用安装位置", "准备已验证导入包", "按插件或项目级入口导入"],
            ["Cherry Studio", "应用安装位置", "准备已验证通用包", "在官方导入界面选择 ZIP 或文件夹"],
        ],
        [1450, 1800, 2800, 3310],
        font_size=8.9,
        centered_columns={0},
    )
    start_new_page_section(document)


def build_workbuddy_page(document: Document) -> None:
    add_heading(document, "3  WorkBuddy 与远程签名适配器", 1)

    add_heading(document, "WorkBuddy 使用同一个跨平台包", 2)
    add_body(
        document,
        "macOS 与 Windows 不再分别发布 WorkBuddy 安装包，也不再运行固定 .command、.cmd、PowerShell 或外部 CLI。"
        "管理器只下载、验签并定位同一个跨平台插件市场 ZIP。",
    )
    add_numbered_steps(
        document,
        [
            ("下载并验签。 ", "在平台页点击“准备安装”，等待签名与逐文件哈希通过。"),
            ("定位并解压。 ", "打开 ZIP 位置，解压到可长期保留的本地目录。"),
            ("添加本地市场。 ", "在 WorkBuddy 插件市场添加解压后的 jiaotang 目录。"),
            ("安装并启用。 ", "安装 jiaotang-workbuddy-skills@jiaotang；不关闭 WorkBuddy，不运行包外脚本。"),
            ("填写 bootstrap。 ", "门户 API 页复制一次性地址并粘贴到敏感配置；同机已绑定时复用现有凭据。"),
            ("完成验收。 ", "确认插件启用、49 项可见，并完成 MCP 工具列表与一次知识检索。"),
        ],
    )
    add_code_line(
        document,
        "/plugin marketplace add <解压后的 jiaotang 目录>    →    安装 jiaotang-workbuddy-skills@jiaotang",
    )

    add_heading(document, "远程适配器只更新数据，不下发命令", 2)
    add_numbered_steps(
        document,
        [
            ("固定来源。 ", "只从焦糖门户读取平台清单、签名和签名元数据。"),
            ("固定发布者。 ", "内置 Ed25519 公钥和指纹验签，namespace 与 identity 必须匹配。"),
            ("严格 schema。 ", "只允许平台、能力、路径、说明和固定导入模式；命令、脚本、URL 与未知字段一律拒绝。"),
            ("版本保护。 ", "新 revision 验证通过才缓存；失败时继续使用当前已验证版本或内置版本。"),
        ],
    )
    add_callout(
        document,
        "什么时候必须升级客户端",
        "平台只更改应用位置或用户级 Skills 目录时发布签名适配器；涉及新可执行逻辑、私有 API 或扩大写入权限时，"
        "必须升级管理器并重走双端发行门禁。",
        fill=LIGHT_GOLD,
        title_color=GOLD,
    )
    start_new_page_section(document)


def build_security_page(document: Document) -> None:
    add_heading(document, "4  安全、审计、回滚与验收", 1)

    add_heading(document, "两条信任链要分开看", 2)
    add_table(
        document,
        ["对象", "门禁", "结论"],
        [
            ["桌面应用", "未签名本机授权；发布页提供版本、SHA-256 与发行审计", "确认下载来源后，由用户按本机策略决定是否运行"],
            ["Skills 下载", "HTTPS 同源、发布通道 SHA-256", "整体文件被替换时停止"],
            ["发布者", "固定 Ed25519 公钥、指纹、identity 与 namespace", "签名不匹配时停止"],
            ["包内内容", "安全路径、解压上限、签名清单与逐文件 SHA-256", "缺失、夹带或哈希不符时停止"],
            ["平台适配器", "签名、schema、兼容版本和 revision 保护", "失败时不执行远程内容"],
        ],
        [1750, 3500, 4110],
        font_size=9.0,
        centered_columns={0},
    )

    add_heading(document, "查看审计与执行回滚", 2)
    add_bullets(
        document,
        [
            "在“安全中心”点击“打开审计日志”。日志记录扫描、适配器更新、包验签、导入与回滚的开始、结果和失败原因。",
            "访问令牌、bootstrap、密码、设备私钥和其他敏感字段只记录为脱敏占位，不写入普通日志。",
            "每次更新已托管目录前，旧内容进入同盘 backups；回滚前，当前版本先进入 displaced 恢复区。",
            "未知同名目录默认视为冲突并阻断，不自动覆盖个人项目、第三方 Skills 或平台私有配置。",
        ],
    )
    add_heading(
        document,
        "发布后验收清单",
        2,
        page_break_before=True,
        libreoffice_section_anchor=False,
    )
    add_bullets(
        document,
        [
            "管理器“关于”或安全中心显示版本 0.2.0，Skills 清单显示 V1.3.1.3 与 49 项。",
            "点击扫描后，平台数量、命中路径、耗时与适配器 revision 均有结果。",
            "TRAE 或 Kimi Code 的导入计划没有未处理冲突；导入后新建会话能识别 Skills。",
            "WorkBuddy 使用通用跨平台包，并在应用内完成市场添加、安装、启用与 bootstrap 配置。",
            "审计日志存在本轮扫描、验签和导入事件；托管目录存在可恢复的回滚点。",
            "系统或企业策略阻止未签名应用时，停止原生安装并改用 HTTPS PWA。",
        ],
        check=True,
    )

    add_callout(
        document,
        "遇到问题先这样处理",
        "未发现平台：确认应用已安装后重新扫描，或改走平台官方导入界面。"
        "出现目录冲突：不要强制覆盖，先确认目录归属并备份。"
        "签名或 SHA-256 不一致：立即停止，不重新下载同名文件反复尝试。"
        "回滚记录缺失：不要手工删除当前目录，先从审计日志确认目标路径和上一次成功操作。",
        fill=LIGHT_GOLD,
        title_color=GOLD,
    )


def build_document() -> Document:
    document = Document()
    configure_styles(document)
    configure_page(document)

    properties = document.core_properties
    properties.title = "企业全生命周期助手 Skills 管理器 0.2.0 用户手册"
    properties.subject = "macOS 与 Windows 本地授权安装、Agent 扫描、一键导入与安全验收"
    properties.author = "焦糖"
    properties.last_modified_by = "焦糖"
    properties.keywords = "Skills 管理器, Agent, WorkBuddy, macOS, Windows, 用户手册"
    properties.comments = "正式用户手册，适用管理器 0.2.0 与 Skills V1.3.1.3"

    build_cover(document)
    build_install_page(document)
    build_scan_page(document)
    build_workbuddy_page(document)
    build_security_page(document)
    return document


def main() -> None:
    document = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
